"""Document generation: Italian civil procedure opposition act (atto di opposizione a decreto ingiuntivo)."""

from __future__ import annotations

import io
import json
import logging
import os
import re
from typing import Any, Dict, List, Optional

from fastapi import HTTPException
from langchain_core.messages import HumanMessage, SystemMessage

from .ai_chat import _call_chat

try:
    from pypdf import PdfReader as _PdfReader
    _PDF_SUPPORT = True
except ImportError:
    _PdfReader = None
    _PDF_SUPPORT = False

logger = logging.getLogger(__name__)

DOCUMENT_TYPE_REGISTRY: dict = {}


def _placeholder(lang: str) -> str:
    if lang == "es":
        return "[COMPLETAR]"
    if lang == "en":
        return "[TO BE COMPLETED]"
    return "[DA COMPILARE]"

_GENERATION_COMBOS = [
    (
        {
            "genera", "generami", "scrivi", "scrivimi", "redigi", "redigimi",
            "crea", "creami", "prepara", "preparami", "drafta",
            "elabora", "elaborami", "formulami", "formula",
        },
        {"opposizione", "atto", "ricorso", "memoria", "decreto", "ingiuntivo"},
    ),
]


def is_generation_request(message: str) -> bool:
    msg = message.lower()
    action_verbs = [
        "genera", "generami", "scrivi", "scrivimi", "redigi", "redigimi",
        "crea", "creami", "prepara", "preparami", "stendi", "stendimi",
        "elabora", "elaborami", "formula", "formulami", "produce", "producimi",
        "fammi", "draft", "write", "create", "redacta",
        "voglio", "vorrei", "ho bisogno di", "necesito", "quiero",
    ]
    strong_triggers = [
        # Italian with mi
        "redigimi", "generami", "scrivimi", "preparami", "creami", "fammi",
        "elaborami", "stendimi", "formulami", "producimi", "drafta",
        # Italian without mi — "un/una" makes them unambiguous
        "scrivi un", "scrivi una", "redigi un", "redigi una",
        "genera un", "genera una", "crea un", "crea una",
        "prepara un", "prepara una", "stendi un", "stendi una",
        # Italian without article — verb + specific document type word
        "scrivi atto", "scrivi contratto", "scrivi lettera",
        "scrivi documento", "scrivi memoria", "scrivi ricorso",
        "scrivi istanza", "scrivi verbale", "scrivi diffida",
        "scrivi procura", "scrivi appello", "scrivi dichiarazione",
        "redigi atto", "redigi contratto", "redigi lettera",
        "redigi documento", "redigi ricorso", "redigi memoria",
        "genera atto", "genera documento", "genera contratto",
        "generami documento", "generami atto", "generami contratto",
        # Common phrases
        "fammi un atto", "fammi un contratto", "fammi una lettera",
        "fammi un documento", "fammi una memoria", "fammi un verbale",
        "ho bisogno di un contratto", "ho bisogno di una lettera",
        "ho bisogno di un atto", "ho bisogno di una memoria",
        "voglio un contratto", "voglio una lettera", "voglio un atto",
        "vorrei un contratto", "vorrei una lettera", "vorrei un atto",
        # Spanish
        "redacta un", "redacta una",
        "necesito un contrato", "necesito una carta",
        # English
        "draft a", "draft an", "write a", "write an",
        "create a", "create an", "generate a", "generate an",
        "i need a contract", "i need a letter",
        # Unambiguous Italian situation phrases
        "mi hanno licenziato",
        # Noun-form document requests — user names the document type directly
        # without an action verb (e.g. "Nomina del difensore di fiducia per...")
        "nomina del difensore",
        "nomina difensore",
        "procura alle liti",
        "atto di citazione",
        "atto di appello",
        "ricorso per",
        "istanza di",
        "memoria difensiva",
        "memorie difensive",
        "contratto di locazione",
        "contratto di compravendita",
        "contratto di lavoro",
        "lettera di licenziamento",
        "diffida ad adempiere",
        "messa in mora",
        "verbale di assemblea",
        "dichiarazione sostitutiva",
        "atto di opposizione",
        "querela contro",
        "denuncia contro",
        "rinuncia al mandato",
        "revoca della procura",
    ]
    if any(t in msg for t in strong_triggers):
        return True
    # Weak path: action verb + any significant word from the catalog
    # (DOCUMENT_TYPE_REGISTRY is now empty — use catalog labels instead)
    _catalog_words = frozenset(
        word
        for entry in SYSTEM_TEMPLATES_CATALOG
        for phrase in [entry.get("tipo_atto", ""), entry.get("label", "")]
        for word in phrase.lower().split()
        if len(word) > 5
    )
    has_verb = any(v in msg for v in action_verbs)
    has_keyword = any(k in msg for k in _catalog_words)
    return has_verb and has_keyword


def classify_document_type(message: str, lang: str) -> str:
    # DOCUMENT_TYPE_REGISTRY is intentionally empty — the legacy 20-type
    # registry has been retired. Skip the LLM call entirely and return
    # "unknown" so the caller falls through to classify_system_template.
    if not DOCUMENT_TYPE_REGISTRY:
        return "unknown"

    type_list = "\n".join(
        f"- {key}: {entry['label']}"
        for key, entry in DOCUMENT_TYPE_REGISTRY.items()
    )
    rental_hints = (
        "Per distinguere i tipi di locazione:\n"
        "- cedolare secca o affitto privato semplice → rental_basic\n"
        "- studenti, universitari, fuori sede → rental_student\n"
        "- transitorio, temporaneo, esigenza transitoria → rental_transitional\n"
        "- canone libero, 4+4 → rental_free_rent\n"
        "- commerciale, ufficio, negozio, B2B → rental_commercial\n"
        "- locazione abitativa generica senza altri segnali → rental_standard\n"
        "\nPer distinguere atti difensivi e di opposizione:\n"
        "- memoria difensiva, risposta denuncia, atto difensivo → memoria_difensiva\n"
        "- opposizione decreto, decreto ingiuntivo, opporsi → opposition_act\n"
    )
    system = (
        "Sei un classificatore di richieste di documenti legali. "
        "Dato un messaggio utente, restituisci esattamente una delle seguenti chiavi "
        "che meglio corrisponde al tipo di documento richiesto:\n\n"
        f"{type_list}\n"
        "- unknown: nessuno dei precedenti\n\n"
        f"{rental_hints}\n"
        "Restituisci SOLO la chiave, senza testo aggiuntivo."
    )
    human = f"Messaggio utente:\n{message}"
    try:
        result = _call_chat(
            [SystemMessage(content=system), HumanMessage(content=human)],
            max_tokens=20,
        ).strip().lower()
        if result in DOCUMENT_TYPE_REGISTRY or result == "unknown":
            return result
        # fuzzy match — model may have added punctuation or minor variation
        for key in DOCUMENT_TYPE_REGISTRY:
            if key in result or result in key:
                return key
        logger.warning(f"classify_document_type got unexpected value: {result!r}, falling back to unknown")
        return "unknown"
    except Exception as e:
        logger.warning(f"classify_document_type failed: {e}, falling back to unknown")
        return "unknown"


def extract_document_fields(user_message: str, doc_type: str, lang: str) -> Dict[str, str]:
    entry = DOCUMENT_TYPE_REGISTRY[doc_type]
    fields = entry["fields"]
    label = entry["label"]
    fields_list = ", ".join(fields)

    if lang == "es":
        system = (
            f"Extrae del mensaje del usuario los valores para los siguientes campos del documento '{label}'.\n"
            f"Devuelve SOLO un objeto JSON válido, sin texto antes ni después, sin markdown.\n"
            f"Campos requeridos: {json.dumps(fields, ensure_ascii=False)}\n"
            "Reglas:\n"
            "- Si un valor está explícitamente mencionado en el mensaje, úsalo exactamente\n"
            "- Si un valor es claramente deducible del contexto, deducelo (ej: si la duración es 6 meses y la fecha de fin es 15 diciembre 2026, calcula la fecha de inicio como 15 junio 2026)\n"
            "- Si un valor no está presente ni es deducible, usa cadena vacía \"\"\n"
            "- No inventes datos no presentes ni deducibles\n"
            f"Responde SOLO con el JSON, ejemplo: {{\"campo1\": \"valor1\", \"campo2\": \"\"}}"
        )
    elif lang == "en":
        system = (
            f"Extract from the user message the values for the following fields of the document '{label}'.\n"
            f"Return ONLY a valid JSON object, with no text before or after, no markdown.\n"
            f"Required fields: {json.dumps(fields, ensure_ascii=False)}\n"
            "Rules:\n"
            "- If a value is explicitly mentioned in the message, use it exactly\n"
            "- If a value is clearly inferable from context, infer it (e.g. if duration is 6 months and end date is 15 December 2026, calculate start date as 15 June 2026)\n"
            "- If a value is neither present nor inferable, use empty string \"\"\n"
            "- Do not invent data that is not present or inferable\n"
            f"Reply ONLY with the JSON, example: {{\"field1\": \"value1\", \"field2\": \"\"}}"
        )
    else:
        system = (
            f"Estrai dal messaggio dell'utente i valori per i seguenti campi del documento '{label}'.\n"
            f"Restituisci SOLO un oggetto JSON valido, senza testo prima o dopo, senza markdown.\n"
            f"Campi richiesti: {json.dumps(fields, ensure_ascii=False)}\n"
            "Regole:\n"
            "- Se un valore è esplicitamente menzionato nel messaggio, usalo esattamente\n"
            "- Se un valore è chiaramente deducibile dal contesto, deducilo (es: se durata è 6 mesi e data fine è 15 dicembre 2026, calcola data inizio come 15 giugno 2026)\n"
            "- Se un valore non è presente né deducibile, usa stringa vuota \"\"\n"
            "- Non inventare dati non presenti o non deducibili\n"
            f"Rispondi SOLO con il JSON, esempio: {{\"campo1\": \"valore1\", \"campo2\": \"\"}}"
        )

    human = f"Messaggio dell'utente:\n{user_message}"
    raw = _call_chat(
        [SystemMessage(content=system), HumanMessage(content=human)],
        max_tokens=400,
    )

    text = re.sub(r"```(?:json)?\s*", "", raw).strip().rstrip("`").strip()
    try:
        parsed = json.loads(text)
    except (json.JSONDecodeError, ValueError):
        logger.warning("extract_document_fields: JSON parse failed, retrying for doc_type=%r, raw=%r", doc_type, raw[:200])
        retry_response = _call_chat(
            [
                SystemMessage(content="Rispondi SOLO con JSON valido, nessun testo aggiuntivo."),
                HumanMessage(content=f"Estrai questi campi: {fields}\nDal testo: {user_message}\nJSON:"),
            ],
            max_tokens=400,
        )
        clean_retry = re.sub(r"```(?:json)?\s*", "", retry_response).strip().rstrip("`").strip()
        try:
            parsed = json.loads(clean_retry)
        except (json.JSONDecodeError, ValueError):
            logger.error("extract_document_fields: retry also failed for doc_type=%r", doc_type)
            parsed = {}

    return {k: str(parsed.get(k, "") or "") for k in fields}


def extract_case_details(user_message: str) -> Dict[str, str]:
    return extract_document_fields(user_message, "opposition_act", "it")


def _field(value: Optional[str], label: str, lang: str = "it") -> str:
    """Return value if non-empty, else a labelled placeholder."""
    v = (value or "").strip()
    return v if v else f"{_placeholder(lang)} ({label})"


def _format_retrieved_sections(retrieved_sections: List[Any], lang: str = "it") -> str:
    """Format retrieved knowledge-base sections into a readable context block."""
    if not retrieved_sections:
        if lang == "es":
            return "(ninguna seccion recuperada de la base de conocimiento)"
        if lang == "en":
            return "(no sections retrieved from the knowledge base)"
        return "(nessuna sezione recuperata dalla knowledge base)"

    lines: List[str] = []
    for i, item in enumerate(retrieved_sections, 1):
        if isinstance(item, dict):
            title = item.get("title") or item.get("heading") or f"Sezione {i}"
            text = item.get("text") or item.get("text_en") or item.get("content") or ""
            source = item.get("document_title") or item.get("source") or ""
            entry = f"[{i}] {title}"
            if source:
                entry += f" (Fonte: {source})"
            if text:
                entry += f"\n{text[:600]}"
                if len(text) > 600:
                    entry += "..."
        else:
            entry = f"[{i}] {str(item)[:600]}"
        lines.append(entry)

    return "\n\n".join(lines)


def _opposition_system(lang: str) -> str:
    ph = _placeholder(lang)
    if lang == "es":
        return (
            "Eres un abogado experto en derecho civil espanol. "
            "Redacta un escrito de oposicion a decreto monitorio conforme al art. 815 de la Ley de Enjuiciamiento Civil. "
            "El documento debe ser formalmente correcto, en terminologia juridica espanola precisa, y estructurado "
            "en las siguientes secciones en el orden indicado:\n\n"
            "1. ENCABEZAMIENTO — Juzgado competente, partes (oponente y acreedor), numero y fecha del decreto monitorio.\n"
            "2. ANTECEDENTES DE HECHO — Exposicion sintetica de los hechos relevantes.\n"
            "3. MOTIVOS DE OPOSICION — Argumentos juridicos y facticos. "
            "Cita las secciones de la base de conocimiento proporcionadas en el contexto cuando sean pertinentes, "
            "indicando la fuente entre parentesis.\n"
            "4. SUPLICO — Peticion precisa: lo que el oponente solicita al Juzgado "
            "(anulacion/suspension del decreto, desestimacion de las pretensiones de la parte contraria, condena en costas).\n"
            "5. FIRMA Y FECHA — Lugar, fecha, firma del letrado.\n\n"
            f"Usa {ph} para cada dato ausente sin excepción — no inventes ni asumas fechas, nombres, códigos, números ni ningún otro dato no proporcionado explícitamente: usa {ph} para fechas, números de identificación, direcciones y referencias numéricas. "
            f"IMPORTANTE: usa EXCLUSIVAMENTE {ph} para los datos faltantes. "
            f"ESTÁ PROHIBIDO usar marcadores descriptivos como [NOMBRE], [DIRECCIÓN], [FECHA] o similares. "
            f"Solo {ph} y nada más para cualquier dato no proporcionado. "
            "No anadas advertencias meta-legales en el propio documento. "
            "Redacta directamente el escrito, sin preambulos ni explicaciones."
        )
    if lang == "en":
        return (
            "You are an expert civil litigation lawyer. "
            "Draft an opposition to an injunction/payment order under applicable civil procedure rules. "
            "The document must be formally correct, in precise legal language, and structured "
            "in the following sections in the order given:\n\n"
            "1. HEADING — Competent court, parties (opposing party and creditor), reference number and date of the order.\n"
            "2. STATEMENT OF FACTS — Concise factual background.\n"
            "3. GROUNDS OF OPPOSITION — Legal and factual arguments. "
            "Cite sections from the knowledge base provided in the context where relevant, "
            "indicating the source in parentheses.\n"
            "4. RELIEF SOUGHT — Precise petitum: what the opposing party requests from the court "
            "(revocation/suspension of the order, dismissal of the claimant's claims, costs award).\n"
            "5. SIGNATURE AND DATE — Place, date, lawyer's signature.\n\n"
            f"Use {ph} for every missing piece of information without exception — never invent or assume dates, names, codes, numbers or any other data not explicitly provided: use {ph} for dates, tax codes, addresses and reference numbers. "
            f"IMPORTANT: use EXCLUSIVELY {ph} for missing data. "
            f"It is FORBIDDEN to use descriptive placeholders such as [NAME], [ADDRESS], [DATE] or similar. "
            f"Only {ph} and nothing else for any data not provided. "
            "Do not add meta-legal disclaimers to the document itself. "
            "Write the document directly, without preamble or explanation."
        )
    # Default: Italian
    return (
        "Sei un avvocato esperto in diritto civile italiano. "
        "Redigi un atto di opposizione a decreto ingiuntivo ai sensi degli artt. 645 e ss. c.p.c. "
        "Il documento deve essere formalmente corretto, in italiano giuridico preciso, e strutturato "
        "nelle seguenti sezioni nell'ordine indicato:\n\n"
        "1. INTESTAZIONE — Tribunale competente, parti (opponente e opposto), "
        "numero e data del decreto ingiuntivo.\n"
        "2. PREMESSE IN FATTO — Ricostruzione sintetica dei fatti rilevanti.\n"
        "3. MOTIVI DI OPPOSIZIONE — Argomenti giuridici e fattuali. "
        "Cita le sezioni della knowledge base fornite nel contesto ove pertinenti, "
        "indicando la fonte tra parentesi.\n"
        "4. CONCLUSIONI — Petitum preciso: cosa chiede l'opponente al Tribunale "
        "(revoca/sospensione del decreto, rigetto delle domande avversarie, condanna alle spese).\n"
        "5. FIRMA E DATA — Luogo, data, firma del difensore.\n\n"
        f"Usa {ph} per ogni dato mancante senza eccezioni — non inventare né assumere date, nomi, codici, numeri o qualsiasi altro dato non esplicitamente fornito: usa {ph} per date, codici fiscali, indirizzi e riferimenti numerici. "
        f"IMPORTANTE: usa ESCLUSIVAMENTE {ph} per i dati mancanti. "
        f"È VIETATO usare placeholder descrittivi come [NOME], [INDIRIZZO], [DATA] o simili. "
        f"Solo {ph} e nient'altro per qualsiasi dato non fornito. "
        "Non aggiungere avvertenze meta-legali sul documento stesso. "
        "Scrivi direttamente l'atto, senza prefazioni o spiegazioni."
    )


def _memoria_difensiva_system(lang: str) -> str:
    ph = _placeholder(lang)
    if lang == "es":
        return (
            "Eres un abogado experto en defensa penal y administrativa. "
            "Redacta una memoria defensiva formal en nombre del investigado o denunciado. "
            "El documento debe ser formalmente correcto, en terminología jurídica precisa, y estructurado "
            "en las siguientes secciones en el orden indicado:\n\n"
            "1. ENCABEZAMIENTO — Autoridad o tribunal destinatario, datos del defendido y del denunciante, "
            "número de referencia del expediente o denuncia.\n"
            "2. ANTECEDENTES DE HECHO — Exposición de los hechos desde la perspectiva de la defensa.\n"
            "3. MOTIVOS DE DEFENSA — Argumentos jurídicos y fácticos a descargo. "
            "Cita las secciones de la base de conocimiento proporcionadas en el contexto cuando sean pertinentes, "
            "indicando la fuente entre paréntesis.\n"
            "4. SUPLICO — Petición precisa: sobreseimiento, archivo de las actuaciones o desestimación "
            "de los cargos, con reserva de todas las acciones legales pertinentes.\n"
            "5. FIRMA Y FECHA — Lugar, fecha, firma del letrado o del defendido.\n\n"
            f"Usa {ph} para cada dato ausente sin excepción — no inventes ni asumas fechas, nombres, códigos, números ni ningún otro dato no proporcionado explícitamente: usa {ph} para fechas, números de identificación, direcciones y referencias numéricas. "
            f"IMPORTANTE: usa EXCLUSIVAMENTE {ph} para los datos faltantes. "
            f"ESTÁ PROHIBIDO usar marcadores descriptivos como [NOMBRE], [DIRECCIÓN], [FECHA] o similares. "
            f"Solo {ph} y nada más para cualquier dato no proporcionado. "
            "No añadas advertencias meta-legales en el propio documento. "
            "Redacta directamente el escrito, sin preámbulos ni explicaciones."
        )
    if lang == "en":
        return (
            "You are an expert criminal and administrative defence lawyer. "
            "Draft a formal defence memorial on behalf of the accused or investigated party. "
            "The document must be formally correct, in precise legal language, and structured "
            "in the following sections in the order given:\n\n"
            "1. HEADING — Competent authority or court, details of the defendant and complainant, "
            "case or complaint reference number.\n"
            "2. STATEMENT OF FACTS — Account of events from the defence perspective.\n"
            "3. GROUNDS OF DEFENCE — Legal and factual arguments in favour of the defendant. "
            "Cite sections from the knowledge base provided in the context where relevant, "
            "indicating the source in parentheses.\n"
            "4. RELIEF SOUGHT — Precise petitum: dismissal of proceedings, closure of the case, "
            "or rejection of all charges, with reservation of all applicable legal remedies.\n"
            "5. SIGNATURE AND DATE — Place, date, lawyer's or defendant's signature.\n\n"
            f"Use {ph} for every missing piece of information without exception — never invent or assume dates, names, codes, numbers or any other data not explicitly provided: use {ph} for dates, tax codes, addresses and reference numbers. "
            f"IMPORTANT: use EXCLUSIVELY {ph} for missing data. "
            f"It is FORBIDDEN to use descriptive placeholders such as [NAME], [ADDRESS], [DATE] or similar. "
            f"Only {ph} and nothing else for any data not provided. "
            "Do not add meta-legal disclaimers to the document itself. "
            "Write the document directly, without preamble or explanation."
        )
    # Default: Italian
    return (
        "Sei un avvocato esperto in difesa penale e amministrativa. "
        "Redigi una memoria difensiva formale nell'interesse dell'indagato o del denunciato. "
        "Il documento deve essere formalmente corretto, in italiano giuridico preciso, e strutturato "
        "nelle seguenti sezioni nell'ordine indicato:\n\n"
        "1. INTESTAZIONE — Autorità o tribunale destinatario, dati del difeso e del denunciante, "
        "numero di riferimento del procedimento o della denuncia.\n"
        "2. PREMESSE IN FATTO — Ricostruzione dei fatti dal punto di vista della difesa.\n"
        "3. MOTIVI DI DIFESA — Argomenti giuridici e fattuali a discarico. "
        "Cita le sezioni della knowledge base fornite nel contesto ove pertinenti, "
        "indicando la fonte tra parentesi.\n"
        "4. CONCLUSIONI — Petitum preciso: proscioglimento, archiviazione del procedimento o rigetto "
        "di ogni addebito, con riserva di ogni azione legale opportuna.\n"
        "5. FIRMA E DATA — Luogo, data, firma del difensore o del difeso.\n\n"
        f"Usa {ph} per ogni dato mancante senza eccezioni — non inventare né assumere date, nomi, codici, numeri o qualsiasi altro dato non esplicitamente fornito: usa {ph} per date, codici fiscali, indirizzi e riferimenti numerici. "
        f"IMPORTANTE: usa ESCLUSIVAMENTE {ph} per i dati mancanti. "
        f"È VIETATO usare placeholder descrittivi come [NOME], [INDIRIZZO], [DATA] o simili. "
        f"Solo {ph} e nient'altro per qualsiasi dato non fornito. "
        "Non aggiungere avvertenze meta-legali sul documento stesso. "
        "Scrivi direttamente l'atto, senza prefazioni o spiegazioni."
    )


def _rental_basic_system(lang: str) -> str:
    ph = _placeholder(lang)
    if lang == "es":
        return (
            "Eres un abogado experto en arrendamientos urbanos españoles. "
            "Redacta un contrato de arrendamiento de vivienda con tributación simplificada. "
            "El documento debe ser formalmente correcto y estructurado en las siguientes secciones:\n\n"
            "1. ENCABEZAMIENTO — Arrendador, arrendatario, D.N.I./N.I.E., fecha y lugar de firma.\n"
            "2. OBJETO Y DURACIÓN — Descripción del inmueble, duración en años, fecha de inicio, "
            "prórroga tácita con preaviso de 3 meses.\n"
            "3. RENTA Y RÉGIMEN FISCAL — Importe mensual, día de pago, régimen fiscal aplicable, "
            "sin actualización por IPC salvo pacto expreso.\n"
            "4. OBLIGACIONES DEL ARRENDATARIO — Uso residencial, prohibición de subarriendo, "
            "mantenimiento ordinario, suministros a cargo del arrendatario.\n"
            "5. DEPÓSITO Y CLÁUSULAS FINALES — Depósito máximo 2 mensualidades, gastos contractuales, "
            "mediación obligatoria, firmas con doble suscripción de cláusulas generales.\n\n"
            f"Usa {ph} para cada dato ausente sin excepción — no inventes ni asumas fechas, nombres, códigos, números ni ningún otro dato no proporcionado explícitamente: usa {ph} para fechas, números de identificación, direcciones y referencias numéricas. "
            f"IMPORTANTE: usa EXCLUSIVAMENTE {ph} para los datos faltantes. "
            f"ESTÁ PROHIBIDO usar marcadores descriptivos como [NOMBRE], [DIRECCIÓN], [FECHA] o similares. "
            f"Solo {ph} y nada más para cualquier dato no proporcionado. "
            "No añadas advertencias meta-legales. "
            "Redacta directamente el contrato, sin preámbulos ni explicaciones."
        )
    if lang == "en":
        return (
            "You are an expert in residential tenancy law. "
            "Draft a basic residential tenancy agreement with simplified tax treatment. "
            "The document must be formally correct and structured in the following sections:\n\n"
            "1. HEADING — Landlord, tenant, tax/ID numbers, date and place of signing.\n"
            "2. PROPERTY AND TERM — Description of the property, duration in years, start date, "
            "tacit renewal with 3-month notice to terminate.\n"
            "3. RENT AND TAX REGIME — Monthly amount, payment date, applicable tax regime, "
            "no index-linking unless expressly agreed.\n"
            "4. TENANT OBLIGATIONS — Residential use only, no subletting, routine maintenance, "
            "utilities at tenant's expense.\n"
            "5. DEPOSIT AND FINAL CLAUSES — Deposit capped at 2 monthly rents, contract costs, "
            "mandatory mediation, signatures with express approval of general clauses.\n\n"
            f"Use {ph} for every missing piece of information without exception — never invent or assume dates, names, codes, numbers or any other data not explicitly provided: use {ph} for dates, tax codes, addresses and reference numbers. "
            f"IMPORTANT: use EXCLUSIVELY {ph} for missing data. "
            f"It is FORBIDDEN to use descriptive placeholders such as [NAME], [ADDRESS], [DATE] or similar. "
            f"Only {ph} and nothing else for any data not provided. "
            "Do not add meta-legal disclaimers. "
            "Write the document directly, without preamble or explanation."
        )
    return (
        "Sei un avvocato esperto in locazioni abitative italiane. "
        "Redigi un contratto di locazione con cedolare secca. "
        "Il documento deve essere formalmente corretto e strutturato nelle seguenti sezioni:\n\n"
        "1. INTESTAZIONE — Locatore, conduttore, C.F., data e luogo di stipula.\n"
        "2. OGGETTO E DURATA — Descrizione dell'immobile, durata in anni, data di inizio, "
        "rinnovo tacito con disdetta 3 mesi prima.\n"
        "3. CANONE E CEDOLARE SECCA — Importo mensile, giorno di pagamento, regime cedolare secca, "
        "nessun aggiornamento ISTAT.\n"
        "4. OBBLIGHI DEL CONDUTTORE — Uso abitativo, divieto di sublocazione, manutenzione ordinaria, "
        "utenze a carico del conduttore.\n"
        "5. DEPOSITO E CLAUSOLE FINALI — Deposito massimo 2 mensilità, spese contrattuali, "
        "mediazione obbligatoria, firme con doppia sottoscrizione artt. 1341-1342 c.c.\n\n"
        f"Usa {ph} per ogni dato mancante senza eccezioni — non inventare né assumere date, nomi, codici, numeri o qualsiasi altro dato non esplicitamente fornito: usa {ph} per date, codici fiscali, indirizzi e riferimenti numerici. "
        f"IMPORTANTE: usa ESCLUSIVAMENTE {ph} per i dati mancanti. "
        f"È VIETATO usare placeholder descrittivi come [NOME], [INDIRIZZO], [DATA] o simili. "
        f"Solo {ph} e nient'altro per qualsiasi dato non fornito. "
        "Non aggiungere avvertenze meta-legali sul documento stesso. "
        "Scrivi direttamente il contratto, senza prefazioni o spiegazioni."
    )


def _rental_standard_system(lang: str) -> str:
    ph = _placeholder(lang)
    if lang == "es":
        return (
            "Eres un abogado experto en arrendamientos urbanos españoles. "
            "Redacta un contrato de arrendamiento de vivienda habitual de renta pactada (3+2 años). "
            "El documento debe ser formalmente correcto y estructurado en las siguientes secciones:\n\n"
            "1. ENCABEZAMIENTO — Arrendador y arrendatario con datos completos, referencias catastrales, "
            "certificado de eficiencia energética.\n"
            "2. DURACIÓN Y PRÓRROGA — 3 años + prórroga de 2 años, condiciones de desistimiento, "
            "prórroga tácita.\n"
            "3. RENTA Y DEPÓSITO — Renta anual pactada, IBAN, pagos mensuales, actualización por IPC 75%, "
            "depósito máximo 3 mensualidades con intereses.\n"
            "4. CARGAS Y OBLIGACIONES — Distribución de gastos conforme a normativa, uso residencial, "
            "prohibición de subarriendo, entrega y devolución, obras con consentimiento escrito.\n"
            "5. CLÁUSULAS FINALES Y FIRMAS — Comisión paritaria, RGPD, fuero competente, "
            "doble suscripción de cláusulas generales.\n\n"
            f"Usa {ph} para cada dato ausente sin excepción — no inventes ni asumas fechas, nombres, códigos, números ni ningún otro dato no proporcionado explícitamente: usa {ph} para fechas, números de identificación, direcciones y referencias numéricas. "
            f"IMPORTANTE: usa EXCLUSIVAMENTE {ph} para los datos faltantes. "
            f"ESTÁ PROHIBIDO usar marcadores descriptivos como [NOMBRE], [DIRECCIÓN], [FECHA] o similares. "
            f"Solo {ph} y nada más para cualquier dato no proporcionado. "
            "No añadas advertencias meta-legales. "
            "Redacta directamente el contrato, sin preámbulos ni explicaciones."
        )
    if lang == "en":
        return (
            "You are an expert in residential tenancy law. "
            "Draft a standard assured shorthold tenancy agreement (3+2 year term). "
            "The document must be formally correct and structured in the following sections:\n\n"
            "1. HEADING — Landlord and tenant with full details, land registry references, "
            "energy performance certificate.\n"
            "2. TERM AND RENEWAL — 3-year term plus 2-year extension, termination conditions, "
            "tacit renewal.\n"
            "3. RENT AND DEPOSIT — Annual rent as agreed, IBAN, monthly instalments, "
            "75% index-linking, deposit capped at 3 monthly rents with interest.\n"
            "4. OBLIGATIONS — Schedule of charges per applicable regulations, residential use only, "
            "no subletting, handover and return, alterations require written consent.\n"
            "5. FINAL CLAUSES AND SIGNATURES — Joint committee, GDPR, jurisdiction, "
            "express approval of general clauses.\n\n"
            f"Use {ph} for every missing piece of information without exception — never invent or assume dates, names, codes, numbers or any other data not explicitly provided: use {ph} for dates, tax codes, addresses and reference numbers. "
            f"IMPORTANT: use EXCLUSIVELY {ph} for missing data. "
            f"It is FORBIDDEN to use descriptive placeholders such as [NAME], [ADDRESS], [DATE] or similar. "
            f"Only {ph} and nothing else for any data not provided. "
            "Do not add meta-legal disclaimers. "
            "Write the document directly, without preamble or explanation."
        )
    return (
        "Sei un avvocato esperto in locazioni abitative italiane. "
        "Redigi un contratto di locazione abitativa ai sensi della L. 431/98 (3+2 anni). "
        "Il documento deve essere formalmente corretto e strutturato nelle seguenti sezioni:\n\n"
        "1. INTESTAZIONE — Locatore e conduttore con dati completi, riferimenti catastali, "
        "prestazione energetica.\n"
        "2. DURATA E RINNOVO — 3 anni + proroga 2 anni L. 431/98, condizioni di disdetta, "
        "rinnovo tacito.\n"
        "3. CANONE E DEPOSITO — Canone annuale da accordo territoriale, IBAN, rate mensili, "
        "ISTAT 75%, deposito massimo 3 mensilità con interessi.\n"
        "4. ONERI E OBBLIGHI — Tabella oneri D.M. infrastrutture, uso abitativo, "
        "divieto di sublocazione, consegna e riconsegna, modifiche con consenso scritto.\n"
        "5. CLAUSOLE FINALI E FIRME — Commissione paritetica, GDPR, foro competente, "
        "doppia sottoscrizione artt. 1341-1342 c.c.\n\n"
        f"Usa {ph} per ogni dato mancante senza eccezioni — non inventare né assumere date, nomi, codici, numeri o qualsiasi altro dato non esplicitamente fornito: usa {ph} per date, codici fiscali, indirizzi e riferimenti numerici. "
        f"IMPORTANTE: usa ESCLUSIVAMENTE {ph} per i dati mancanti. "
        f"È VIETATO usare placeholder descrittivi come [NOME], [INDIRIZZO], [DATA] o simili. "
        f"Solo {ph} e nient'altro per qualsiasi dato non fornito. "
        "Non aggiungere avvertenze meta-legali sul documento stesso. "
        "Scrivi direttamente il contratto, senza prefazioni o spiegazioni."
    )


def _rental_student_system(lang: str) -> str:
    ph = _placeholder(lang)
    if lang == "es":
        return (
            "Eres un abogado experto en arrendamientos para estudiantes universitarios. "
            "Redacta un contrato de arrendamiento de temporada para estudiante universitario. "
            "El documento debe ser formalmente correcto y estructurado en las siguientes secciones:\n\n"
            "1. ENCABEZAMIENTO — Arrendador y arrendatario con datos completos, referencias catastrales.\n"
            "2. DURACIÓN Y NATURALEZA TRANSITORIA — De 6 meses a 3 años, renovación automática, "
            "declaración expresa de matriculación universitaria con nombre del curso y centro, "
            "acuerdo territorial aplicable.\n"
            "3. RENTA Y DEPÓSITO — Renta mensual, transferencia anticipada antes del día 5, "
            "depósito máximo 3 mensualidades, eventuales garantías adicionales.\n"
            "4. OBLIGACIONES DE LAS PARTES — Suministros a cargo del arrendatario, uso residencial, "
            "prohibición de subarriendo con resolución de pleno derecho, desistimiento por causa grave "
            "con preaviso de 3 meses, entrega y devolución del inmueble.\n"
            "5. CLÁUSULAS FINALES Y FIRMAS — Comisión paritaria, RGPD, fuero competente, "
            "doble suscripción de cláusulas generales.\n\n"
            f"Usa {ph} para cada dato ausente sin excepción — no inventes ni asumas fechas, nombres, códigos, números ni ningún otro dato no proporcionado explícitamente: usa {ph} para fechas, números de identificación, direcciones y referencias numéricas. "
            f"IMPORTANTE: usa EXCLUSIVAMENTE {ph} para los datos faltantes. "
            f"ESTÁ PROHIBIDO usar marcadores descriptivos como [NOMBRE], [DIRECCIÓN], [FECHA] o similares. "
            f"Solo {ph} y nada más para cualquier dato no proporcionado. "
            "No añadas advertencias meta-legales. "
            "Redacta directamente el contrato, sin preámbulos ni explicaciones."
        )
    if lang == "en":
        return (
            "You are an expert in student residential tenancy law. "
            "Draft a student tenancy agreement for a university student. "
            "The document must be formally correct and structured in the following sections:\n\n"
            "1. HEADING — Landlord and tenant with full details, land registry references.\n"
            "2. TERM AND TRANSITIONAL NATURE — 6 months to 3 years, automatic renewal, "
            "explicit declaration of university enrolment with course name and institution, "
            "applicable territorial agreement.\n"
            "3. RENT AND DEPOSIT — Monthly rent, advance bank transfer by day 5, "
            "deposit capped at 3 monthly rents, any additional guarantees.\n"
            "4. OBLIGATIONS OF THE PARTIES — Utilities at tenant's expense, residential use only, "
            "no subletting with automatic termination clause, early termination for serious cause "
            "with 3-month notice, handover and return of the property.\n"
            "5. FINAL CLAUSES AND SIGNATURES — Joint committee, GDPR, jurisdiction, "
            "express approval of general clauses.\n\n"
            f"Use {ph} for every missing piece of information without exception — never invent or assume dates, names, codes, numbers or any other data not explicitly provided: use {ph} for dates, tax codes, addresses and reference numbers. "
            f"IMPORTANT: use EXCLUSIVELY {ph} for missing data. "
            f"It is FORBIDDEN to use descriptive placeholders such as [NAME], [ADDRESS], [DATE] or similar. "
            f"Only {ph} and nothing else for any data not provided. "
            "Do not add meta-legal disclaimers. "
            "Write the document directly, without preamble or explanation."
        )
    return (
        "Sei un avvocato esperto in locazioni abitative italiane. "
        "Redigi un contratto di locazione abitativa per studenti universitari ai sensi della L. 431/98. "
        "Il documento deve essere formalmente corretto e strutturato nelle seguenti sezioni:\n\n"
        "1. INTESTAZIONE — Locatore e conduttore con dati completi, riferimenti catastali.\n"
        "2. DURATA E NATURA TRANSITORIA — Da 6 mesi a 3 anni, rinnovo automatico, "
        "dichiarazione esplicita di frequenza del corso universitario con nome del corso e ateneo, "
        "accordo territoriale applicabile.\n"
        "3. CANONE E DEPOSITO — Canone mensile, bonifico anticipato entro il giorno 5, "
        "deposito massimo 3 mensilità, eventuali garanzie aggiuntive.\n"
        "4. OBBLIGHI DELLE PARTI — Utenze a carico del conduttore, uso abitativo, "
        "divieto di sublocazione con risoluzione di diritto, recesso per gravi motivi con preavviso "
        "di 3 mesi, consegna e riconsegna dell'immobile.\n"
        "5. CLAUSOLE FINALI E FIRME — Commissione paritetica, GDPR, foro competente, "
        "doppia sottoscrizione artt. 1341-1342 c.c.\n\n"
        f"Usa {ph} per ogni dato mancante senza eccezioni — non inventare né assumere date, nomi, codici, numeri o qualsiasi altro dato non esplicitamente fornito: usa {ph} per date, codici fiscali, indirizzi e riferimenti numerici. "
        f"IMPORTANTE: usa ESCLUSIVAMENTE {ph} per i dati mancanti. "
        f"È VIETATO usare placeholder descrittivi come [NOME], [INDIRIZZO], [DATA] o simili. "
        f"Solo {ph} e nient'altro per qualsiasi dato non fornito. "
        "Non aggiungere avvertenze meta-legali sul documento stesso. "
        "Scrivi direttamente il contratto, senza prefazioni o spiegazioni."
    )


def _rental_transitional_system(lang: str) -> str:
    ph = _placeholder(lang)
    if lang == "es":
        return (
            "Eres un abogado experto en arrendamientos de temporada y transitorios. "
            "Redacta un contrato de arrendamiento de naturaleza transitoria. "
            "El documento debe ser formalmente correcto y estructurado en las siguientes secciones:\n\n"
            "1. ENCABEZAMIENTO — Arrendador y arrendatario con datos completos, referencias catastrales.\n"
            "2. DURACIÓN Y NECESIDAD TRANSITORIA — De 1 a 18 meses, cese automático sin denuncia, "
            "motivación transitoria documentada obligatoria para duraciones superiores a 30 días, "
            "acuerdo territorial aplicable, consecuencias del incumplimiento de las modalidades de firma.\n"
            "3. RENTA Y DEPÓSITO — Renta mensual, transferencia anticipada antes del día 5, "
            "depósito máximo 3 mensualidades, eventuales garantías adicionales.\n"
            "4. OBLIGACIONES DE LAS PARTES — Suministros a cargo del arrendatario, uso residencial "
            "con lista de convivientes, prohibición de subarriendo, desistimiento por causa grave "
            "con preaviso de 30 días no aplicable a contratos de 30 días o menos, "
            "entrega y devolución del inmueble.\n"
            "5. CLÁUSULAS FINALES Y FIRMAS — Comisión paritaria, RGPD, fuero competente, "
            "doble suscripción de cláusulas generales.\n\n"
            f"Usa {ph} para cada dato ausente sin excepción — no inventes ni asumas fechas, nombres, códigos, números ni ningún otro dato no proporcionado explícitamente: usa {ph} para fechas, números de identificación, direcciones y referencias numéricas. "
            f"IMPORTANTE: usa EXCLUSIVAMENTE {ph} para los datos faltantes. "
            f"ESTÁ PROHIBIDO usar marcadores descriptivos como [NOMBRE], [DIRECCIÓN], [FECHA] o similares. "
            f"Solo {ph} y nada más para cualquier dato no proporcionado. "
            "No añadas advertencias meta-legales. "
            "Redacta directamente el contrato, sin preámbulos ni explicaciones."
        )
    if lang == "en":
        return (
            "You are an expert in short-term and transitional residential tenancy law. "
            "Draft a transitional residential tenancy agreement. "
            "The document must be formally correct and structured in the following sections:\n\n"
            "1. HEADING — Landlord and tenant with full details, land registry references.\n"
            "2. TERM AND TRANSITIONAL NEED — 1 to 18 months, automatic expiry without notice, "
            "documented transitional reason mandatory for terms exceeding 30 days, "
            "applicable territorial agreement, consequences of non-compliance with signing requirements.\n"
            "3. RENT AND DEPOSIT — Monthly rent, advance bank transfer by day 5, "
            "deposit capped at 3 monthly rents, any additional guarantees.\n"
            "4. OBLIGATIONS OF THE PARTIES — Utilities at tenant's expense, residential use with list "
            "of occupants, no subletting, early termination for serious cause with 30-day notice "
            "not applicable to agreements of 30 days or less, handover and return of the property.\n"
            "5. FINAL CLAUSES AND SIGNATURES — Joint committee, GDPR, jurisdiction, "
            "express approval of general clauses.\n\n"
            f"Use {ph} for every missing piece of information without exception — never invent or assume dates, names, codes, numbers or any other data not explicitly provided: use {ph} for dates, tax codes, addresses and reference numbers. "
            f"IMPORTANT: use EXCLUSIVELY {ph} for missing data. "
            f"It is FORBIDDEN to use descriptive placeholders such as [NAME], [ADDRESS], [DATE] or similar. "
            f"Only {ph} and nothing else for any data not provided. "
            "Do not add meta-legal disclaimers. "
            "Write the document directly, without preamble or explanation."
        )
    return (
        "Sei un avvocato esperto in locazioni abitative italiane. "
        "Redigi un contratto di locazione abitativa di natura transitoria ai sensi della L. 431/98. "
        "Il documento deve essere formalmente corretto e strutturato nelle seguenti sezioni:\n\n"
        "1. INTESTAZIONE — Locatore e conduttore con dati completi, riferimenti catastali.\n"
        "2. DURATA ED ESIGENZA TRANSITORIA — Da 1 a 18 mesi, cessazione automatica senza disdetta, "
        "motivazione transitoria documentata obbligatoria per durate superiori a 30 giorni, "
        "accordo territoriale applicabile, conseguenze dell'inadempimento delle modalità di stipula.\n"
        "3. CANONE E DEPOSITO — Canone mensile, bonifico anticipato entro il giorno 5, "
        "deposito massimo 3 mensilità, eventuali garanzie aggiuntive.\n"
        "4. OBBLIGHI DELLE PARTI — Utenze a carico del conduttore, uso abitativo con elenco dei "
        "conviventi, divieto di sublocazione, recesso per gravi motivi con preavviso di 30 giorni "
        "non applicabile a contratti di durata ≤ 30 giorni, consegna e riconsegna dell'immobile.\n"
        "5. CLAUSOLE FINALI E FIRME — Commissione paritetica, GDPR, foro competente, "
        "doppia sottoscrizione artt. 1341-1342 c.c.\n\n"
        f"Usa {ph} per ogni dato mancante senza eccezioni — non inventare né assumere date, nomi, codici, numeri o qualsiasi altro dato non esplicitamente fornito: usa {ph} per date, codici fiscali, indirizzi e riferimenti numerici. "
        f"IMPORTANTE: usa ESCLUSIVAMENTE {ph} per i dati mancanti. "
        f"È VIETATO usare placeholder descrittivi come [NOME], [INDIRIZZO], [DATA] o simili. "
        f"Solo {ph} e nient'altro per qualsiasi dato non fornito. "
        "Non aggiungere avvertenze meta-legali sul documento stesso. "
        "Scrivi direttamente il contratto, senza prefazioni o spiegazioni."
    )


def _rental_free_rent_system(lang: str) -> str:
    ph = _placeholder(lang)
    if lang == "es":
        return (
            "Eres un abogado experto en arrendamientos urbanos de renta libre. "
            "Redacta un contrato de arrendamiento de vivienda a renta libre (4+4 años). "
            "El documento debe ser formalmente correcto y estructurado en las siguientes secciones:\n\n"
            "1. ENCABEZAMIENTO Y PREMISAS — Arrendador y arrendatario con datos completos, "
            "declaraciones preliminares de propiedad, referencias catastrales, certificado energético.\n"
            "2. DURACIÓN Y DESISTIMIENTO — 4+4 años, preaviso de desistimiento del arrendador "
            "de 6 meses con causa tasada, desistimiento del arrendatario con preaviso de 6 meses "
            "por carta certificada.\n"
            "3. OBLIGACIONES Y RENTA — Uso residencial, prohibición de subarriendo, renta anual "
            "con pagos mensuales anticipados antes del día 5 por transferencia IBAN, "
            "actualización IPC 75% desde el mes 13, gastos ordinarios a cargo del arrendatario, "
            "gastos extraordinarios a cargo del arrendador.\n"
            "4. DEPÓSITO Y GARANTÍAS — Depósito máximo 3 mensualidades con intereses legales anuales, "
            "no imputable a renta, eventual garantía suplementaria.\n"
            "5. DISPOSICIONES FINALES Y FIRMAS — Fecha de eficacia, derecho español, RGPD, "
            "registro 50% cada parte, fuero competente, doble suscripción de cláusulas generales.\n\n"
            f"Usa {ph} para cada dato ausente sin excepción — no inventes ni asumas fechas, nombres, códigos, números ni ningún otro dato no proporcionado explícitamente: usa {ph} para fechas, números de identificación, direcciones y referencias numéricas. "
            f"IMPORTANTE: usa EXCLUSIVAMENTE {ph} para los datos faltantes. "
            f"ESTÁ PROHIBIDO usar marcadores descriptivos como [NOMBRE], [DIRECCIÓN], [FECHA] o similares. "
            f"Solo {ph} y nada más para cualquier dato no proporcionado. "
            "No añadas advertencias meta-legales. "
            "Redacta directamente el contrato, sin preámbulos ni explicaciones."
        )
    if lang == "en":
        return (
            "You are an expert in open-market residential tenancy law. "
            "Draft an open-market tenancy agreement (4+4 year term). "
            "The document must be formally correct and structured in the following sections:\n\n"
            "1. HEADING AND RECITALS — Landlord and tenant with full details, ownership declarations, "
            "land registry references, energy performance certificate.\n"
            "2. TERM AND TERMINATION — 4+4 years, landlord's 6-month notice with statutory grounds, "
            "tenant's 6-month notice by recorded delivery.\n"
            "3. OBLIGATIONS AND RENT — Residential use only, no subletting, annual rent with monthly "
            "advance payments by day 5 via IBAN bank transfer, 75% index-linking from month 13, "
            "routine repairs at tenant's expense, major repairs at landlord's expense.\n"
            "4. DEPOSIT AND GUARANTEES — Deposit capped at 3 monthly rents with annual statutory "
            "interest, not offset against rent, any supplementary guarantee.\n"
            "5. FINAL PROVISIONS AND SIGNATURES — Effective date, governing law, GDPR, registration "
            "costs split 50/50, jurisdiction, express approval of general clauses.\n\n"
            f"Use {ph} for every missing piece of information without exception — never invent or assume dates, names, codes, numbers or any other data not explicitly provided: use {ph} for dates, tax codes, addresses and reference numbers. "
            f"IMPORTANT: use EXCLUSIVELY {ph} for missing data. "
            f"It is FORBIDDEN to use descriptive placeholders such as [NAME], [ADDRESS], [DATE] or similar. "
            f"Only {ph} and nothing else for any data not provided. "
            "Do not add meta-legal disclaimers. "
            "Write the document directly, without preamble or explanation."
        )
    return (
        "Sei un avvocato esperto in locazioni abitative italiane. "
        "Redigi un contratto di locazione ad uso abitativo a canone libero ai sensi della L. 431/98 (4+4 anni). "
        "Il documento deve essere formalmente corretto e strutturato nelle seguenti sezioni:\n\n"
        "1. INTESTAZIONE E PREMESSE — Locatore e conduttore con dati completi, "
        "dichiarazioni preliminari di proprietà, riferimenti catastali, attestato energetico D.Lgs. 28/2011.\n"
        "2. DURATA E RECESSO — 4+4 anni L. 431/98, preavviso di disdetta del locatore 6 mesi "
        "con motivazione tassativa, recesso del conduttore con preavviso di 6 mesi tramite raccomandata.\n"
        "3. OBBLIGHI E CANONE — Uso abitativo, divieto di sublocazione, canone annuale con rate "
        "mensili anticipate entro il giorno 5 tramite bonifico IBAN, ISTAT 75% dal tredicesimo mese, "
        "spese ordinarie a carico del conduttore, spese straordinarie a carico del locatore.\n"
        "4. DEPOSITO E GARANZIE — Deposito massimo 3 mensilità con interessi legali annuali "
        "art. 11 L. 392/78, non imputabile a canone, eventuale garanzia supplementare.\n"
        "5. DISPOSIZIONI FINALI E FIRME — Data di efficacia, diritto italiano, GDPR, "
        "registrazione 50% ciascuna parte, foro competente, doppia sottoscrizione artt. 1341-1342 c.c.\n\n"
        f"Usa {ph} per ogni dato mancante senza eccezioni — non inventare né assumere date, nomi, codici, numeri o qualsiasi altro dato non esplicitamente fornito: usa {ph} per date, codici fiscali, indirizzi e riferimenti numerici. "
        f"IMPORTANTE: usa ESCLUSIVAMENTE {ph} per i dati mancanti. "
        f"È VIETATO usare placeholder descrittivi come [NOME], [INDIRIZZO], [DATA] o simili. "
        f"Solo {ph} e nient'altro per qualsiasi dato non fornito. "
        "Non aggiungere avvertenze meta-legali sul documento stesso. "
        "Scrivi direttamente il contratto, senza prefazioni o spiegazioni."
    )


def _rental_commercial_system(lang: str) -> str:
    ph = _placeholder(lang)
    if lang == "es":
        return (
            "Eres un abogado experto en arrendamientos de locales de negocio. "
            "Redacta un contrato de arrendamiento para uso distinto del de vivienda (6+6 años). "
            "El documento debe ser formalmente correcto y estructurado en las siguientes secciones:\n\n"
            "1. ENCABEZAMIENTO — Parte arrendadora y arrendataria con razón social, N.I.F., "
            "representante legal, referencias catastrales, declaraciones de conformidad del inmueble, "
            "certificado energético.\n"
            "2. DURACIÓN Y DESISTIMIENTO — 6+6 años, preaviso de denuncia de 12 meses, "
            "renuncia del arrendador a denegar la renovación en el primer vencimiento, "
            "desistimiento de la arrendataria con preaviso de 6 meses.\n"
            "3. RENTA Y PAGO — Renta anual más IVA, gastos de conducción y calefacción, "
            "4 pagos trimestrales anticipados 01/01-01/04-01/07-01/10, "
            "actualización IPC 75% desde el segundo año.\n"
            "4. USO MEJORAS Y CARGAS — Destino de uso especificado, prohibición de cambio de destino, "
            "mejoras solo con consentimiento escrito quedan en beneficio del arrendador, "
            "prohibición de subarriendo salvo grupo societario, todos los gastos a cargo de la "
            "arrendataria incluidos tributos locales.\n"
            "5. DEPÓSITO VARIOS Y FIRMAS — Depósito de 3 mensualidades por transferencia, "
            "visita del arrendador con preaviso de 24 horas, comunicaciones por burofax, "
            "timbre y registro 50% arrendataria, fuero competente exclusivo, "
            "firmas con aprobación de cláusulas esenciales.\n\n"
            f"Usa {ph} para cada dato ausente sin excepción — no inventes ni asumas fechas, nombres, códigos, números ni ningún otro dato no proporcionado explícitamente: usa {ph} para fechas, números de identificación, direcciones y referencias numéricas. "
            f"IMPORTANTE: usa EXCLUSIVAMENTE {ph} para los datos faltantes. "
            f"ESTÁ PROHIBIDO usar marcadores descriptivos como [NOMBRE], [DIRECCIÓN], [FECHA] o similares. "
            f"Solo {ph} y nada más para cualquier dato no proporcionado. "
            "No añadas advertencias meta-legales. "
            "Redacta directamente el contrato, sin preámbulos ni explicaciones."
        )
    if lang == "en":
        return (
            "You are an expert in commercial property tenancy law. "
            "Draft a commercial lease agreement (6+6 year term). "
            "The document must be formally correct and structured in the following sections:\n\n"
            "1. HEADING — Landlord and tenant with company name, VAT number, legal representative, "
            "land registry references, compliance declarations, energy performance certificate.\n"
            "2. TERM AND TERMINATION — 6+6 years, 12-month notice to terminate, landlord waives "
            "right to refuse renewal at first expiry, tenant's early termination with 6-month notice.\n"
            "3. RENT AND PAYMENT — Annual rent plus VAT, service and heating charges, "
            "4 quarterly advance instalments 01/01-01/04-01/07-01/10, "
            "75% index-linking from the second year.\n"
            "4. USE IMPROVEMENTS AND CHARGES — Specified permitted use, no change of use, "
            "improvements only with written consent and revert to landlord, no subletting except "
            "within group companies, all outgoings at tenant's expense including local taxes.\n"
            "5. DEPOSIT MISCELLANEOUS AND SIGNATURES — 3-month deposit by bank transfer, "
            "landlord access with 24-hour notice, notices by recorded delivery or PEC, "
            "stamp duty and registration 50% tenant, exclusive jurisdiction clause, "
            "signatures with express approval of essential clauses.\n\n"
            f"Use {ph} for every missing piece of information without exception — never invent or assume dates, names, codes, numbers or any other data not explicitly provided: use {ph} for dates, tax codes, addresses and reference numbers. "
            f"IMPORTANT: use EXCLUSIVELY {ph} for missing data. "
            f"It is FORBIDDEN to use descriptive placeholders such as [NAME], [ADDRESS], [DATE] or similar. "
            f"Only {ph} and nothing else for any data not provided. "
            "Do not add meta-legal disclaimers. "
            "Write the document directly, without preamble or explanation."
        )
    return (
        "Sei un avvocato esperto in locazioni commerciali italiane. "
        "Redigi un contratto di locazione ad uso commerciale ai sensi della L. 392/78 (6+6 anni). "
        "Il documento deve essere formalmente corretto e strutturato nelle seguenti sezioni:\n\n"
        "1. INTESTAZIONE — Parte locatrice e conduttrice con ragione sociale, P.IVA, "
        "rappresentante legale, riferimenti catastali, dichiarazioni di conformità dell'immobile, "
        "attestato energetico.\n"
        "2. DURATA E RECESSO — 6+6 anni L. 392/78, preavviso di disdetta 12 mesi, "
        "rinuncia della locatrice al diniego di rinnovo alla prima scadenza, "
        "recesso della conduttrice con preavviso di 6 mesi.\n"
        "3. CANONE E PAGAMENTO — Canone annuale oltre IVA, spese di conduzione e riscaldamento, "
        "4 rate trimestrali anticipate 01/01-01/04-01/07-01/10, "
        "ISTAT 75% dal secondo anno.\n"
        "4. USO MIGLIORIE E ONERI — Destinazione d'uso specificata, divieto di cambio di destinazione, "
        "migliorie solo con consenso scritto e restano al locatore, divieto di sublocazione salvo "
        "gruppo societario, tutti gli oneri a carico della conduttrice inclusi TARI e TASI.\n"
        "5. DEPOSITO VARIE E FIRME — Deposito 3 mensilità tramite bonifico, visita della locatrice "
        "con preavviso di 24 ore, comunicazioni tramite PEC, bollo e registro 50% conduttrice, "
        "foro competente esclusivo, firme con approvazione delle clausole essenziali artt. 1341-1342 c.c.\n\n"
        f"Usa {ph} per ogni dato mancante senza eccezioni — non inventare né assumere date, nomi, codici, numeri o qualsiasi altro dato non esplicitamente fornito: usa {ph} per date, codici fiscali, indirizzi e riferimenti numerici. "
        f"IMPORTANTE: usa ESCLUSIVAMENTE {ph} per i dati mancanti. "
        f"È VIETATO usare placeholder descrittivi come [NOME], [INDIRIZZO], [DATA] o simili. "
        f"Solo {ph} e nient'altro per qualsiasi dato non fornito. "
        "Non aggiungere avvertenze meta-legali sul documento stesso. "
        "Scrivi direttamente il contratto, senza prefazioni o spiegazioni."
    )


def _rental_cancellation_system(lang: str) -> str:
    ph = _placeholder(lang)
    if lang == "es":
        return (
            "Eres un abogado experto en arrendamientos urbanos. "
            "Redacta una carta de desistimiento de contrato de arrendamiento de vivienda. "
            "El documento debe ser formalmente correcto y estructurado en las siguientes secciones:\n\n"
            "1. ENCABEZAMIENTO — Datos del arrendatario (nombre, DNI, domicilio), datos del arrendador "
            "(nombre, domicilio), medio de envío (carta certificada con acuse de recibo o correo electrónico "
            "certificado), objeto de la comunicación.\n"
            "2. DECLARACIÓN DE DESISTIMIENTO — Manifestación expresa de desistir del contrato de "
            "arrendamiento, fecha de firma del contrato, número de referencia de registro, "
            "dirección del inmueble arrendado.\n"
            "3. PLAZO DE ENTREGA — Fecha en la que el inmueble será entregado libre de personas y enseres, "
            "plazo de preaviso en meses/días desde la recepción de la presente comunicación.\n"
            "4. CIERRE Y FIRMA — Fórmula de cierre, lugar, fecha, firma del arrendatario.\n\n"
            f"Usa {ph} para cada dato ausente sin excepción — no inventes ni asumas fechas, nombres, códigos, números ni ningún otro dato no proporcionado explícitamente: usa {ph} para fechas, números de identificación, direcciones y referencias numéricas. "
            f"IMPORTANTE: usa EXCLUSIVAMENTE {ph} para los datos faltantes. "
            f"ESTÁ PROHIBIDO usar marcadores descriptivos como [NOMBRE], [DIRECCIÓN], [FECHA] o similares. "
            f"Solo {ph} y nada más para cualquier dato no proporcionado. "
            "No añadas advertencias meta-legales. "
            "Redacta directamente la carta, sin preámbulos ni explicaciones."
        )
    if lang == "en":
        return (
            "You are an expert in residential tenancy law. "
            "Draft a notice of termination of a residential tenancy agreement. "
            "The document must be formally correct and structured in the following sections:\n\n"
            "1. HEADING — Tenant's details (name, ID, address), landlord's details (name, address), "
            "method of delivery (recorded delivery letter or certified email), subject of the notice.\n"
            "2. NOTICE OF TERMINATION — Express statement of intent to terminate the tenancy agreement, "
            "date the agreement was signed, land registry or tax registration reference, "
            "address of the rented property.\n"
            "3. HANDOVER DATE — Date by which the property will be returned vacant and free of belongings, "
            "notice period in months/days from receipt of this communication.\n"
            "4. CLOSING AND SIGNATURE — Closing formula, place, date, tenant's signature.\n\n"
            f"Use {ph} for every missing piece of information without exception — never invent or assume dates, names, codes, numbers or any other data not explicitly provided: use {ph} for dates, tax codes, addresses and reference numbers. "
            f"IMPORTANT: use EXCLUSIVELY {ph} for missing data. "
            f"It is FORBIDDEN to use descriptive placeholders such as [NAME], [ADDRESS], [DATE] or similar. "
            f"Only {ph} and nothing else for any data not provided. "
            "Do not add meta-legal disclaimers. "
            "Write the document directly, without preamble or explanation."
        )
    return (
        "Sei un avvocato esperto in locazioni abitative italiane. "
        "Redigi una disdetta di contratto di locazione da parte del conduttore. "
        "Il documento deve essere formalmente corretto e strutturato nelle seguenti sezioni:\n\n"
        "1. INTESTAZIONE — Dati del conduttore (nome, cognome, C.F., indirizzo), dati del locatore "
        "(nome, cognome, indirizzo), modalità di invio (raccomandata A/R o PEC), oggetto della comunicazione.\n"
        "2. DICHIARAZIONE DI RECESSO — Manifestazione espressa di voler recedere anticipatamente dal "
        "contratto di locazione, data di stipula del contratto, data di registrazione presso l'Agenzia "
        "delle Entrate, indirizzo dell'immobile locato.\n"
        "3. TERMINE DI CONSEGNA — Data entro cui l'immobile sarà consegnato libero da cose e persone, "
        "preavviso in mesi/giorni dalla ricezione della presente comunicazione.\n"
        "4. CHIUSURA E FIRMA — Formula di chiusura, luogo, data, firma del conduttore.\n\n"
        f"Usa {ph} per ogni dato mancante senza eccezioni — non inventare né assumere date, nomi, codici, numeri o qualsiasi altro dato non esplicitamente fornito: usa {ph} per date, codici fiscali, indirizzi e riferimenti numerici. "
        f"IMPORTANTE: usa ESCLUSIVAMENTE {ph} per i dati mancanti. "
        f"È VIETATO usare placeholder descrittivi come [NOME], [INDIRIZZO], [DATA] o simili. "
        f"Solo {ph} e nient'altro per qualsiasi dato non fornito. "
        "Non aggiungere avvertenze meta-legali sul documento stesso. "
        "Scrivi direttamente la lettera, senza prefazioni o spiegazioni."
    )


def _insurance_cancellation_system(lang: str) -> str:
    ph = _placeholder(lang)
    if lang == "es":
        return (
            "Eres un experto en derecho de seguros. "
            "Redacta una carta de desistimiento o no renovación de póliza de seguro por vencimiento natural. "
            "El documento debe ser formalmente correcto y estructurado en las siguientes secciones:\n\n"
            "1. ENCABEZAMIENTO — Lugar y fecha, destinatario (nombre compañía aseguradora, dirección), "
            "objeto con número de póliza y fecha de vencimiento.\n"
            "2. CUERPO DE LA CARTA — Solicitud formal de no renovación de la póliza por vencimiento natural, "
            "referencia a las condiciones generales de la póliza y a la normativa vigente, "
            "petición de acuse de recibo.\n"
            "3. FIRMA — Fórmula de cierre cordial, firma del asegurado.\n\n"
            f"Usa {ph} para cada dato ausente sin excepción — no inventes ni asumas fechas, nombres, códigos, números ni ningún otro dato no proporcionado explícitamente: usa {ph} para fechas, números de identificación, direcciones y referencias numéricas. "
            f"IMPORTANTE: usa EXCLUSIVAMENTE {ph} para los datos faltantes. "
            f"ESTÁ PROHIBIDO usar marcadores descriptivos como [NOMBRE], [DIRECCIÓN], [FECHA] o similares. "
            f"Solo {ph} y nada más para cualquier dato no proporcionado. "
            "No añadas advertencias meta-legales. "
            "Redacta directamente la carta, sin preámbulos ni explicaciones."
        )
    if lang == "en":
        return (
            "You are an expert in insurance law. "
            "Draft a notice of cancellation or non-renewal of an insurance policy at natural expiry. "
            "The document must be formally correct and structured in the following sections:\n\n"
            "1. HEADING — Place and date, addressee (insurance company name, address), "
            "subject with policy number and expiry date.\n"
            "2. BODY OF THE LETTER — Formal request not to renew the policy at natural expiry, "
            "reference to the general policy conditions and applicable regulations, "
            "request for acknowledgement of receipt.\n"
            "3. SIGNATURE — Cordial closing formula, policyholder's signature.\n\n"
            f"Use {ph} for every missing piece of information without exception — never invent or assume dates, names, codes, numbers or any other data not explicitly provided: use {ph} for dates, tax codes, addresses and reference numbers. "
            f"IMPORTANT: use EXCLUSIVELY {ph} for missing data. "
            f"It is FORBIDDEN to use descriptive placeholders such as [NAME], [ADDRESS], [DATE] or similar. "
            f"Only {ph} and nothing else for any data not provided. "
            "Do not add meta-legal disclaimers. "
            "Write the document directly, without preamble or explanation."
        )
    return (
        "Sei un esperto in diritto assicurativo. "
        "Redigi una lettera di disdetta di polizza assicurativa per naturale scadenza. "
        "Il documento deve essere formalmente corretto e strutturato nelle seguenti sezioni:\n\n"
        "1. INTESTAZIONE — Luogo e data, destinatario (nome compagnia assicurativa, indirizzo), "
        "oggetto con numero polizza e data di scadenza.\n"
        "2. CORPO DELLA LETTERA — Richiesta formale di disdetta per naturale scadenza, riferimento "
        "alle condizioni di polizza e alle norme vigenti, richiesta di ricevuta di ritorno.\n"
        "3. FIRMA — Formula di chiusura cordiale, firma dell'assicurato.\n\n"
        f"Usa {ph} per ogni dato mancante senza eccezioni — non inventare né assumere date, nomi, codici, numeri o qualsiasi altro dato non esplicitamente fornito: usa {ph} per date, codici fiscali, indirizzi e riferimenti numerici. "
        f"IMPORTANTE: usa ESCLUSIVAMENTE {ph} per i dati mancanti. "
        f"È VIETATO usare placeholder descrittivi come [NOME], [INDIRIZZO], [DATA] o simili. "
        f"Solo {ph} e nient'altro per qualsiasi dato non fornito. "
        "Non aggiungere avvertenze meta-legali sul documento stesso. "
        "Scrivi direttamente la lettera, senza prefazioni o spiegazioni."
    )


def _insurance_declaration_system(lang: str) -> str:
    ph = _placeholder(lang)
    if lang == "es":
        return (
            "Eres un experto en derecho de seguros y documentación oficial. "
            "Redacta una declaración sustitutiva de póliza de seguro para una actividad artístico-cultural. "
            "El documento debe ser formalmente correcto y estructurado en las siguientes secciones:\n\n"
            "1. ENCABEZAMIENTO — Fac-símil en papel con membrete, título declaración sustitutiva.\n"
            "2. DATOS DEL DECLARANTE — Nombre y apellidos, lugar y fecha de nacimiento, domicilio, "
            "calidad en que actúa, nombre de la organización, sede.\n"
            "3. PERÍODO DE REFERENCIA — Fechas desde/hasta incluidos días de montaje/desmontaje, "
            "actividades autorizadas objeto de la cobertura.\n"
            "4. DECLARACIONES — Póliza contra accidentes para actividades artístico-culturales, "
            "cobertura de daños a cosas de terceros, aviso APL sobre señalización de espacios.\n"
            "5. CONSENTIMIENTO Y FIRMA — Consentimiento para el tratamiento de datos personales, "
            "fecha, firma, adjuntar documento de identidad.\n\n"
            f"Usa {ph} para cada dato ausente sin excepción — no inventes ni asumas fechas, nombres, códigos, números ni ningún otro dato no proporcionado explícitamente: usa {ph} para fechas, números de identificación, direcciones y referencias numéricas. "
            f"IMPORTANTE: usa EXCLUSIVAMENTE {ph} para los datos faltantes. "
            f"ESTÁ PROHIBIDO usar marcadores descriptivos como [NOMBRE], [DIRECCIÓN], [FECHA] o similares. "
            f"Solo {ph} y nada más para cualquier dato no proporcionado. "
            "No añadas advertencias meta-legales. "
            "Redacta directamente el documento, sin preámbulos ni explicaciones."
        )
    if lang == "en":
        return (
            "You are an expert in insurance law and official documentation. "
            "Draft a substitute declaration of insurance policy for an artistic-cultural activity. "
            "The document must be formally correct and structured in the following sections:\n\n"
            "1. HEADING — Facsimile on headed paper, title of the substitute declaration.\n"
            "2. DECLARANT'S DETAILS — Full name, place and date of birth, address, "
            "capacity in which acting, name of the organisation, registered office.\n"
            "3. REFERENCE PERIOD — Dates from/to including set-up/dismantling days, "
            "authorised activities covered by the policy.\n"
            "4. DECLARATIONS — Accident policy for artistic-cultural activities, "
            "coverage for damage to third-party property, APL notice regarding signage of venues.\n"
            "5. CONSENT AND SIGNATURE — Consent to personal data processing, "
            "date, signature, attach identity document.\n\n"
            f"Use {ph} for every missing piece of information without exception — never invent or assume dates, names, codes, numbers or any other data not explicitly provided: use {ph} for dates, tax codes, addresses and reference numbers. "
            f"IMPORTANT: use EXCLUSIVELY {ph} for missing data. "
            f"It is FORBIDDEN to use descriptive placeholders such as [NAME], [ADDRESS], [DATE] or similar. "
            f"Only {ph} and nothing else for any data not provided. "
            "Do not add meta-legal disclaimers. "
            "Write the document directly, without preamble or explanation."
        )
    return (
        "Sei un esperto in diritto assicurativo e documentazione ufficiale. "
        "Redigi una dichiarazione sostitutiva di polizza assicurativa per un'attività artistico-culturale. "
        "Il documento deve essere formalmente corretto e strutturato nelle seguenti sezioni:\n\n"
        "1. INTESTAZIONE — Fac-simile su carta intestata, titolo dichiarazione sostitutiva.\n"
        "2. DATI DEL DICHIARANTE — Nome e cognome, luogo e data di nascita, residente in, "
        "qualità in cui agisce, nome dell'organizzazione, sede.\n"
        "3. PERIODO DI RIFERIMENTO — Date dal/al inclusi giorni di allestimento/disallestimento, "
        "attività autorizzate oggetto della copertura.\n"
        "4. DICHIARAZIONI — Polizza contro infortuni per attività artistico-culturali, copertura danni "
        "a cose di terzi, avvertenza APL su segnalazione degli spazi.\n"
        "5. CONSENSO E FIRMA — Consenso al trattamento dei dati personali ai sensi della L. 196/2003, "
        "data, firma, allegare documento di identità.\n\n"
        f"Usa {ph} per ogni dato mancante senza eccezioni — non inventare né assumere date, nomi, codici, numeri o qualsiasi altro dato non esplicitamente fornito: usa {ph} per date, codici fiscali, indirizzi e riferimenti numerici. "
        f"IMPORTANTE: usa ESCLUSIVAMENTE {ph} per i dati mancanti. "
        f"È VIETATO usare placeholder descrittivi come [NOME], [INDIRIZZO], [DATA] o simili. "
        f"Solo {ph} e nient'altro per qualsiasi dato non fornito. "
        "Non aggiungere avvertenze meta-legali sul documento stesso. "
        "Scrivi direttamente il documento, senza prefazioni o spiegazioni."
    )


def _employment_dismissal_appeal_system(lang: str) -> str:
    ph = _placeholder(lang)
    if lang == "es":
        return (
            "Eres un experto en derecho laboral. "
            "Redacta una impugnación formal de despido para proteger los derechos del trabajador. "
            "El documento debe ser formalmente correcto y estructurado en las siguientes secciones:\n\n"
            "1. ENCABEZAMIENTO — Trabajador remitente con domicilio, empleador destinatario con dirección, "
            "modalidad de envío (carta certificada con acuse de recibo o PEC), objeto: impugnación de despido.\n"
            "2. DATOS DE LA RELACIÓN LABORAL — Período de trabajo desde/hasta, categoría y puesto, "
            "fecha y modalidad del despido (si fue solo verbal, especificarlo).\n"
            "3. MOTIVOS DE ILEGITIMIDAD — Exposición de las razones por las que el despido se considera nulo y/o ilegítimo.\n"
            "4. IMPUGNACIÓN Y REQUERIMIENTO — Impugnación formal del despido, requerimiento al empleador "
            "para que garantice la reincorporación en un plazo de 7 días o abone las indemnizaciones económicas correspondientes, "
            "advertencia de recurso ante la autoridad judicial en caso de incumplimiento.\n"
            "5. RESERVAS Y FIRMA — Reserva de reclamar todos los conceptos pendientes no abonados, fecha, firma del trabajador.\n\n"
            f"Usa {ph} para cada dato ausente sin excepción — no inventes ni asumas fechas, nombres, códigos, números ni ningún otro dato no proporcionado explícitamente: usa {ph} para fechas, números de identificación, direcciones y referencias numéricas. "
            f"IMPORTANTE: usa EXCLUSIVAMENTE {ph} para los datos faltantes. "
            f"ESTÁ PROHIBIDO usar marcadores descriptivos como [NOMBRE], [DIRECCIÓN], [FECHA] o similares. "
            f"Solo {ph} y nada más para cualquier dato no proporcionado. "
            "No añadas advertencias meta-legales. "
            "Redacta directamente el documento, sin preámbulos ni explicaciones."
        )
    if lang == "en":
        return (
            "You are an expert in employment law. "
            "Draft a formal dismissal appeal letter to protect the worker's rights. "
            "The document must be formally correct and structured in the following sections:\n\n"
            "1. HEADING — Worker as sender with address, employer as recipient with address, "
            "delivery method (registered letter with acknowledgement of receipt or PEC), subject: dismissal appeal.\n"
            "2. EMPLOYMENT DETAILS — Period of employment from/to, job title and duties, "
            "date and method of dismissal (if only verbal, state so explicitly).\n"
            "3. GROUNDS OF ILLEGITIMACY — Statement of the reasons why the dismissal is considered void and/or unlawful.\n"
            "4. FORMAL APPEAL AND NOTICE — Formal contestation of the dismissal, notice to the employer "
            "to ensure reinstatement within 7 days or pay the applicable financial compensation, "
            "warning of judicial recourse in the event of non-compliance.\n"
            "5. RESERVATIONS AND SIGNATURE — Reservation of the right to claim all outstanding entitlements not yet paid, date, worker's signature.\n\n"
            f"Use {ph} for every missing piece of information without exception — never invent or assume dates, names, codes, numbers or any other data not explicitly provided: use {ph} for dates, tax codes, addresses and reference numbers. "
            f"IMPORTANT: use EXCLUSIVELY {ph} for missing data. "
            f"It is FORBIDDEN to use descriptive placeholders such as [NAME], [ADDRESS], [DATE] or similar. "
            f"Only {ph} and nothing else for any data not provided. "
            "Do not add meta-legal disclaimers. "
            "Write the document directly, without preamble or explanation."
        )
    return (
        "Sei un esperto in diritto del lavoro. "
        "Redigi un'impugnativa formale di licenziamento a tutela dei diritti del lavoratore. "
        "Il documento deve essere formalmente corretto e strutturato nelle seguenti sezioni:\n\n"
        "1. INTESTAZIONE — Lavoratore mittente con indirizzo, datore di lavoro destinatario con indirizzo, "
        "modalità di invio (raccomandata A/R o PEC), oggetto: impugnativa di licenziamento.\n"
        "2. DATI DEL RAPPORTO DI LAVORO — Periodo di lavoro dal/al, qualifica e mansione, "
        "data e modalità del licenziamento (se solo verbale, specificarlo).\n"
        "3. MOTIVI DI ILLEGITTIMITÀ — Esposizione delle ragioni per cui il licenziamento è ritenuto nullo e/o illegittimo.\n"
        "4. IMPUGNATIVA E DIFFIDA — Impugnazione formale del licenziamento, diffida al datore di lavoro "
        "a garantire la reintegrazione entro 7 giorni o a corrispondere le indennità economiche spettanti, "
        "avvertimento di ricorso all'autorità giudiziaria in caso di inottemperanza.\n"
        "5. RISERVE E FIRMA — Riserva di rivendicare tutte le spettanze non corrisposte, data, firma del lavoratore.\n\n"
        f"Usa {ph} per ogni dato mancante senza eccezioni — non inventare né assumere date, nomi, codici, numeri o qualsiasi altro dato non esplicitamente fornito: usa {ph} per date, codici fiscali, indirizzi e riferimenti numerici. "
        f"IMPORTANTE: usa ESCLUSIVAMENTE {ph} per i dati mancanti. "
        f"È VIETATO usare placeholder descrittivi come [NOME], [INDIRIZZO], [DATA] o simili. "
        f"Solo {ph} e nient'altro per qualsiasi dato non fornito. "
        "Non aggiungere avvertenze meta-legali sul documento stesso. "
        "Scrivi direttamente il documento, senza prefazioni o spiegazioni."
    )


def _employment_termination_system(lang: str) -> str:
    ph = _placeholder(lang)
    if lang == "es":
        return (
            "Eres un experto en derecho laboral. "
            "Redacta una carta de despido por causa justificada en nombre del empleador. "
            "El documento debe ser formalmente correcto y estructurado en las siguientes secciones:\n\n"
            "1. ENCABEZAMIENTO — Empleador remitente con dirección, empleado destinatario con dirección, "
            "modalidad de envío (carta certificada con acuse de recibo o entrega en mano), objeto: despido por causa justificada.\n"
            "2. COMUNICACIÓN DE DESPIDO — Comunicación oficial de resolución inmediata de la relación laboral "
            "conforme al art. 54 del Estatuto de los Trabajadores, indicación de la causa justificada específica.\n"
            "3. EFECTOS Y CIERRE — Efecto inmediato del despido, invitación a recoger los efectos personales, "
            "lugar, fecha, firma del empleador.\n\n"
            f"Usa {ph} para cada dato ausente sin excepción — no inventes ni asumas fechas, nombres, códigos, números ni ningún otro dato no proporcionado explícitamente: usa {ph} para fechas, números de identificación, direcciones y referencias numéricas. "
            f"IMPORTANTE: usa EXCLUSIVAMENTE {ph} para los datos faltantes. "
            f"ESTÁ PROHIBIDO usar marcadores descriptivos como [NOMBRE], [DIRECCIÓN], [FECHA] o similares. "
            f"Solo {ph} y nada más para cualquier dato no proporcionado. "
            "No añadas advertencias meta-legales. "
            "Redacta directamente el documento, sin preámbulos ni explicaciones."
        )
    if lang == "en":
        return (
            "You are an expert in employment law. "
            "Draft a letter of dismissal for just cause on behalf of the employer. "
            "The document must be formally correct and structured in the following sections:\n\n"
            "1. HEADING — Employer as sender with address, employee as recipient with address, "
            "delivery method (registered letter with acknowledgement of receipt or hand delivery), subject: dismissal for just cause.\n"
            "2. NOTICE OF DISMISSAL — Official communication of immediate termination of the employment relationship "
            "pursuant to art. 2119 of the Civil Code, specifying the particular just cause.\n"
            "3. EFFECTS AND CLOSING — Immediate effect of the dismissal, invitation to collect personal belongings, "
            "place, date, employer's signature.\n\n"
            f"Use {ph} for every missing piece of information without exception — never invent or assume dates, names, codes, numbers or any other data not explicitly provided: use {ph} for dates, tax codes, addresses and reference numbers. "
            f"IMPORTANT: use EXCLUSIVELY {ph} for missing data. "
            f"It is FORBIDDEN to use descriptive placeholders such as [NAME], [ADDRESS], [DATE] or similar. "
            f"Only {ph} and nothing else for any data not provided. "
            "Do not add meta-legal disclaimers. "
            "Write the document directly, without preamble or explanation."
        )
    return (
        "Sei un esperto in diritto del lavoro. "
        "Redigi una lettera di licenziamento per giusta causa per conto del datore di lavoro. "
        "Il documento deve essere formalmente corretto e strutturato nelle seguenti sezioni:\n\n"
        "1. INTESTAZIONE — Datore di lavoro mittente con indirizzo, dipendente destinatario con indirizzo, "
        "modalità di invio (raccomandata A/R o consegna a mano), oggetto: licenziamento per giusta causa.\n"
        "2. COMUNICAZIONE DI LICENZIAMENTO — Comunicazione ufficiale di risoluzione immediata del rapporto di lavoro "
        "ai sensi dell'art. 2119 c.c., indicazione della giusta causa specifica.\n"
        "3. EFFETTI E CHIUSURA — Effetto immediato del licenziamento, invito a ritirare gli effetti personali, "
        "luogo, data, firma del datore di lavoro.\n\n"
        f"Usa {ph} per ogni dato mancante senza eccezioni — non inventare né assumere date, nomi, codici, numeri o qualsiasi altro dato non esplicitamente fornito: usa {ph} per date, codici fiscali, indirizzi e riferimenti numerici. "
        f"IMPORTANTE: usa ESCLUSIVAMENTE {ph} per i dati mancanti. "
        f"È VIETATO usare placeholder descrittivi come [NOME], [INDIRIZZO], [DATA] o simili. "
        f"Solo {ph} e nient'altro per qualsiasi dato non fornito. "
        "Non aggiungere avvertenze meta-legali sul documento stesso. "
        "Scrivi direttamente il documento, senza prefazioni o spiegazioni."
    )


def _franchising_contract_system(lang: str) -> str:
    ph = _placeholder(lang)
    if lang == "es":
        return (
            "Eres un experto en derecho mercantil y contratos de franquicia. "
            "Redacta un contrato de franquicia completo y formalmente correcto estructurado en las siguientes secciones:\n\n"
            "1. ENCABEZAMIENTO Y CONSIDERANDOS — Franquiciador y franquiciado con datos personales completos, "
            "declaraciones preliminares (sector de actividad, know-how, marca, local comercial del franquiciado).\n"
            "2. OBJETO, EXCLUSIVIDAD Y CANON — Concesión de know-how y marca con exclusividad territorial, "
            "canon de entrada con servicios incluidos, servicios adicionales e importe correspondiente, "
            "autorizaciones intuitu personae no cedibles.\n"
            "3. ROYALTIES Y OBLIGACIONES — Porcentaje de royalties sobre volumen de negocio anual, "
            "facturación mínima anual, obligaciones del franquiciador (know-how, procedimientos, manuales, asistencia, formación), "
            "obligaciones del franquiciado (gestión autónoma, prohibición de traslado de sede, confidencialidad, respeto de la política empresarial).\n"
            "4. DURACIÓN, RESOLUCIÓN Y CONTROLES — Duración del contrato con prórroga tácita y preaviso de 6 meses, "
            "cláusula resolutoria expresa, efectos de la resolución (cese del uso de marca y signos distintivos), "
            "nulidad por información falsa, controles del franquiciador en el punto de venta.\n"
            "5. DISPOSICIONES FINALES Y FIRMAS — Resolución de controversias mediante Cámara de Comercio con foro competente, "
            "tratamiento de datos RGPD, lugar, fecha, firmas del franquiciador y del franquiciado.\n\n"
            f"Usa {ph} para cada dato ausente sin excepción — no inventes ni asumas fechas, nombres, códigos, números ni ningún otro dato no proporcionado explícitamente: usa {ph} para fechas, números de identificación, direcciones y referencias numéricas. "
            f"IMPORTANTE: usa EXCLUSIVAMENTE {ph} para los datos faltantes. "
            f"ESTÁ PROHIBIDO usar marcadores descriptivos como [NOMBRE], [DIRECCIÓN], [FECHA] o similares. "
            f"Solo {ph} y nada más para cualquier dato no proporcionado. "
            "No añadas advertencias meta-legales. "
            "Redacta directamente el documento, sin preámbulos ni explicaciones."
        )
    if lang == "en":
        return (
            "You are an expert in commercial law and franchise agreements. "
            "Draft a complete and formally correct franchise contract structured in the following sections:\n\n"
            "1. HEADING AND RECITALS — Franchisor and franchisee with full personal details, "
            "preliminary declarations (sector of activity, know-how, trademark, franchisee's commercial premises).\n"
            "2. SUBJECT MATTER, EXCLUSIVITY AND FEE — Grant of know-how and trademark with territorial exclusivity, "
            "entry fee with included services, additional services and corresponding amount, "
            "intuitu personae authorisations not transferable.\n"
            "3. ROYALTIES AND OBLIGATIONS — Royalty percentage on annual turnover, "
            "annual minimum turnover, franchisor's obligations (know-how, procedures, manuals, assistance, training), "
            "franchisee's obligations (independent management, prohibition on relocating premises, confidentiality, compliance with company policy).\n"
            "4. DURATION, TERMINATION AND CONTROLS — Contract duration with tacit renewal and 6-month prior notice for termination, "
            "express termination clause, effects of termination (cessation of trademark and distinctive signs), "
            "voidance for false information, franchisor's controls at the point of sale.\n"
            "5. FINAL PROVISIONS AND SIGNATURES — Dispute resolution through the Chamber of Commerce with competent jurisdiction, "
            "GDPR data processing, place, date, signatures of franchisor and franchisee.\n\n"
            f"Use {ph} for every missing piece of information without exception — never invent or assume dates, names, codes, numbers or any other data not explicitly provided: use {ph} for dates, tax codes, addresses and reference numbers. "
            f"IMPORTANT: use EXCLUSIVELY {ph} for missing data. "
            f"It is FORBIDDEN to use descriptive placeholders such as [NAME], [ADDRESS], [DATE] or similar. "
            f"Only {ph} and nothing else for any data not provided. "
            "Do not add meta-legal disclaimers. "
            "Write the document directly, without preamble or explanation."
        )
    return (
        "Sei un esperto in diritto commerciale e contratti di franchising. "
        "Redigi un contratto di franchising completo e formalmente corretto strutturato nelle seguenti sezioni:\n\n"
        "1. INTESTAZIONE E PREMESSE — Franchisor e franchisee con dati anagrafici completi, "
        "dichiarazioni preliminari (settore attività, know-how, marchio, locale commerciale franchisee).\n"
        "2. OGGETTO, ESCLUSIVA E FEE — Concessione know-how e marchio con esclusiva per area, "
        "fee di ingresso con servizi inclusi, servizi ulteriori e relativo importo, "
        "autorizzazioni intuitu personae non cedibili.\n"
        "3. ROYALTIES E OBBLIGHI — Percentuale royalties su volume d'affari annuo, "
        "incasso minimo annuale, obblighi del franchisor (know-how, procedure, manuali, assistenza, formazione), "
        "obblighi del franchisee (gestione autonoma, divieto trasferimento sede, riservatezza, rispetto policy aziendale).\n"
        "4. DURATA, RISOLUZIONE E CONTROLLI — Durata contratto con rinnovo tacito e disdetta 6 mesi prima, "
        "clausola risolutiva espressa, effetti risoluzione (cessazione marchio e segni distintivi), "
        "annullamento per false informazioni, controlli franchisor sul punto vendita.\n"
        "5. DISPOSIZIONI FINALI E FIRME — Risoluzione controversie tramite Camera di Commercio con foro competente, "
        "trattamento dati GDPR, luogo, data, firme franchisor e franchisee.\n\n"
        f"Usa {ph} per ogni dato mancante senza eccezioni — non inventare né assumere date, nomi, codici, numeri o qualsiasi altro dato non esplicitamente fornito: usa {ph} per date, codici fiscali, indirizzi e riferimenti numerici. "
        f"IMPORTANTE: usa ESCLUSIVAMENTE {ph} per i dati mancanti. "
        f"È VIETATO usare placeholder descrittivi come [NOME], [INDIRIZZO], [DATA] o simili. "
        f"Solo {ph} e nient'altro per qualsiasi dato non fornito. "
        "Non aggiungere avvertenze meta-legali sul documento stesso. "
        "Scrivi direttamente il documento, senza prefazioni o spiegazioni."
    )


def _verbale_assemblea_system(lang: str) -> str:
    ph = _placeholder(lang)
    if lang == "es":
        return (
            "Eres un experto en derecho de la propiedad horizontal y comunidades de propietarios. "
            "Redacta un acta de junta de propietarios formalmente correcta y estructurada en las siguientes secciones:\n\n"
            "1. ENCABEZAMIENTO — Nombre de la comunidad de propietarios, domicilio, fecha, hora y lugar de celebracion.\n"
            "2. ASISTENTES Y QUORUM — Lista de propietarios asistentes o representados con cuotas de participacion, "
            "comprobacion del quorum necesario, nombre del presidente y del secretario de la junta.\n"
            "3. ORDEN DEL DIA — Enumeracion de los puntos del orden del dia tal como fueron convocados.\n"
            "4. DESARROLLO Y ACUERDOS — Para cada punto del orden del dia: resumen del debate, "
            "acuerdo adoptado con resultado de la votacion (votos a favor, en contra, abstenciones), "
            "indicacion de si el acuerdo es vinculante conforme a la LPH.\n"
            "5. CIERRE Y FIRMAS — Hora de cierre, firma del presidente y del secretario.\n\n"
            f"Usa {ph} para cada dato ausente sin excepcion - no inventes ni asumas fechas, nombres, cuotas ni ningun otro dato. "
            f"IMPORTANTE: usa EXCLUSIVAMENTE {ph} para los datos faltantes. "
            f"ESTÁ PROHIBIDO usar marcadores descriptivos como [NOMBRE], [DIRECCIÓN], [FECHA] o similares. "
            f"Solo {ph} y nada más para cualquier dato no proporcionado. "
            "No anadas advertencias meta-legales. "
            "Redacta directamente el acta, sin preambulos ni explicaciones."
        )
    if lang == "en":
        return (
            "You are an expert in condominium and homeowners association law. "
            "Draft a formally correct condominium assembly minutes document structured in the following sections:\n\n"
            "1. HEADING — Name of the condominium, address, date, time and place of the meeting.\n"
            "2. ATTENDEES AND QUORUM — List of owners present or represented with their ownership shares, "
            "confirmation that the required quorum is met, name of the chairperson and secretary of the meeting.\n"
            "3. AGENDA — Numbered list of agenda items as notified in the convening notice.\n"
            "4. PROCEEDINGS AND RESOLUTIONS — For each agenda item: summary of the discussion, "
            "resolution passed with voting result (votes in favour, against, abstentions), "
            "indication of whether the resolution is binding.\n"
            "5. CLOSING AND SIGNATURES — Time of closure, signatures of the chairperson and secretary.\n\n"
            f"Use {ph} for every missing piece of information without exception - never invent or assume dates, names, ownership shares or any other data. "
            f"IMPORTANT: use EXCLUSIVELY {ph} for missing data. "
            f"It is FORBIDDEN to use descriptive placeholders such as [NAME], [ADDRESS], [DATE] or similar. "
            f"Only {ph} and nothing else for any data not provided. "
            "Do not add meta-legal disclaimers. "
            "Write the document directly, without preamble or explanation."
        )
    return (
        "Sei un esperto in diritto condominiale italiano. "
        "Redigi un verbale di assemblea condominiale formalmente corretto e strutturato nelle seguenti sezioni:\n\n"
        "1. INTESTAZIONE — Denominazione del condominio, indirizzo, data, ora e luogo di svolgimento dell'assemblea.\n"
        "2. PRESENTI E QUORUM — Elenco dei condominii presenti o rappresentati con relative quote millesimali, "
        "verifica del quorum costitutivo, nome del presidente e del segretario dell'assemblea.\n"
        "3. ORDINE DEL GIORNO — Elencazione dei punti all'ordine del giorno come convocati.\n"
        "4. SVOLGIMENTO E DELIBERE — Per ogni punto all'ordine del giorno: sintesi della discussione, "
        "delibera assunta con esito della votazione (voti favorevoli, contrari, astenuti), "
        "indicazione se la delibera e vincolante ai sensi del codice civile.\n"
        "5. CHIUSURA E FIRME — Ora di chiusura, firma del presidente e del segretario.\n\n"
        f"Usa {ph} per ogni dato mancante senza eccezioni - non inventare ne assumere date, nomi, quote millesimali o qualsiasi altro dato. "
        f"IMPORTANTE: usa ESCLUSIVAMENTE {ph} per i dati mancanti. "
        f"È VIETATO usare placeholder descrittivi come [NOME], [INDIRIZZO], [DATA] o simili. "
        f"Solo {ph} e nient'altro per qualsiasi dato non fornito. "
        "Non aggiungere avvertenze meta-legali sul documento stesso. "
        "Scrivi direttamente il verbale, senza prefazioni o spiegazioni."
    )


def _demand_letter_system(lang: str) -> str:
    ph = _placeholder(lang)
    if lang == "es":
        return (
            "Eres un abogado experto en derecho civil y mercantil. "
            "Redacta una carta de requerimiento o burofax formal estructurada en las siguientes secciones:\n\n"
            "1. ENCABEZAMIENTO — Datos del remitente (nombre, domicilio), datos del destinatario "
            "(nombre, domicilio), modalidad de envio (burofax con acuse de recibo o carta certificada), "
            "objeto de la comunicacion.\n"
            "2. ANTECEDENTES DE HECHO — Descripcion concisa y cronologica de los hechos relevantes "
            "que fundamentan el requerimiento.\n"
            "3. FUNDAMENTO JURIDICO — Base legal de la reclamacion con referencia a los principios "
            "generales del derecho aplicables.\n"
            "4. REQUERIMIENTO — Descripcion precisa de lo que se exige al destinatario: cumplimiento "
            "de la obligacion, pago de cantidad, cese de conducta, entrega de documentacion u otro.\n"
            "5. PLAZO Y CONSECUENCIAS — Plazo concedido para el cumplimiento (dias habiles desde la "
            "recepcion), advertencia expresa de las acciones legales que se emprendera en caso de "
            "incumplimiento (reclamacion judicial, denuncia, ejercicio de otros derechos).\n"
            "6. FIRMA Y FECHA — Lugar, fecha, firma del remitente o su representante legal.\n\n"
            f"Usa {ph} para cada dato ausente sin excepcion - no inventes ni asumas fechas, nombres, "
            "importes ni ningun otro dato no proporcionado explicitamente. "
            f"IMPORTANTE: usa EXCLUSIVAMENTE {ph} para los datos faltantes. "
            f"ESTÁ PROHIBIDO usar marcadores descriptivos como [NOMBRE], [DIRECCIÓN], [FECHA] o similares. "
            f"Solo {ph} y nada más para cualquier dato no proporcionado. "
            "No anadas advertencias meta-legales. "
            "Redacta directamente la carta, sin preambulos ni explicaciones."
        )
    if lang == "en":
        return (
            "You are an expert in civil and commercial law. "
            "Draft a formal legal demand letter or notice structured in the following sections:\n\n"
            "1. HEADING — Sender's details (name, address), recipient's details (name, address), "
            "method of delivery (registered letter with acknowledgement of receipt or certified email), "
            "subject of the communication.\n"
            "2. STATEMENT OF FACTS — Concise chronological description of the relevant facts "
            "underpinning the demand.\n"
            "3. LEGAL BASIS — Legal foundation of the claim with reference to applicable general "
            "principles of law.\n"
            "4. DEMAND — Precise description of what is required from the recipient: performance of "
            "an obligation, payment of a sum, cessation of conduct, delivery of documents or other.\n"
            "5. DEADLINE AND CONSEQUENCES — Time allowed for compliance (working days from receipt), "
            "express warning of the legal action that will be taken in case of non-compliance "
            "(judicial claim, complaint, exercise of other rights).\n"
            "6. SIGNATURE AND DATE — Place, date, signature of the sender or their legal representative.\n\n"
            f"Use {ph} for every missing piece of information without exception - never invent or assume "
            "dates, names, amounts or any other data not explicitly provided. "
            f"IMPORTANT: use EXCLUSIVELY {ph} for missing data. "
            f"It is FORBIDDEN to use descriptive placeholders such as [NAME], [ADDRESS], [DATE] or similar. "
            f"Only {ph} and nothing else for any data not provided. "
            "Do not add meta-legal disclaimers. "
            "Write the document directly, without preamble or explanation."
        )
    return (
        "Sei un avvocato esperto in diritto civile e commerciale italiano. "
        "Redigi una lettera di diffida e messa in mora formalmente corretta e strutturata nelle seguenti sezioni:\n\n"
        "1. INTESTAZIONE — Dati del mittente (nome, cognome/ragione sociale, indirizzo), dati del "
        "destinatario (nome, cognome/ragione sociale, indirizzo), modalita di invio (raccomandata A/R "
        "o PEC), oggetto della comunicazione.\n"
        "2. PREMESSE IN FATTO — Descrizione concisa e cronologica dei fatti rilevanti che fondano "
        "la diffida.\n"
        "3. FONDAMENTO GIURIDICO — Base legale della pretesa con riferimento ai principi generali "
        "del diritto applicabili.\n"
        "4. DIFFIDA E RICHIESTA — Descrizione precisa di cio che si esige dal destinatario: "
        "adempimento dell'obbligazione, pagamento di somma, cessazione di condotta, consegna di "
        "documentazione o altro.\n"
        "5. TERMINE E CONSEGUENZE — Termine concesso per l'adempimento (giorni dalla ricezione), "
        "avvertimento espresso delle azioni legali che saranno intraprese in caso di inottemperanza "
        "(azione giudiziaria, esposto, esercizio di altri diritti).\n"
        "6. FIRMA E DATA — Luogo, data, firma del mittente o del suo rappresentante legale.\n\n"
        f"Usa {ph} per ogni dato mancante senza eccezioni - non inventare ne assumere date, nomi, "
        "importi o qualsiasi altro dato non esplicitamente fornito. "
        f"IMPORTANTE: usa ESCLUSIVAMENTE {ph} per i dati mancanti. "
        f"È VIETATO usare placeholder descrittivi come [NOME], [INDIRIZZO], [DATA] o simili. "
        f"Solo {ph} e nient'altro per qualsiasi dato non fornito. "
        "Non aggiungere avvertenze meta-legali sul documento stesso. "
        "Scrivi direttamente la lettera, senza prefazioni o spiegazioni."
    )


def _appeal_system(lang: str) -> str:
    ph = _placeholder(lang)
    if lang == "es":
        return (
            "Eres un abogado experto en derecho procesal civil y administrativo. "
            "Redacta un recurso o escrito de apelacion formal estructurado en las siguientes secciones:\n\n"
            "1. ENCABEZAMIENTO — Organo jurisdiccional o autoridad destinataria, recurrente con datos "
            "completos y representacion letrada, parte recurrida con datos completos, numero de "
            "procedimiento y resolucion impugnada.\n"
            "2. ANTECEDENTES DE HECHO — Descripcion concisa del procedimiento previo, de la resolucion "
            "impugnada y de los hechos relevantes.\n"
            "3. FUNDAMENTOS DE DERECHO — Normas procesales que habilitan el recurso, jurisprudencia "
            "aplicable y principios generales del derecho en que se apoya el recurrente.\n"
            "4. MOTIVOS DEL RECURSO — Enumeracion y desarrollo de cada motivo de impugnacion: "
            "infracciones de norma procesal o sustantiva, error en la valoracion de la prueba, "
            "incongruencia u otras irregularidades.\n"
            "5. SUPLICO — Peticion precisa al organo: admision del recurso, revocacion o anulacion "
            "de la resolucion impugnada, resolucion conforme a derecho, condena en costas si procede.\n"
            "6. FIRMA Y FECHA — Lugar, fecha, firma del letrado.\n\n"
            f"Usa {ph} para cada dato ausente sin excepcion - no inventes ni asumas fechas, nombres, "
            "numeros de procedimiento ni ningun otro dato no proporcionado explicitamente. "
            f"IMPORTANTE: usa EXCLUSIVAMENTE {ph} para los datos faltantes. "
            f"ESTÁ PROHIBIDO usar marcadores descriptivos como [NOMBRE], [DIRECCIÓN], [FECHA] o similares. "
            f"Solo {ph} y nada más para cualquier dato no proporcionado. "
            "No anadas advertencias meta-legales. "
            "Redacta directamente el escrito, sin preambulos ni explicaciones."
        )
    if lang == "en":
        return (
            "You are an expert in civil and administrative procedural law. "
            "Draft a formal appeal or recourse document structured in the following sections:\n\n"
            "1. HEADING — Court or authority addressed, appellant with full details and legal "
            "representation, respondent with full details, case number and decision being appealed.\n"
            "2. STATEMENT OF FACTS — Concise account of the prior proceedings, the decision being "
            "challenged, and the relevant facts.\n"
            "3. LEGAL GROUNDS — Procedural rules authorising the appeal, applicable case law and "
            "general principles of law relied upon by the appellant.\n"
            "4. GROUNDS OF APPEAL — Enumeration and development of each ground of challenge: "
            "breaches of procedural or substantive rules, errors in the assessment of evidence, "
            "inconsistency or other irregularities.\n"
            "5. RELIEF SOUGHT — Precise request to the court: admission of the appeal, revocation "
            "or annulment of the challenged decision, decision in accordance with the law, "
            "costs award where applicable.\n"
            "6. SIGNATURE AND DATE — Place, date, lawyer's signature.\n\n"
            f"Use {ph} for every missing piece of information without exception - never invent or assume "
            "dates, names, case numbers or any other data not explicitly provided. "
            f"IMPORTANT: use EXCLUSIVELY {ph} for missing data. "
            f"It is FORBIDDEN to use descriptive placeholders such as [NAME], [ADDRESS], [DATE] or similar. "
            f"Only {ph} and nothing else for any data not provided. "
            "Do not add meta-legal disclaimers. "
            "Write the document directly, without preamble or explanation."
        )
    return (
        "Sei un avvocato esperto in diritto processuale civile e amministrativo italiano. "
        "Redigi un ricorso o atto di appello formalmente corretto e strutturato nelle seguenti sezioni:\n\n"
        "1. INTESTAZIONE — Autorita giudiziaria o amministrativa destinataria, ricorrente con dati "
        "completi e difensore, resistente con dati completi, numero di procedimento e provvedimento "
        "impugnato.\n"
        "2. PREMESSE IN FATTO — Descrizione concisa del procedimento pregresso, del provvedimento "
        "impugnato e dei fatti rilevanti.\n"
        "3. FONDAMENTI DI DIRITTO — Norme processuali che abilitano l'impugnazione, giurisprudenza "
        "applicabile e principi generali del diritto su cui si fonda il ricorrente.\n"
        "4. MOTIVI DEL RICORSO — Enumerazione e sviluppo di ciascun motivo di impugnazione: "
        "violazioni di norma processuale o sostanziale, errore nella valutazione delle prove, "
        "contraddittorieta o altre irregolarita.\n"
        "5. PETITUM — Richiesta precisa all'organo: ammissione del ricorso, revoca o annullamento "
        "del provvedimento impugnato, pronuncia conforme a diritto, condanna alle spese ove dovuta.\n"
        "6. FIRMA E DATA — Luogo, data, firma del difensore.\n\n"
        f"Usa {ph} per ogni dato mancante senza eccezioni - non inventare ne assumere date, nomi, "
        "numeri di procedimento o qualsiasi altro dato non esplicitamente fornito. "
        f"IMPORTANTE: usa ESCLUSIVAMENTE {ph} per i dati mancanti. "
        f"È VIETATO usare placeholder descrittivi come [NOME], [INDIRIZZO], [DATA] o simili. "
        f"Solo {ph} e nient'altro per qualsiasi dato non fornito. "
        "Non aggiungere avvertenze meta-legali sul documento stesso. "
        "Scrivi direttamente l'atto, senza prefazioni o spiegazioni."
    )


def _power_of_attorney_system(lang: str) -> str:
    ph = _placeholder(lang)
    if lang == "es":
        return (
            "Eres un notario o abogado experto en derecho notarial y representacion legal. "
            "Redacta una escritura de poder notarial o documento de delegacion formal estructurado "
            "en las siguientes secciones:\n\n"
            "1. ENCABEZAMIENTO — Titulo del documento (poder especial / poder general / delegacion), "
            "lugar y fecha de otorgamiento.\n"
            "2. DATOS DEL PODERDANTE — Nombre completo, DNI/NIE/NIF, domicilio, capacidad legal "
            "en que actua.\n"
            "3. DATOS DEL APODERADO — Nombre completo, DNI/NIE/NIF, domicilio, capacidad legal.\n"
            "4. OBJETO Y PODERES CONFERIDOS — Descripcion precisa del objeto de la delegacion y "
            "enumeracion detallada de los actos que el apoderado queda facultado a realizar en "
            "nombre del poderdante.\n"
            "5. LIMITACIONES Y DURACION — Eventuales limitaciones al ejercicio del poder, duracion "
            "(indefinida o con fecha de vencimiento), condiciones de revocacion y modalidades.\n"
            "6. FIRMA Y AUTENTICACION — Firma del poderdante, firma y sello del notario o fedatario "
            "publico si requerido, datos del notario y numero de protocolo.\n\n"
            f"Usa {ph} para cada dato ausente sin excepcion - no inventes ni asumas nombres, numeros "
            "de identificacion, poderes ni ningun otro dato no proporcionado explicitamente. "
            f"IMPORTANTE: usa EXCLUSIVAMENTE {ph} para los datos faltantes. "
            f"ESTÁ PROHIBIDO usar marcadores descriptivos como [NOMBRE], [DIRECCIÓN], [FECHA] o similares. "
            f"Solo {ph} y nada más para cualquier dato no proporcionado. "
            "No anadas advertencias meta-legales. "
            "Redacta directamente el documento, sin preambulos ni explicaciones."
        )
    if lang == "en":
        return (
            "You are a notary or expert in notarial law and legal representation. "
            "Draft a power of attorney or formal delegation document structured in the following sections:\n\n"
            "1. HEADING — Title of the document (special power of attorney / general power of attorney "
            "/ delegation), place and date of execution.\n"
            "2. PRINCIPAL'S DETAILS — Full name, ID/tax number, address, legal capacity in which acting.\n"
            "3. ATTORNEY'S DETAILS — Full name, ID/tax number, address, legal capacity.\n"
            "4. SCOPE AND POWERS GRANTED — Precise description of the subject matter of the delegation "
            "and detailed enumeration of the acts the attorney is authorised to perform on behalf "
            "of the principal.\n"
            "5. LIMITATIONS AND DURATION — Any limitations on the exercise of the power, duration "
            "(indefinite or with expiry date), conditions and manner of revocation.\n"
            "6. SIGNATURE AND AUTHENTICATION — Principal's signature, notary's signature and seal "
            "if required, notary's details and deed number.\n\n"
            f"Use {ph} for every missing piece of information without exception - never invent or assume "
            "names, identification numbers, powers or any other data not explicitly provided. "
            f"IMPORTANT: use EXCLUSIVELY {ph} for missing data. "
            f"It is FORBIDDEN to use descriptive placeholders such as [NAME], [ADDRESS], [DATE] or similar. "
            f"Only {ph} and nothing else for any data not provided. "
            "Do not add meta-legal disclaimers. "
            "Write the document directly, without preamble or explanation."
        )
    return (
        "Sei un notaio o avvocato esperto in diritto notarile e rappresentanza legale italiano. "
        "Redigi una procura speciale o generale formalmente corretta e strutturata nelle seguenti sezioni:\n\n"
        "1. INTESTAZIONE — Titolo del documento (procura speciale / procura generale / delega), "
        "luogo e data del conferimento.\n"
        "2. DATI DEL CONFERENTE — Nome e cognome completi, codice fiscale, domicilio, qualita "
        "in cui agisce.\n"
        "3. DATI DEL PROCURATORE — Nome e cognome completi, codice fiscale, domicilio, qualita.\n"
        "4. OGGETTO E POTERI CONFERITI — Descrizione precisa dell'oggetto della delega ed "
        "enumerazione dettagliata degli atti che il procuratore e autorizzato a compiere in nome "
        "e per conto del conferente.\n"
        "5. LIMITAZIONI E DURATA — Eventuali limitazioni all'esercizio della procura, durata "
        "(a tempo indeterminato o con data di scadenza), condizioni e modalita di revoca.\n"
        "6. FIRMA E AUTENTICAZIONE — Firma del conferente, firma e timbro del notaio o pubblico "
        "ufficiale se richiesto, dati del notaio e numero di repertorio.\n\n"
        f"Usa {ph} per ogni dato mancante senza eccezioni - non inventare ne assumere nomi, codici "
        "fiscali, poteri o qualsiasi altro dato non esplicitamente fornito. "
        f"IMPORTANTE: usa ESCLUSIVAMENTE {ph} per i dati mancanti. "
        f"È VIETATO usare placeholder descrittivi come [NOME], [INDIRIZZO], [DATA] o simili. "
        f"Solo {ph} e nient'altro per qualsiasi dato non fornito. "
        "Non aggiungere avvertenze meta-legali sul documento stesso. "
        "Scrivi direttamente il documento, senza prefazioni o spiegazioni."
    )


def _sale_agreement_system(lang: str) -> str:
    ph = _placeholder(lang)
    if lang == "es":
        return (
            "Eres un abogado experto en derecho civil patrimonial y contratos de compraventa. "
            "Redacta un contrato de compraventa formalmente correcto y estructurado en las siguientes secciones:\n\n"
            "1. ENCABEZAMIENTO — Vendedor y comprador con datos completos (nombre, DNI/NIF, domicilio, "
            "capacidad legal), lugar y fecha de firma del contrato.\n"
            "2. OBJETO DE LA VENTA — Descripcion detallada del bien vendido (bien mueble o inmueble, "
            "referencias identificativas, estado de conservacion, pertenencias incluidas y excluidas).\n"
            "3. PRECIO Y MODALIDADES DE PAGO — Precio total de venta, modalidad de pago (contado, "
            "plazos, transferencia bancaria, cheque), fecha o fechas de pago, consecuencias del "
            "impago.\n"
            "4. GARANTIAS Y DECLARACIONES — Declaraciones del vendedor sobre titularidad, ausencia "
            "de cargas y gravamenes, garantia por eviccion y vicios ocultos, obligaciones de "
            "saneamiento conforme a la normativa aplicable.\n"
            "5. ENTREGA Y TRANSMISION DE PROPIEDAD — Fecha y modalidades de entrega del bien, "
            "momento del traspaso de la propiedad y del riesgo, obligaciones de las partes hasta "
            "la entrega.\n"
            "6. CLAUSULAS FINALES Y FIRMAS — Resolucion de controversias, ley aplicable y fuero "
            "competente, gastos e impuestos a cargo de cada parte, firmas de vendedor y comprador.\n\n"
            f"Usa {ph} para cada dato ausente sin excepcion - no inventes ni asumas precios, datos "
            "identificativos del bien ni ningun otro dato no proporcionado explicitamente. "
            f"IMPORTANTE: usa EXCLUSIVAMENTE {ph} para los datos faltantes. "
            f"ESTÁ PROHIBIDO usar marcadores descriptivos como [NOMBRE], [DIRECCIÓN], [FECHA] o similares. "
            f"Solo {ph} y nada más para cualquier dato no proporcionado. "
            "No anadas advertencias meta-legales. "
            "Redacta directamente el contrato, sin preambulos ni explicaciones."
        )
    if lang == "en":
        return (
            "You are an expert in property law and sale and purchase agreements. "
            "Draft a formally correct sale and purchase agreement structured in the following sections:\n\n"
            "1. HEADING — Seller and buyer with full details (name, ID/tax number, address, legal "
            "capacity), place and date of signing.\n"
            "2. SUBJECT MATTER — Detailed description of the goods sold (movable or immovable property, "
            "identifying references, condition, included and excluded accessories).\n"
            "3. PRICE AND PAYMENT TERMS — Total sale price, payment method (lump sum, instalments, "
            "bank transfer, cheque), payment date or dates, consequences of non-payment.\n"
            "4. WARRANTIES AND REPRESENTATIONS — Seller's representations as to title, freedom from "
            "encumbrances, warranty against eviction and latent defects, remedies under applicable law.\n"
            "5. DELIVERY AND TRANSFER OF TITLE — Date and manner of delivery, moment of transfer of "
            "ownership and risk, obligations of the parties until delivery.\n"
            "6. FINAL CLAUSES AND SIGNATURES — Dispute resolution, governing law and jurisdiction, "
            "costs and taxes at each party's expense, signatures of seller and buyer.\n\n"
            f"Use {ph} for every missing piece of information without exception - never invent or assume "
            "prices, identifying details of the goods or any other data not explicitly provided. "
            f"IMPORTANT: use EXCLUSIVELY {ph} for missing data. "
            f"It is FORBIDDEN to use descriptive placeholders such as [NAME], [ADDRESS], [DATE] or similar. "
            f"Only {ph} and nothing else for any data not provided. "
            "Do not add meta-legal disclaimers. "
            "Write the document directly, without preamble or explanation."
        )
    return (
        "Sei un avvocato esperto in diritto civile patrimoniale e contratti di compravendita italiani. "
        "Redigi un contratto di compravendita formalmente corretto e strutturato nelle seguenti sezioni:\n\n"
        "1. INTESTAZIONE — Venditore e acquirente con dati completi (nome e cognome/ragione sociale, "
        "codice fiscale/P.IVA, domicilio, qualita in cui agisce), luogo e data di stipula.\n"
        "2. OGGETTO DELLA VENDITA — Descrizione dettagliata del bene venduto (bene mobile o immobile, "
        "riferimenti identificativi, stato di conservazione, pertinenze incluse ed escluse).\n"
        "3. PREZZO E MODALITA DI PAGAMENTO — Prezzo totale di vendita, modalita di pagamento "
        "(contante, rate, bonifico bancario, assegno), data o date di pagamento, conseguenze "
        "del mancato pagamento.\n"
        "4. GARANZIE E DICHIARAZIONI — Dichiarazioni del venditore su titolarita, assenza di "
        "vincoli e gravami, garanzia per evizione e vizi occulti, obblighi di risanamento ai "
        "sensi della normativa applicabile.\n"
        "5. CONSEGNA E TRASFERIMENTO DELLA PROPRIETA — Data e modalita di consegna del bene, "
        "momento del passaggio di proprieta e del rischio, obblighi delle parti fino alla consegna.\n"
        "6. CLAUSOLE FINALI E FIRME — Risoluzione delle controversie, legge applicabile e foro "
        "competente, spese e imposte a carico di ciascuna parte, firme di venditore e acquirente.\n\n"
        f"Usa {ph} per ogni dato mancante senza eccezioni - non inventare ne assumere prezzi, dati "
        "identificativi del bene o qualsiasi altro dato non esplicitamente fornito. "
        f"IMPORTANTE: usa ESCLUSIVAMENTE {ph} per i dati mancanti. "
        f"È VIETATO usare placeholder descrittivi come [NOME], [INDIRIZZO], [DATA] o simili. "
        f"Solo {ph} e nient'altro per qualsiasi dato non fornito. "
        "Non aggiungere avvertenze meta-legali sul documento stesso. "
        "Scrivi direttamente il contratto, senza prefazioni o spiegazioni."
    )


def _nota_contestazione_system(lang: str) -> str:
    ph = _placeholder(lang)
    if lang == "es":
        return (
            "Eres un experto en derecho administrativo y de tráfico. "
            "Redacta una nota formal de contestación a una denuncia de tráfico, "
            "declarando la no participación del remitente en el incidente. "
            "El documento debe ser formalmente correcto y estructurado en las siguientes secciones:\n\n"
            "1. ENCABEZAMIENTO — Datos del remitente (nombre, apellidos, C.F./D.N.I., domicilio), "
            "autoridad destinataria con dirección, número de referencia de la denuncia, fecha de la denuncia.\n"
            "2. REFERENCIA A LA DENUNCIA RECIBIDA — Indicación precisa de la denuncia recibida: "
            "número de referencia, fecha, matrícula del vehículo implicado según la denuncia, "
            "infracción o hecho contestado.\n"
            "3. DECLARACIÓN DE NO PARTICIPACIÓN — Declaración expresa y formal de no haber estado "
            "presente ni haber participado en el incidente o infracción indicados, "
            "con alusión a la matrícula o circunstancias erróneas.\n"
            "4. DESCRIPCIÓN DE LOS HECHOS — Exposición clara y cronológica de los hechos tal y como "
            "se desarrollaron según el remitente: lugar donde se encontraba, vehículo utilizado, "
            "cualquier circunstancia relevante que excluya su participación.\n"
            "5. SOLICITUD DE ARCHIVO — Solicitud formal de archivo de las actuaciones por falta de "
            "fundamento, con reserva de ejercitar las acciones legales oportunas en caso de "
            "continuación indebida del procedimiento.\n"
            "6. FIRMA — Fórmula de cierre, lugar, fecha, firma del remitente.\n\n"
            f"Usa {ph} para cada dato ausente sin excepción — no inventes ni asumas fechas, matrículas, "
            "nombres, números de expediente ni ningún otro dato no proporcionado explícitamente. "
            f"IMPORTANTE: usa EXCLUSIVAMENTE {ph} para los datos faltantes. "
            f"ESTÁ PROHIBIDO usar marcadores descriptivos como [NOMBRE], [DIRECCIÓN], [FECHA] o similares. "
            f"Solo {ph} y nada más para cualquier dato no proporcionado. "
            "No añadas advertencias meta-legales. "
            "Redacta directamente la nota, sin preámbulos ni explicaciones."
        )
    if lang == "en":
        return (
            "You are an expert in administrative and traffic law. "
            "Draft a formal letter contesting a road traffic dispute notice, "
            "asserting the sender's non-involvement in the incident. "
            "The document must be formally correct and structured in the following sections:\n\n"
            "1. HEADING — Sender's details (full name, tax/ID number, address), "
            "authority as addressee with address, reference number of the notice, date of the notice.\n"
            "2. REFERENCE TO THE NOTICE RECEIVED — Precise identification of the notice received: "
            "reference number, date, vehicle registration plate cited in the notice, "
            "infraction or event being contested.\n"
            "3. DECLARATION OF NON-INVOLVEMENT — Express and formal declaration that the sender was "
            "not present and did not participate in the incident or infraction indicated, "
            "with reference to the erroneous registration or circumstances.\n"
            "4. DESCRIPTION OF FACTS — Clear chronological account of events as they actually occurred "
            "according to the sender: location at the time, vehicle used, any relevant circumstances "
            "excluding their participation.\n"
            "5. REQUEST FOR DISMISSAL — Formal request to close the proceedings for lack of grounds, "
            "with reservation of all legal rights should the proceedings continue improperly.\n"
            "6. SIGNATURE — Closing formula, place, date, sender's signature.\n\n"
            f"Use {ph} for every missing piece of information without exception — never invent or assume "
            "dates, registration plates, names, case numbers or any other data not explicitly provided. "
            f"IMPORTANT: use EXCLUSIVELY {ph} for missing data. "
            f"It is FORBIDDEN to use descriptive placeholders such as [NAME], [ADDRESS], [DATE] or similar. "
            f"Only {ph} and nothing else for any data not provided. "
            "Do not add meta-legal disclaimers. "
            "Write the document directly, without preamble or explanation."
        )
    return (
        "Sei un esperto in diritto amministrativo e della circolazione stradale. "
        "Redigi una nota formale alla contestazione di un verbale stradale, "
        "dichiarando il non coinvolgimento del mittente nell'incidente. "
        "Il documento deve essere formalmente corretto e strutturato nelle seguenti sezioni:\n\n"
        "1. INTESTAZIONE — Dati del mittente (nome, cognome, C.F., indirizzo), "
        "ente contestante destinatario con indirizzo, numero di riferimento della contestazione, "
        "data della contestazione.\n"
        "2. RIFERIMENTO ALLA CONTESTAZIONE RICEVUTA — Indicazione precisa della contestazione ricevuta: "
        "numero di riferimento, data, targa del veicolo indicata nel verbale, "
        "infrazione o fatto contestato.\n"
        "3. DICHIARAZIONE DI NON COINVOLGIMENTO — Dichiarazione espressa e formale di non essere "
        "stati presenti né di aver preso parte all'incidente o all'infrazione indicati, "
        "con riferimento alla targa o alle circostanze erronee.\n"
        "4. DESCRIZIONE DEI FATTI — Esposizione chiara e cronologica dei fatti così come si sono "
        "svolti secondo il mittente: luogo in cui si trovava, veicolo utilizzato, "
        "ogni circostanza rilevante che escluda il proprio coinvolgimento.\n"
        "5. RICHIESTA DI ARCHIVIAZIONE — Richiesta formale di archiviazione del procedimento per "
        "infondatezza, con riserva di esercitare ogni azione legale opportuna in caso di "
        "ingiustificata prosecuzione del procedimento.\n"
        "6. FIRMA — Formula di chiusura, luogo, data, firma del mittente.\n\n"
        f"Usa {ph} per ogni dato mancante senza eccezioni — non inventare né assumere date, targhe, "
        "nomi, numeri di protocollo o qualsiasi altro dato non esplicitamente fornito. "
        f"IMPORTANTE: usa ESCLUSIVAMENTE {ph} per i dati mancanti. "
        f"È VIETATO usare placeholder descrittivi come [NOME], [INDIRIZZO], [DATA] o simili. "
        f"Solo {ph} e nient'altro per qualsiasi dato non fornito. "
        "Non aggiungere avvertenze meta-legali sul documento stesso. "
        "Scrivi direttamente la nota, senza prefazioni o spiegazioni."
    )


_SYSTEM_TEMPLATES_CATALOG_PATH = os.getenv(
    "SYSTEM_TEMPLATES_CATALOG_PATH",
    "/opt/chatbot/data/system_templates/catalog_enriched.json",
)


def _load_system_templates_catalog() -> List[Dict[str, Any]]:
    try:
        with open(_SYSTEM_TEMPLATES_CATALOG_PATH, encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        logger.warning(f"Could not load system templates catalog from {_SYSTEM_TEMPLATES_CATALOG_PATH}: {e}")
        return []


SYSTEM_TEMPLATES_CATALOG: List[Dict[str, Any]] = _load_system_templates_catalog()
SYSTEM_TEMPLATES_BY_KEY: Dict[str, Dict[str, Any]] = {
    (entry["filename"][:-5] if entry["filename"].endswith(".docx") else entry["filename"]): entry
    for entry in SYSTEM_TEMPLATES_CATALOG
}
logger.info(f"Loaded {len(SYSTEM_TEMPLATES_CATALOG)} system templates from catalog")


_CODICE_DESCRIPTIONS = {
    "Codice di procedura civile": "atti processuali civili: citazioni, ricorsi, memorie, opposizioni, esecuzioni, tutele cautelari",
    "Codice di procedura penale": "atti processuali penali: difese, istanze, ricorsi, memorie penali, misure cautelari penali",
    "Codice del processo amministrativo": "atti processuali amministrativi: ricorsi TAR, istanze sospensive, ottemperanza, accesso agli atti",
    "Contratti e Atti Stragiudiziali": "contratti privati e atti non processuali: locazioni, polizze, lavoro, compravendite, franchising, diffide, messa in mora, verbali, contestazioni",
    "Arbitrato e Procedure Alternative": "Arbitrato rituale e irrituale, lodi arbitrali, clausole arbitrali, ADR, mediazione arbitrale",
}


def _normalize_tokens(s: str) -> set:
    """Lowercase, strip punctuation, split to tokens, drop stopwords."""
    _STOP = {
        "di", "del", "della", "dei", "degli", "delle", "a", "ad", "al",
        "alla", "alle", "agli", "ai", "da", "dal", "dalla", "dai", "dagli",
        "dalle", "in", "nel", "nella", "nei", "negli", "nelle", "su", "sul",
        "sulla", "sui", "sugli", "sulle", "con", "per", "tra", "fra",
        "e", "o", "il", "lo", "la", "i", "gli", "le", "un", "uno", "una",
        "che", "non", "si", "è", "uso", "ad",
    }
    tokens = re.sub(r"[^a-z0-9\s]", " ", s.lower()).split()
    return {t for t in tokens if t not in _STOP and len(t) > 1}


def _overlap_score(a: str, b: str) -> float:
    """Jaccard-style overlap between token sets of two strings."""
    ta, tb = _normalize_tokens(a), _normalize_tokens(b)
    if not ta or not tb:
        return 0.0
    return len(ta & tb) / len(ta | tb)


def _best_categoria_for_codice(
    message: str, matched_codice: str, entries_in_codice: List[Dict[str, Any]]
) -> str | None:
    """Stage 2: classify which categoria (within one codice) the message belongs to."""
    categorie = sorted({c for e in entries_in_codice for c in e["categorie"]})
    categoria_system = (
        f"Sei un classificatore di richieste di atti giuridici italiani. "
        f"L'utente sta richiedendo un atto del {matched_codice}. Determina "
        "a quale fase/categoria processuale appartiene la richiesta.\n\n"
        "Categorie disponibili:\n"
        + "\n".join(f"- {c}" for c in categorie)
        + "\n\nREGOLE DI CLASSIFICAZIONE:\n"
        "- Se l'utente NON specifica la fase processuale (es. non menziona "
        "'dibattimento', 'udienza', 'appello', 'cassazione', 'indagini'), "
        "scegli la categoria più generale, di solito '1. ATTI GENERALI DEL "
        "DIFENSORE (TUTTE LE PARTI)' se disponibile\n"
        "- Solo se l'utente menziona esplicitamente una fase specifica "
        "(es. 'per il dibattimento', 'in appello', 'durante le indagini'), "
        "scegli la categoria corrispondente a quella fase\n"
        "- In caso di dubbio, scegli sempre la categoria più generale\n\n"
        "Restituisci SOLO il nome esatto della categoria, nient'altro."
    )
    try:
        categoria_result = _call_chat(
            [SystemMessage(content=categoria_system), HumanMessage(content=message)],
            max_tokens=40,
        ).strip()
    except Exception as e:
        logger.warning(f"classify_system_template stage 2 failed for {matched_codice!r}: {e}")
        return None
    return next(
        (c for c in categorie if c.lower() in categoria_result.lower() or categoria_result.lower() in c.lower()),
        None,
    )


_STEM_SUFFIXES = ("zioni", "zione", "menti", "mento", "ali", "ale")


def _stem_word(word: str) -> str:
    """Strip a common Italian noun/adjective ending, only if the remaining stem is > 3 chars."""
    for suffix in _STEM_SUFFIXES:
        if word.endswith(suffix) and len(word) - len(suffix) > 3:
            return word[: -len(suffix)]
    stem = re.sub(r"[oaie]s?$", "", word)
    if stem != word and len(stem) > 3:
        return stem
    return word


def _stem_text(s: str) -> str:
    """Lowercase, strip punctuation, collapse whitespace, then stem each word."""
    text = re.sub(r"[^a-zà-ÿ0-9\s]", " ", s.lower())
    words = re.sub(r"\s+", " ", text).strip().split()
    return " ".join(_stem_word(w) for w in words)


def _regex_match_catalog(message: str) -> List[Dict[str, Any]]:
    """
    Broad, no-LLM candidate matching: stems every word (> 3 chars) of the user
    message and, separately, each catalog entry's tipo_atto and label, then
    keeps any entry where at least one stemmed query word appears as a
    substring of the stemmed tipo_atto or label.
    """
    query_stems = {w for w in _stem_text(message).split() if len(w) > 3}
    if not query_stems:
        return []

    matches = []
    for entry in SYSTEM_TEMPLATES_CATALOG:
        tipo_stemmed = _stem_text(entry.get("tipo_atto", ""))
        label_stemmed = _stem_text(entry.get("label", ""))
        if any(qs in tipo_stemmed or qs in label_stemmed for qs in query_stems):
            matches.append({
                "filename": entry["filename"],
                "tipo_atto": entry.get("tipo_atto", ""),
                "label": entry.get("label", ""),
                "codice": entry.get("codice", ""),
                "description": entry.get("description", ""),
            })
    return matches


def _llm_rank_candidates(message: str, candidates: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Single LLM call to rank/filter the regex-matched candidates by relevance
    to the user's message. On call failure or unparseable output, falls back
    to all candidates sorted alphabetically by label.
    """
    options_text = "\n".join(
        f"{i}. {c['tipo_atto']}: {c.get('description', '')}"
        for i, c in enumerate(candidates, start=1)
    )
    system = (
        "Sei un assistente legale italiano. L'utente vuole generare un "
        "documento legale. Dato il messaggio dell'utente e la lista di tipi "
        "di documento disponibili, restituisci SOLO un array JSON con i "
        "numeri (1-based) di TUTTI i documenti in ordine di "
        "rilevanza, dal più al meno pertinente. Includi tutti i documenti "
        "nell'array, anche quelli meno pertinenti. Esempio: [2, 5, 1, 3, 4]\n\n"
        "Documenti disponibili:\n" + options_text
    )
    try:
        raw = _call_chat(
            [SystemMessage(content=system), HumanMessage(content=message)],
            max_tokens=200,
        ).strip()
        json_match = re.search(r"\[.*\]", raw, re.DOTALL)
        indices = json.loads(json_match.group(0) if json_match else raw)
        if not isinstance(indices, list):
            raise ValueError(f"LLM ranking response is not a JSON array: {raw!r}")
    except Exception as e:
        logger.warning(f"classify_system_template LLM ranking failed: {e}")
        return sorted(candidates, key=lambda c: c.get("label", ""))

    return [
        candidates[idx - 1]
        for idx in indices
        if isinstance(idx, int) and 1 <= idx <= len(candidates)
    ]


def _slugify(s: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", s.lower()).strip("-")


# Applied (in order) to a title-cased sublabel to restore legal-citation
# punctuation that title-casing a hyphenated slug strips out. Word-boundary
# matched so "Cpp"/"Cpc"/"Cp" (distinct whole tokens) never collide.
_SUBLABEL_ABBREVIATIONS = [
    (re.compile(r"\bCpp\b"), "c.p.p."),
    (re.compile(r"\bCpc\b"), "c.p.c."),
    (re.compile(r"\bCp\b"), "c.p."),
    (re.compile(r"\bArt\b"), "Art."),
]

# Tried longest-first against the START of remainder only, so "atto-di-" is
# stripped whole rather than leaving a dangling "di-" behind.
_SUBLABEL_PREFIX_STRIP = (
    "atto-di-", "atto-", "per-", "di-", "del-", "della-", "delle-", "degli-",
)

# Italian prepositions/articles that .title() wrongly capitalises when they
# appear mid-sublabel; left capitalised only as the first word.
_SUBLABEL_LOWERCASE_WORDS = {
    "di", "del", "della", "dei", "degli", "delle", "per", "con", "su", "sul",
    "sulla", "in", "nel", "nella", "e", "o", "a", "al", "alla",
}


def _derive_sublabel(key: str, label: str) -> str:
    """
    Derive a sublabel from a catalog key by stripping the codice prefix (the part
    before "__") and then the base label slug from the start of what remains.
    E.g. key "penale__memoria-difensiva-ex-art-415-bis-cpp" with label "Memoria
    difensiva" -> sublabel "Ex Art. 415 Bis c.p.p.". "Generale" if nothing
    remains (the entry IS the base template).
    """
    remainder = key.split("__", 1)[1] if "__" in key else key
    for prefix in _SUBLABEL_PREFIX_STRIP:
        if remainder.startswith(prefix):
            remainder = remainder[len(prefix):]
            break
    label_slug = _slugify(label)
    if label_slug and remainder.startswith(label_slug):
        remainder = remainder[len(label_slug):].lstrip("-")
    if not remainder:
        return "Generale"

    words = remainder.replace("-", " ").title().split(" ")
    sublabel = " ".join(
        w.lower() if i > 0 and w.lower() in _SUBLABEL_LOWERCASE_WORDS else w
        for i, w in enumerate(words)
    )
    for pattern, replacement in _SUBLABEL_ABBREVIATIONS:
        sublabel = pattern.sub(replacement, sublabel)
    # Letter-by-letter slugs (e.g. "c-p-p") survive the loop above as "C P P"
    # since each letter is its own token -- collapse those too. Longest first
    # so " C P P"/" C P C" don't get partially matched by " C P".
    for old, new in ((" C P P", " c.p.p."), (" C P C", " c.p.c."), (" C P", " c.p.")):
        sublabel = sublabel.replace(old, new)
    return sublabel


def classify_system_template(
    message: str, lang: str, top_k: int = 1
) -> str | List[Dict[str, Any]]:
    """
    3-stage LLM classification against the 383-template system catalog
    (codice -> categoria -> tipo di atto). Used as a fallback when
    classify_document_type() returns "unknown" for the legacy 20-type
    registry. Returns the matched template's key (catalog filename
    without .docx extension) or "unknown".

    When top_k > 1, runs a different pipeline: a no-LLM regex/stemming pass
    (_regex_match_catalog) finds every catalog entry that shares a stemmed
    word with the message. 0 matches -> empty list. Exactly 1 match ->
    returned directly as a single-item list (score 1.0). 2+ matches -> one
    LLM call (_llm_rank_candidates) ranks/filters them by relevance; results
    are returned in that order as {"key", "label", "codice", "sublabel",
    "score"} dicts, with score 1.0 for the first result and -0.1 per
    subsequent rank (floored at 0.0). top_k == 1 (the default) is completely
    unchanged.
    """
    if not SYSTEM_TEMPLATES_CATALOG:
        return [] if top_k > 1 else "unknown"

    codici = sorted({e["codice"] for e in SYSTEM_TEMPLATES_CATALOG})

    if top_k > 1:
        matches = _regex_match_catalog(message)
        logger.info("DEBUG stage1 regex matches: %s", [m["tipo_atto"] for m in matches])
        if not matches:
            return []
        if len(matches) == 1:
            only = matches[0]
            fname = only["filename"]
            key = fname[:-5] if fname.endswith(".docx") else fname
            label = only.get("label", "")
            return [{
                "key": key,
                "label": label,
                "codice": only["codice"],
                "sublabel": _derive_sublabel(key, label),
                "score": 1.0,
            }]

        ranked = _llm_rank_candidates(message, sorted(matches, key=lambda m: m.get("label", ""))[:30])
        results = []
        for i, entry in enumerate(ranked):
            fname = entry["filename"]
            key = fname[:-5] if fname.endswith(".docx") else fname
            label = entry.get("label", "")
            results.append({
                "key": key,
                "label": label,
                "codice": entry["codice"],
                "sublabel": _derive_sublabel(key, label),
                "score": max(1.0 - 0.1 * i, 0.0),
            })
        return results

    codici_lines = "\n".join(
        f"- {c}: {_CODICE_DESCRIPTIONS.get(c, '')}" for c in codici
    )
    codice_system = (
        "Sei un classificatore di richieste di documenti legali italiani. "
        "Determina a quale categoria appartiene la richiesta dell'utente "
        "scegliendo tra quelle disponibili.\n\n"
        "Categorie disponibili:\n"
        + codici_lines
        + "\n\nRestituisci SOLO il nome esatto della categoria, nient'altro."
    )
    try:
        codice_result = _call_chat(
            [SystemMessage(content=codice_system), HumanMessage(content=message)],
            max_tokens=30,
        ).strip()
    except Exception as e:
        logger.warning(f"classify_system_template stage 1 failed: {e}")
        return "unknown"

    matched_codice = next(
        (c for c in codici if c.lower() in codice_result.lower() or codice_result.lower() in c.lower()),
        None,
    )
    if not matched_codice:
        return "unknown"

    entries_in_codice = [e for e in SYSTEM_TEMPLATES_CATALOG if e["codice"] == matched_codice]
    matched_categoria = _best_categoria_for_codice(message, matched_codice, entries_in_codice)
    if not matched_categoria:
        return "unknown"

    entries_in_categoria = [e for e in entries_in_codice if matched_categoria in e["categorie"]]
    if len(entries_in_categoria) == 1:
        only = entries_in_categoria[0]["filename"]
        return only[:-5] if only.endswith(".docx") else only

    options_text = "\n".join(f"- {e['tipo_atto']}: {e['description']}" for e in entries_in_categoria)
    tipo_system = (
        f"Sei un classificatore di richieste di atti giuridici italiani. "
        f"L'utente sta richiedendo un atto della categoria "
        f"'{matched_categoria}'. Determina esattamente quale tipo di atto "
        "sta richiedendo.\n\nTipi disponibili:\n"
        + options_text
        + "\n\nREGOLE DI CLASSIFICAZIONE:\n"
        "- Se il tipo è espresso al singolare (es. 'Memoria difensiva') "
        "e l'utente usa il singolare, preferisci il tipo singolare\n"
        "- I tipi al plurale (es. 'Memorie difensive') si usano quando "
        "l'utente vuole un atto per la fase dibattimentale o specifica "
        "esplicitamente il plurale\n"
        "- Se l'utente non specifica la fase processuale, scegli il tipo "
        "più generico e applicabile (di solito il singolare)\n"
        "- art. 121 c.p.p. si applica a memorie difensive generali "
        "presentate al PM o al GIP, non al dibattimento\n\n"
        "Restituisci SOLO il nome esatto del tipo di atto, nient'altro."
    )
    try:
        tipo_result = _call_chat(
            [SystemMessage(content=tipo_system), HumanMessage(content=message)],
            max_tokens=40,
        ).strip()
    except Exception as e:
        logger.warning(f"classify_system_template stage 3 failed: {e}")
        return "unknown"

    # Fast path: exact substring match (handles perfect LLM output)
    matched_entry = next(
        (
            e for e in entries_in_categoria
            if e["tipo_atto"].lower() in tipo_result.lower()
            or tipo_result.lower() in e["tipo_atto"].lower()
        ),
        None,
    )

    # Fuzzy fallback: pick catalog entry with highest token overlap
    if not matched_entry:
        scored = [
            (e, _overlap_score(tipo_result, e["tipo_atto"]))
            for e in entries_in_categoria
        ]
        best_entry, best_score = max(scored, key=lambda x: x[1])
        if best_score >= 0.35:
            logger.info(
                "classify_system_template: fuzzy matched %r → %r (score=%.2f)",
                tipo_result, best_entry["tipo_atto"], best_score,
            )
            matched_entry = best_entry
        else:
            logger.warning(
                "classify_system_template: no match for tipo_result=%r "
                "(best candidate %r scored %.2f < 0.35)",
                tipo_result, best_entry["tipo_atto"], best_score,
            )
            return "unknown"

    fname = matched_entry["filename"]
    return fname[:-5] if fname.endswith(".docx") else fname


def _build_system_template_prompt(entry: Dict[str, Any], lang: str) -> str:
    ph = _placeholder(lang)
    sections = entry.get("sections", [])
    sections_text = "\n\n".join(
        f"{s['heading']}\n{s['content']}"
        for s in sections
        if s.get("heading") and s["heading"] != "Campi strutturati"
    )
    label = entry.get("label") or entry.get("tipo_atto", "")
    codice = entry.get("codice", "")

    if lang == "es":
        lang_instruction = (
            "IMPORTANTE: el documento sigue siendo un acto procesal italiano, "
            "presentado ante un tribunal italiano y regido por el derecho "
            "italiano. NO cambies la jurisdicción ni las referencias legales. "
            "Redacta el TEXTO en español, manteniendo intactas (sin traducir) "
            "las referencias a tribunales, códigos y normativa italiana."
        )
    elif lang == "en":
        lang_instruction = (
            "IMPORTANT: this remains an Italian procedural act, filed with an "
            "Italian court and governed by Italian law. Do NOT change the "
            "jurisdiction or legal references. Write the TEXT in English, "
            "while keeping references to Italian courts, codes, and statutes "
            "untranslated."
        )
    else:
        lang_instruction = ""

    return (
        f"Sei un avvocato esperto di diritto processuale italiano con almeno 20 anni di esperienza. "
        f"Devi redigere un atto giuridico italiano completo e professionale del tipo '{label}' "
        f"({codice}).\n\n"
        f"OBIETTIVO: produrre un documento legale vero, completo e formalmente corretto, "
        f"non uno schema da compilare. Ogni sezione deve contenere testo legale autentico, "
        f"formulazioni tecniche appropriate e il linguaggio giuridico italiano preciso "
        f"tipico di questo tipo di atto.\n\n"
        f"STRUTTURA E CONTENUTO DI RIFERIMENTO:\n"
        f"Di seguito la struttura dell'atto con il contenuto di riferimento per ogni sezione. "
        f"Usala come guida per capire cosa deve contenere ogni sezione, "
        f"ma scrivi testo legale completo e professionale — non limitarti a copiare o riempire i placeholder:\n\n"
        f"{sections_text}\n\n"
        f"ISTRUZIONI PER I DATI:\n"
        f"- Usa i dati forniti dall'utente per personalizzare il documento\n"
        f"- REGOLA ASSOLUTA ANTI-ALLUCINAZIONE: se un dato non è esplicitamente presente "
        f"nel messaggio dell'utente o nei campi strutturati, scrivi letteralmente {ph} — "
        f"MAI inventare o dedurre indirizzi, date, nomi, numeri, importi, codici fiscali, "
        f"riferimenti normativi specifici o qualsiasi altro dato non fornito esplicitamente. "
        f"Un documento con {ph} visibili è sempre preferibile a uno con dati inventati.\n"
        f"- Usa tutti i dettagli fattuali presenti nel messaggio originale (nomi, date, luoghi, "
        f"circostanze, testimoni, alibi) sviluppandoli nel contenuto giuridico appropriato\n"
        f"- Sviluppa ogni sezione con tutto il contenuto legale necessario per questo tipo di atto\n"
        f"- Usa formule e clausole appropriate ma NON citare articoli di legge specifici "
        f"a meno che non siano esplicitamente menzionati dall'utente\n"
        f"- Scrivi come un avvocato che redige il documento per un cliente reale\n"
        f"- NON aggiungere note esplicative, commenti o paragrafi dopo la fine dell'atto\n"
        f"- Il documento termina con la formula di sottoscrizione — niente altro dopo\n"
        f"- Per i difensori usa sempre 'Avv.' o 'Avvocato', mai 'Dott.' o 'Dott.ssa'\n\n"
        f"{lang_instruction}"
    )


def extract_system_template_fields(user_message: str, entry: Dict[str, Any], lang: str) -> Dict[str, str]:
    fields = entry.get("fields", [])
    label = entry.get("label") or entry.get("tipo_atto", "")

    if lang == "es":
        system = (
            f"Extrae del mensaje del usuario los valores para los siguientes campos del documento '{label}'.\n"
            f"Devuelve SOLO un objeto JSON válido, sin texto antes ni después, sin markdown.\n"
            f"Campos requeridos: {json.dumps(fields, ensure_ascii=False)}\n"
            "Reglas:\n"
            "- Si un valor está explícitamente mencionado en el mensaje, úsalo exactamente\n"
            "- Deduce valores clara e inequívocamente deducibles del contexto "
            "(ej: si el usuario escribe \"Mario Rossi como arrendatario\", asigna "
            "\"Mario Rossi\" a cualquier campo cuyo nombre contenga \"arrendatario\"; "
            "si tras el nombre de una persona se indica una fecha de nacimiento, "
            "asígnala al campo de fecha de nacimiento de esa persona)\n"
            "- Si el mensaje contiene varios sujetos con roles distintos (ej. arrendador "
            "y arrendatario), asigna los datos de cada sujeto a los campos "
            "correspondientes a su rol — no mezcles datos entre sujetos distintos\n"
            "- Para los campos de fecha, normaliza al formato DD/MM/AAAA si es posible\n"
            "- Si un valor no está presente ni es deducible, usa cadena vacía \"\"\n"
            f"Responde SOLO con el JSON, ejemplo: {{\"campo1\": \"valor1\", \"campo2\": \"\"}}"
        )
    elif lang == "en":
        system = (
            f"Extract from the user message the values for the following fields of the document '{label}'.\n"
            f"Return ONLY a valid JSON object, with no text before or after, no markdown.\n"
            f"Required fields: {json.dumps(fields, ensure_ascii=False)}\n"
            "Rules:\n"
            "- If a value is explicitly mentioned in the message, use it exactly\n"
            "- Infer values that are clearly and unambiguously deducible from context "
            "(e.g. if the user writes \"Mario Rossi as tenant\", map \"Mario Rossi\" to "
            "any field whose name contains \"tenant\"; if a birth date is given right "
            "after a person's name, assign it to that person's birth date field)\n"
            "- If the message contains multiple subjects with distinct roles (e.g. "
            "landlord and tenant), assign each subject's data to the fields matching "
            "their role — do not mix data between different subjects\n"
            "- For date fields, normalize to DD/MM/YYYY format where possible\n"
            "- If a value is neither present nor inferable, use empty string \"\"\n"
            f"Reply ONLY with the JSON, example: {{\"field1\": \"value1\", \"field2\": \"\"}}"
        )
    else:
        system = (
            f"Estrai dal messaggio dell'utente i valori per i seguenti campi del documento '{label}'.\n"
            f"Restituisci SOLO un oggetto JSON valido, senza testo prima o dopo, senza markdown.\n"
            f"Campi richiesti: {json.dumps(fields, ensure_ascii=False)}\n"
            "Regole:\n"
            "- Se un valore è esplicitamente menzionato nel messaggio, usalo esattamente\n"
            "- Deduci valori chiaramente e inequivocabilmente ricavabili dal contesto "
            "(es. se l'utente scrive \"Mario Rossi come conduttore\", assegna \"Mario Rossi\" "
            "a qualsiasi campo il cui nome contiene \"conduttore\"; se dopo il nome di una "
            "persona è indicata una data di nascita, assegnala al campo data di nascita "
            "di quella persona)\n"
            "- Se il messaggio contiene più soggetti con ruoli distinti (es. locatore e "
            "conduttore), assegna i dati di ciascun soggetto ai campi corrispondenti al "
            "suo ruolo — non mescolare i dati tra soggetti diversi\n"
            "- Per i campi data, normalizza nel formato GG/MM/AAAA se possibile\n"
            "- Se un valore non è presente né deducibile, usa stringa vuota \"\"\n"
            f"Rispondi SOLO con il JSON, esempio: {{\"campo1\": \"valore1\", \"campo2\": \"\"}}"
        )

    human = f"Messaggio dell'utente:\n{user_message}"
    raw = _call_chat(
        [SystemMessage(content=system), HumanMessage(content=human)],
        max_tokens=800,
    )

    text = re.sub(r"```(?:json)?\s*", "", raw).strip().rstrip("`").strip()
    try:
        parsed = json.loads(text)
    except (json.JSONDecodeError, ValueError):
        logger.warning("extract_system_template_fields: JSON parse failed for label=%r, raw=%r", label, raw[:200])
        retry_response = _call_chat(
            [
                SystemMessage(content="Rispondi SOLO con JSON valido, nessun testo aggiuntivo."),
                HumanMessage(content=f"Estrai questi campi: {fields}\nDal testo: {user_message}\nJSON:"),
            ],
            max_tokens=400,
        )
        clean_retry = re.sub(r"```(?:json)?\s*", "", retry_response).strip().rstrip("`").strip()
        try:
            parsed = json.loads(clean_retry)
        except (json.JSONDecodeError, ValueError):
            logger.error("extract_system_template_fields: retry also failed for label=%r", label)
            parsed = {}

    return {k: str(parsed.get(k, "") or "") for k in fields}


_PLACEHOLDER_RUN_PATTERN = re.compile(
    r'_{3,}|\[.+?\]|\(.+?\)|OMISSIS|\b_+\b', re.IGNORECASE
)


def _extract_docx_elements(path: str) -> List[Dict]:
    """Extract non-empty text elements from a DOCX with sequential indices.
    Processes body paragraphs first, then table cells."""
    from docx import Document as _D
    doc = _D(path)
    elements: List[Dict] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            elements.append({"index": len(elements), "text": text, "type": "paragraph"})
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        elements.append({"index": len(elements), "text": text, "type": "cell"})
    return elements


def _extract_pdf_elements(path: str) -> List[Dict]:
    """Extract non-empty text lines from a PDF file."""
    if not _PDF_SUPPORT:
        raise RuntimeError("pypdf not installed; PDF templates not supported")
    reader = _PdfReader(path)
    elements: List[Dict] = []
    for page in reader.pages:
        for line in (page.extract_text() or "").split("\n"):
            line = line.strip()
            if line:
                elements.append({"index": len(elements), "text": line, "type": "line"})
    return elements


def _is_heading_paragraph(para) -> bool:
    """A paragraph counts as a section heading if it's short and either
    styled as a heading, written in ALL CAPS, or fully bold."""
    text = para.text.strip()
    if not text or len(text) >= 100:
        return False
    style_name = (para.style.name or "") if para.style else ""
    if style_name.lower().startswith("heading") or style_name.lower() == "title":
        return True
    if text.isupper() and any(c.isalpha() for c in text):
        return True
    runs_with_text = [r for r in para.runs if r.text.strip()]
    if runs_with_text and all(r.bold for r in runs_with_text):
        return True
    return False


def _select_relevant_section(docx_path: str, user_message: str, lang: str) -> str | None:
    """Identify the DOCX section heading most relevant to the user's request
    via a single LLM call. Returns None if the document has fewer than 2
    identifiable headings (use the whole document), or on failure."""
    from docx import Document as _D
    try:
        doc = _D(docx_path)
    except Exception as e:
        logger.warning("_select_relevant_section: could not open %r: %s", docx_path, e)
        return None

    headings: List[str] = []
    for para in doc.paragraphs:
        if _is_heading_paragraph(para):
            text = para.text.strip()
            if text not in headings:
                headings.append(text)

    if len(headings) < 2:
        return None

    system = (
        "Sei un assistente legale. Il documento contiene più sezioni. "
        "Basandoti sulla richiesta dell'utente, restituisci SOLO il testo "
        "esatto dell'intestazione della sezione più pertinente, senza nessun altro testo."
    )
    headings_list = "\n".join(f"{i}. {h}" for i, h in enumerate(headings, 1))
    human = f"Richiesta utente: {user_message}\n\nIntestazioni disponibili:\n{headings_list}"

    try:
        result = _call_chat(
            [SystemMessage(content=system), HumanMessage(content=human)],
            max_tokens=100,
        ).strip()
    except Exception as e:
        logger.warning("_select_relevant_section: LLM call failed: %s", e)
        return None

    for h in headings:
        if h.lower() == result.lower():
            return h
    for h in headings:
        if h.lower() in result.lower() or result.lower() in h.lower():
            return h
    return None


def _normalize_for_hint_match(s: str) -> str:
    """Lowercase and strip common punctuation for loose substring matching."""
    return re.sub(r"[^\w\s]", "", s, flags=re.UNICODE).lower().strip()


def _find_heading_by_hint(docx_path: str, section_hint: str) -> str | None:
    """Find a DOCX paragraph whose text contains section_hint, case-insensitive
    and ignoring common punctuation. Returns the paragraph's exact text, or None."""
    from docx import Document as _D
    try:
        doc = _D(docx_path)
    except Exception as e:
        logger.warning("_find_heading_by_hint: could not open %r: %s", docx_path, e)
        return None

    normalized_hint = _normalize_for_hint_match(section_hint)
    if not normalized_hint:
        return None

    for para in doc.paragraphs:
        text = para.text.strip()
        if text and normalized_hint in _normalize_for_hint_match(text):
            return text
    return None


def _extract_section_content(docx_path: str, heading: str) -> str | None:
    """Extract the text of one section — the matching heading paragraph
    through the paragraph before the next heading — from a DOCX template."""
    from docx import Document as _D
    try:
        doc = _D(docx_path)
    except Exception as e:
        logger.warning("_extract_section_content: could not open %r: %s", docx_path, e)
        return None

    paragraphs = doc.paragraphs
    start_idx = next(
        (i for i, p in enumerate(paragraphs) if p.text.strip().lower() == heading.strip().lower()),
        None,
    )
    if start_idx is None:
        return None

    section_lines = [paragraphs[start_idx].text.strip()]
    for para in paragraphs[start_idx + 1:]:
        if _is_heading_paragraph(para):
            break
        text = para.text.strip()
        if text:
            section_lines.append(text)

    return "\n".join(section_lines)


def _fill_template_gaps(
    elements: List[Dict],
    user_message: str,
    carta_intestata: Optional[Dict],
    lang: str,
    session_messages: List[Dict],
    docx_path: Optional[str] = None,
) -> Dict[int, str]:
    """Ask the LLM to identify blanks in the template elements and fill them.
    Returns {element_index: replacement_text} — sparse, only changed elements."""
    ph = _placeholder(lang)

    carta_parts: List[str] = []
    if carta_intestata:
        for key, label in [
            ("legal_name", "Nome studio"),
            ("address_street", "Indirizzo"),
            ("address_city", "Città"),
            ("vat_number", "P.IVA"),
            ("phone", "Tel"),
            ("website", "Web"),
        ]:
            val = carta_intestata.get(key)
            if val:
                carta_parts.append(f"{label}: {val}")
    carta_text = "\n".join(carta_parts)

    # Capped to stay within the max_tokens budget of the _call_chat below —
    # each element line adds to the prompt, and the model also has to echo
    # indices back in its JSON response.
    elements_text = "\n".join(
        f"{e['index']}. {e['text']}" for e in elements[:600]
    )

    lang_note = {
        "es": "Rellena los espacios en español (no cambies el idioma del documento).",
        "en": "Fill the blanks in English (do not change the document language).",
    }.get(lang, "")

    system = (
        "Sei un assistente che compila documenti legali italiani. "
        "Ti viene fornito un elenco numerato di elementi testuali di un documento. "
        "Alcuni elementi contengono spazi vuoti da compilare — possono essere: "
        "trattini bassi (___), parentesi quadre [NOME], parentesi tonde (nome), "
        "spazio vuoto dopo un'etichetta (es. 'Locatore:' senza valore), OMISSIS, "
        "o semplicemente mancanza del dato atteso.\n\n"
        "Regole:\n"
        "1. Identifica gli elementi con spazi vuoti.\n"
        "2. Compila usando le informazioni del messaggio utente e i dati dello studio.\n"
        "2b. Per identificare il dato corretto, ragiona sul significato dell'etichetta "
        "contestuale (es. 'Locatore:' indica il proprietario, 'Conduttore:' indica "
        "l'inquilino, 'Data:' indica una data rilevante dal contesto).\n"
        f"3. Usa '{ph}' per qualsiasi dato non disponibile.\n"
        "4. Non includere elementi che non necessitano di modifica.\n"
        "5. Non inventare dati non forniti esplicitamente.\n"
        + (f"6. {lang_note}\n" if lang_note else "")
        + "\nRestituisci SOLO un oggetto JSON valido:\n"
        "{\"indice\": \"testo_completo_sostituito\", ...}\n"
        "Dove 'indice' è il numero dell'elemento e 'testo_completo_sostituito' "
        "è il testo completo della riga compilata (non solo il valore inserito)."
    )

    history = [m for m in (session_messages or []) if m.get("role") != "system"]
    if history and history[-1].get("content") == user_message:
        history = history[:-1]
    history = history[-10:]
    context_text = "\n".join(f"{m.get('role', '')}: {m.get('content', '')}" for m in history)

    relevant_section_text = None
    if docx_path:
        selected_heading = _select_relevant_section(docx_path, user_message, lang)
        if selected_heading:
            relevant_section_text = _extract_section_content(docx_path, selected_heading)
            if relevant_section_text:
                relevant_section_text = f"{selected_heading}\n\n{relevant_section_text}"

    human_parts = [
        f"Contesto conversazione precedente:\n{context_text}\n\nMessaggio attuale:\n{user_message}"
    ]
    if carta_text:
        human_parts.append(
            f"Dati dello studio (per campi relativi al firmatario/studio):\n{carta_text}"
        )
    human_parts.append(f"Elementi del documento:\n{elements_text}")
    if relevant_section_text:
        human_parts.insert(
            0,
            f"Testo esatto da usare come base:\n{relevant_section_text}\n\n"
            "COPIA questo testo ESATTAMENTE come appare. Sostituisci SOLO i segnaposto "
            "(__________, [DA COMPILARE], spazi vuoti dopo etichette) con i dati forniti "
            "dall'utente. NON riscrivere, NON riformulare, NON aggiungere contenuto. "
            "Se un dato non è disponibile, lascia il segnaposto invariato.",
        )

    raw = _call_chat(
        [SystemMessage(content=system), HumanMessage(content="\n\n".join(human_parts))],
        max_tokens=4000,
    )
    text = re.sub(r"```(?:json)?\s*", "", raw).strip().rstrip("`").strip()
    try:
        parsed = json.loads(text)
        return {int(k): str(v) for k, v in parsed.items()}
    except (json.JSONDecodeError, ValueError):
        logger.warning("_fill_template_gaps: JSON parse failed, raw=%r", raw[:200])
        raise HTTPException(
            status_code=422,
            detail="Il modello non ha restituito un JSON valido — riprovare.",
        )


def _apply_fill_to_docx(source_path: str, fill_map: Dict[int, str]) -> bytes:
    """Apply fill_map to a DOCX file and return the modified bytes.
    Preserves formatting: tries run-level replacement first, falls
    back to rewriting the first run only if no placeholder run is found."""
    from docx import Document as _D
    doc = _D(source_path)

    all_paras: List = []
    for para in doc.paragraphs:
        if para.text.strip():
            all_paras.append(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        all_paras.append(para)

    for idx, replacement in fill_map.items():
        if idx >= len(all_paras):
            continue
        para = all_paras[idx]
        replaced = False
        for run in para.runs:
            if _PLACEHOLDER_RUN_PATTERN.search(run.text) or not run.text.strip():
                run.text = replacement
                replaced = True
                break
        if not replaced and para.runs:
            para.runs[0].text = replacement
            for run in para.runs[1:]:
                run.text = ""

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _build_docx_from_pdf_elements(
    elements: List[Dict], fill_map: Dict[int, str]
) -> bytes:
    """Build a new DOCX from PDF-extracted elements with gap-filling applied.
    Used when the source was a PDF — cannot preserve original PDF formatting."""
    from docx import Document as _D
    doc = _D()
    for element in elements:
        idx = element["index"]
        text = fill_map.get(idx, element["text"])
        if len(text) < 80 and text.upper() == text and text.strip():
            doc.add_heading(text, level=2)
        else:
            doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def generate_document(
    user_message: str,
    doc_type: str,
    lang: str,
    citations: list = None,
    studio_name: str = "",
    section_hint: str = "",
) -> Dict[str, Any]:
    """Generic document generator dispatching on doc_type via DOCUMENT_TYPE_REGISTRY.

    Returns:
        {"draft": str, "case_details": Dict[str, str], "doc_type": str}
    """
    _SYSTEM_FN: dict = {}

    system_fn = _SYSTEM_FN.get(doc_type)
    system_template_entry = None if system_fn else SYSTEM_TEMPLATES_BY_KEY.get(doc_type)
    relevant_section_text = None

    if system_fn is not None:
        system_text = system_fn(lang)
        fields_dict = extract_document_fields(user_message, doc_type, lang)
    elif system_template_entry is not None:
        filename = system_template_entry.get("filename", "")
        if filename:
            docx_path = os.path.join(
                os.path.dirname(_SYSTEM_TEMPLATES_CATALOG_PATH), filename
            )
            if section_hint:
                selected_heading = _find_heading_by_hint(docx_path, section_hint)
            else:
                selected_heading = _select_relevant_section(docx_path, user_message, lang)
            if selected_heading:
                relevant_section_text = _extract_section_content(docx_path, selected_heading)
                if relevant_section_text:
                    relevant_section_text = f"{selected_heading}\n\n{relevant_section_text}"
        system_text = _build_system_template_prompt(system_template_entry, lang)
        fields_dict = extract_system_template_fields(user_message, system_template_entry, lang)
    else:
        logger.error(
            "generate_document: doc_type=%r not found in _SYSTEM_FN or SYSTEM_TEMPLATES_BY_KEY "
            "— catalog key mismatch. Raising ValueError instead of falling back to opposition.",
            doc_type,
        )
        raise ValueError(
            f"Tipo di documento '{doc_type}' non trovato nel catalogo. "
            "Riprova descrivendo il documento in modo più specifico."
        )

    _FORMATTING_RULE = (
        "FORMATTING RULE: Use consistent, professional capitalization throughout. "
        "Document titles and section headers should use Title Case or ALL CAPS consistently. "
        "Never mix uppercase and lowercase within a single word "
        "(e.g. never write 'LOCazione' — write 'Locazione' or 'LOCAZIONE'). "
        "All placeholder text in brackets should be in UPPERCASE: "
        "[DA COMPILARE], [INDIRIZZO IMMOBILE], etc.\n\n"
    )
    system_content = _FORMATTING_RULE + system_text
    if citations:
        citations_text = "\n".join(
            f"- {cit.get('document_name', '')}, sezione {sec.get('name', '')}"
            for cit in citations
            for sec in (cit.get("sections") or [])
            if sec.get("name")
        )
        if citations_text:
            if system_template_entry is not None:
                system_content += (
                    "\nGround the legal arguments in these specific retrieved documents and cite "
                    "them explicitly where the document argues legal grounds (e.g. in fatto / "
                    "in diritto / motivi sections):\n" + citations_text
                )
            else:
                system_content += (
                    "\nGround the legal arguments in these specific retrieved documents and cite them "
                    "explicitly in the Motivi di opposizione section:\n" + citations_text
                )

    ph = _placeholder(lang)
    field_lines = "\n".join(
        f"- {k}: {v if v else ph}"
        for k, v in fields_dict.items()
    )
    human_content = (
        f"Messaggio originale dell'utente:\n{user_message}\n\n"
        f"Campi strutturati estratti:\n{field_lines}\n\n"
        f"Redigi l'atto completo. Per ogni sezione scrivi testo legale vero e professionale — "
        f"usa tutti i dettagli fattuali forniti nel messaggio originale (nomi, date, luoghi, "
        f"circostanze, testimoni, alibi, ecc.) sviluppandoli nel contenuto giuridico appropriato. "
        f"Usa {ph} SOLO per dati non presenti né nel messaggio né nei campi strutturati. "
        f"Non aggiungere note, spiegazioni o avvertenze fuori dall'atto."
    )
    if relevant_section_text:
        human_content = (
            f"Testo esatto da usare come base:\n{relevant_section_text}\n\n"
            "COPIA questo testo ESATTAMENTE come appare. Sostituisci SOLO i segnaposto "
            "(__________, [DA COMPILARE], spazi vuoti dopo etichette) con i dati forniti "
            "dall'utente. NON riscrivere, NON riformulare, NON aggiungere contenuto. "
            "Se un dato non è disponibile, lascia il segnaposto invariato.\n\n"
        ) + human_content

    raw_output = _call_chat(
        [SystemMessage(content=system_content), HumanMessage(content=human_content)],
        max_tokens=4000,
    )

    raw_output = re.sub(r'\[[^\]]*\]', '[DA COMPILARE]', raw_output)

    # Remove the "Campi strutturati" section and everything after it —
    # it's a skeleton artifact (field names list) not part of the legal document.
    raw_output = re.split(
        r'\n+(?:\*{0,2}(?:Campi strutturati|CAMPI STRUTTURATI)\*{0,2})\s*[\n:]',
        raw_output,
        maxsplit=1,
        flags=re.IGNORECASE,
    )[0].rstrip()

    if studio_name:
        header = f"{studio_name}\n{'─' * len(studio_name)}\n\n"
        raw_output = header + raw_output

    return {"draft": raw_output, "case_details": fields_dict, "doc_type": doc_type}


def generate_opposition_act(
    case_details: Dict[str, Any],
    retrieved_sections: List[Any],
    session_lang: str = "it",
    citations: list = None,
) -> str:
    """Thin wrapper around generate_document for backwards compatibility."""
    lang = session_lang or "it"
    lines = [f"{k}: {v}" for k, v in case_details.items() if v]
    if retrieved_sections:
        lines.append(
            f"\nRelevant knowledge base sections:\n{_format_retrieved_sections(retrieved_sections, lang)}"
        )
    return generate_document("\n".join(lines), "opposition_act", lang, citations)["draft"]
