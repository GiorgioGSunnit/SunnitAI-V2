"""FastAPI backend for the SunnitAI chatbot.

Endpoints:
    POST   /api/chat              — Send a message (creates session if needed)
    POST   /api/sessions          — Create a new session
    GET    /api/sessions          — List all sessions
    GET    /api/sessions/{id}     — Get session history
    DELETE /api/sessions/{id}     — Delete a session
    GET    /api/health            — Health check
    GET    /api/documents         — List all documents with section counts
    GET    /api/documents/{id}/sections/{name} — Get full section content
"""

import asyncio
import io
import logging
logging.getLogger("src.rag").setLevel(logging.INFO)
import os
import re
import time
import urllib.parse
import uuid
from contextlib import asynccontextmanager
from functools import partial
from typing import List, Optional

from fastapi import Depends, FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session

from .session import (
    ChatBot,
    ChatSession,
    _generate_session_title,
    last_pending_calculation,
)
from ..db.base import get_db
from ..db import crud
from ..db.models import Tenant
from ..db.crud import find_user_document_by_name, get_user_document, find_expired_user_document_by_name, create_user_document
from ..rag.main import run as rag_run, driver as neo4j_driver, NEO4J_DATABASE
from ..rag.verbose_logger import vlog
from ..rag.document_generation import (
    DOCUMENT_TYPE_REGISTRY,
    SYSTEM_TEMPLATES_BY_KEY,
    _placeholder,
    _extract_docx_elements,
    _extract_pdf_elements,
    _fill_template_gaps,
    _apply_fill_to_docx,
    _build_docx_from_pdf_elements,
    _summarise_da_compilare,
    classify_document_type,
    classify_system_template,
    generate_document,
    is_generation_request,
    _call_chat,
)
from ..rag.graph_nodes import _extract_citations
from ..rag.prompts import legal_consultant_system_prefix, _LENGTH
from langchain_core.messages import SystemMessage, HumanMessage
from .auth import get_current_user, require_user, create_access_token, verify_password, hash_password
from .billing import enforce_tenant_product_access, serialize_subscription
from .user_store import (get_user_by_email, get_user_by_id,
    get_tenant_by_id, get_tenant_profile_full, upsert_tenant_profile,
    get_user_document_for_generation, create_studio_and_admin,
    create_user_with_invite, get_tenant_invite_code, update_user_profile,
    update_user_credentials, get_user_settings)

logger = logging.getLogger(__name__)


async def _background_embedding_job():
    """Periodically generate embeddings for Section nodes missing them."""
    _running = False
    while True:
        await asyncio.sleep(120)  # every 2 minutes
        if _running:
            logger.debug("Background embedding job: previous run still in progress, skipping")
            continue
        _running = True
        try:
            from ..preprocessing.generate_embeddings import embed_missing
            count = await asyncio.get_event_loop().run_in_executor(
                None, embed_missing, NEO4J_DATABASE
            )
            if count > 0:
                logger.info(f"Background embedding job: generated {count} embeddings")
        except Exception as e:
            logger.warning(f"Background embedding job failed: {e}")
        finally:
            _running = False


async def _background_doc_expiry_job():
    """Delete expired user documents once per hour."""
    while True:
        await asyncio.sleep(3600)
        try:
            from ..db.base import SessionLocal
            from ..db.crud import cleanup_expired_documents
            db = SessionLocal()
            try:
                count = cleanup_expired_documents(db)
                if count:
                    logger.info("Doc expiry job: removed %d expired document(s)", count)
            finally:
                db.close()
        except Exception as e:
            logger.warning("Doc expiry job failed: %s", e)


@asynccontextmanager
async def _lifespan(app: FastAPI):
    from ..rag.cypher_logger import ensure_cypher_log_ready

    log_path = ensure_cypher_log_ready()
    logger.info("Cypher query log file: %s", log_path)
    task = asyncio.create_task(_background_embedding_job())
    expiry_task = asyncio.create_task(_background_doc_expiry_job())
    yield
    task.cancel()
    expiry_task.cancel()


# ---------------------------------------------------------------------------
# App & chatbot singleton
# ---------------------------------------------------------------------------

app = FastAPI(
    title="SunnitAI ChatBot API",
    description="RAG-powered chatbot over legal documents stored in a Neo4j knowledge graph.",
    version="1.0.0",
    lifespan=_lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

from .routes.auth import router as auth_router
from .routes.totp import router as totp_router
from .routes.users import router as users_router
from .routes.documents import router as documents_router
from .routes.billing import router as billing_router
from .routes.admin import router as admin_router
from .tracking_ws import router as tracking_router
app.include_router(auth_router, prefix="/api")
app.include_router(totp_router, prefix="/api")
app.include_router(users_router, prefix="/api")
app.include_router(documents_router, prefix="/api")
app.include_router(billing_router, prefix="/api")
app.include_router(admin_router, prefix="/api")
app.include_router(tracking_router, prefix="/api")

chatbot = ChatBot()


# ---------------------------------------------------------------------------
# Request / response models
# ---------------------------------------------------------------------------

class ChatRequest(BaseModel):
    message: str = Field(..., min_length=1, description="The user's message")
    session_id: Optional[str] = Field(
        None,
        description="Session ID for multi-turn conversation. Omit to auto-create a new session.",
    )
    document_id: Optional[str] = Field(
        None,
        description="UUID of a user-uploaded document to analyse in this message. "
                    "Used when the user refers to a document without naming it explicitly.",
    )
    document_ids: Optional[List[str]] = Field(
        None,
        description="UUIDs of multiple user-uploaded documents. "
                    "Takes priority over document_id when both are supplied.",
    )


DEFAULT_LOGO_PATH = "/opt/chatbot/assets/studio_logo.jpeg"
TENANT_LOGOS_BASE = "/opt/chatbot/data/tenant_logos"
ALLOWED_LOGO_EXTENSIONS = {".png", ".jpg", ".jpeg"}
MAX_LOGO_SIZE_BYTES = 5 * 1024 * 1024  # 5MB


class GenerateRequest(BaseModel):
    message: str = Field(..., min_length=1, description="Free-text request describing the opposition case.")
    session_id: Optional[str] = Field(None, description="Session ID for context. Omit to auto-create.")
    studio_name: str = Field("", description="Legal studio or organization name for document header.")
    studio_logo_path: str = Field("", description="Absolute server path to logo image file.")
    doc_type: str = Field("", description="Pre-classified document type; skips classify_document_type if provided.")
    draft: str = Field("", description="Pre-generated draft text; skips generation if provided.")
    section_hint: str = ""


class GenerateResponse(BaseModel):
    draft: str
    case_details: dict
    sources: list
    session_id: str
    doc_type: str = ""


class ChatResponse(BaseModel):
    session_id: str
    answer: str
    original_query: str
    resolved_query: str
    session_language: str = Field(
        default="it",
        description="Active session language code: it, en, or es.",
    )
    status_messages: list = Field(
        default_factory=list,
        description="Pipeline status lines (e.g. retrieval evaluation phase).",
    )
    citations: list = Field(
        default_factory=list,
        description="Structured citation list: [{document_name, document_id, sections}].",
    )
    awaiting_clarification: bool = Field(
        default=False,
        description="True when this turn's answer ends with a clarifying question and the "
        "next user message will be used to re-rank the previously retrieved sections "
        "instead of re-running retrieval.",
    )
    draft: str = Field(
        default="",
        description="Pre-generated draft — used by FE when user picks system template in picker.",
    )
    title: str = Field(
        default="Nuova conversazione",
        description="Auto-generated session title (set after first message).",
    )
    is_comparison: bool = Field(
        default=False,
        description="True when the answer was produced by the comparison retrieval path.",
    )
    generated_document_id: Optional[str] = Field(
        default=None,
        description="UUID of the saved generated document in user_documents. Present when generation_mode.",
    )
    generated_document_name: Optional[str] = Field(
        default=None,
        description="Original filename of the saved generated document.",
    )
    generation_candidates: list = Field(
        default_factory=list,
        description="Ranked candidate system templates ({key, label, codice, score}) when "
        "classify_system_template finds 2+ close matches; FE should prompt the user to pick one.",
    )


class SessionResponse(BaseModel):
    session_id: str
    created_at: str
    message_count: int
    title: str = "Nuova conversazione"


class LoginRequest(BaseModel):
    email: str
    password: str


class AuthResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    user_id: str
    email: str
    studio_name: str = ""
    subscription: Optional[dict] = None


class UserProfileResponse(BaseModel):
    user_id: str
    email: str
    role: str
    studio_name: str = ""
    first_name: str = ""
    last_name: str = ""
    tenant_id: str
    subscription: Optional[dict] = None


class UpdateProfileRequest(BaseModel):
    first_name: Optional[str] = None
    last_name: Optional[str] = None
    display_name: Optional[str] = None
    professional_title: Optional[str] = None
    phone: Optional[str] = None


class CredentialsUpdateRequest(BaseModel):
    current_password: str
    new_password: Optional[str] = None
    username: Optional[str] = None


class RegisterStudioRequest(BaseModel):
    email: str
    password: str
    first_name: str
    last_name: str
    studio_name: str


class RegisterUserRequest(BaseModel):
    email: str
    password: str
    first_name: str
    last_name: str
    invite_code: str


class RegisterStudioResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    user_id: str
    email: str
    studio_name: str
    invite_code: str


_PH_PATTERN = re.compile(r"\[[A-ZÀÁÂÄÉÈÊËÍÌÎÏÓÒÔÖÚÙÛÜ\s]+\](?:\s*\([^)]*\))?")

_DOC_FILENAMES = {
    "opposition_act": ("ATTO DI OPPOSIZIONE A DECRETO INGIUNTIVO", "ESCRITO DE OPOSICIÓN A DECRETO MONITORIO", "OPPOSITION TO PAYMENT ORDER", "atto_opposizione"),
    "rental_basic": ("CONTRATTO DI LOCAZIONE CON CEDOLARE SECCA", "CONTRATO DE ARRENDAMIENTO", "RENTAL AGREEMENT", "contratto_locazione_cedolare"),
    "rental_standard": ("CONTRATTO DI LOCAZIONE ABITATIVA", "CONTRATO DE ARRENDAMIENTO RESIDENCIAL", "RESIDENTIAL RENTAL AGREEMENT", "contratto_locazione_abitativa"),
    "rental_student": ("LOCAZIONE ABITATIVA PER STUDENTI UNIVERSITARI", "CONTRATO DE ARRENDAMIENTO PARA ESTUDIANTES", "STUDENT RENTAL AGREEMENT", "contratto_locazione_studenti"),
    "rental_transitional": ("LOCAZIONE ABITATIVA DI NATURA TRANSITORIA", "CONTRATO DE ARRENDAMIENTO TRANSITORIO", "TRANSITIONAL RENTAL AGREEMENT", "contratto_locazione_transitoria"),
    "rental_free_rent": ("CONTRATTO DI LOCAZIONE A CANONE LIBERO", "CONTRATO DE ARRENDAMIENTO A PRECIO LIBRE", "FREE RENT AGREEMENT", "contratto_locazione_canone_libero"),
    "rental_commercial": ("CONTRATTO DI LOCAZIONE AD USO COMMERCIALE", "CONTRATO DE ARRENDAMIENTO COMERCIAL", "COMMERCIAL LEASE AGREEMENT", "contratto_locazione_commerciale"),
    "rental_cancellation": ("DISDETTA CONTRATTO DI LOCAZIONE", "RESCISIÓN CONTRATO DE ARRENDAMIENTO", "RENTAL CANCELLATION NOTICE", "disdetta_locazione"),
    "insurance_cancellation": ("DISDETTA POLIZZA ASSICURATIVA", "RESCISIÓN PÓLIZA DE SEGUROS", "INSURANCE CANCELLATION NOTICE", "disdetta_polizza"),
    "insurance_declaration": ("DICHIARAZIONE SOSTITUTIVA DI POLIZZA ASSICURATIVA", "DECLARACIÓN SUSTITUTIVA DE PÓLIZA", "INSURANCE SUBSTITUTIVE DECLARATION", "dichiarazione_polizza"),
    "employment_dismissal_appeal": ("IMPUGNATIVA DI LICENZIAMENTO", "IMPUGNACIÓN DE DESPIDO", "DISMISSAL APPEAL", "impugnativa_licenziamento"),
    "employment_termination": ("LETTERA DI LICENZIAMENTO PER GIUSTA CAUSA", "CARTA DE DESPIDO POR CAUSA JUSTIFICADA", "TERMINATION LETTER FOR CAUSE", "lettera_licenziamento"),
    "franchising_contract": ("CONTRATTO DI FRANCHISING", "CONTRATO DE FRANQUICIA", "FRANCHISING AGREEMENT", "contratto_franchising"),
    "demand_letter": ("LETTERA DI DIFFIDA", "CARTA DE REQUERIMIENTO", "DEMAND LETTER", "lettera_diffida"),
    "appeal": ("RICORSO", "RECURSO", "APPEAL", "ricorso"),
    "power_of_attorney": ("PROCURA", "PODER NOTARIAL", "POWER OF ATTORNEY", "procura"),
    "sale_agreement": ("CONTRATTO DI COMPRAVENDITA", "CONTRATO DE COMPRAVENTA", "SALE AGREEMENT", "contratto_compravendita"),
    "verbale_assemblea": ("VERBALE DI ASSEMBLEA CONDOMINIALE", "ACTA DE JUNTA DE PROPIETARIOS", "CONDOMINIUM ASSEMBLY MINUTES", "verbale_assemblea"),
    "nota_contestazione": ("NOTA ALLA CONTESTAZIONE", "NOTA A LA CONTESTACIÓN", "NOTICE CONTESTING TRAFFIC VIOLATION", "nota_contestazione"),
    "comparison": ("CONFRONTO TRA DOCUMENTI", "COMPARACIÓN DE DOCUMENTOS", "DOCUMENT COMPARISON", "confronto_documenti"),
}




def _strip_markdown(text: str) -> str:
    text = re.sub(r'\*\*|__', '', text)
    text = re.sub(r'[*_]', '', text)
    return text


# ---------------------------------------------------------------------------
# Generation helpers
# ---------------------------------------------------------------------------

def _raw_result_to_sections(raw_result: list) -> list:
    """Convert raw Neo4j result records to flat dicts for _format_retrieved_sections."""
    sections = []
    seen: set = set()
    for record in raw_result:
        for value in record.values():
            if not isinstance(value, dict) or "properties" not in value:
                continue
            props = value["properties"]
            labels = value.get("labels", [])
            title = props.get("heading") or props.get("title") or ""
            text = props.get("text_en") or props.get("text_it") or props.get("text") or ""
            source = props.get("document_title") or props.get("document_id") or ""
            key = (title, source)
            if key in seen or not (title or text):
                continue
            seen.add(key)
            sections.append({"title": title, "text": text, "document_title": source, "labels": labels})
    return sections


def _get_cached_sections(session) -> Optional[list]:
    """Return converted sections from the most recent RAG-backed assistant message (last 2 turns).

    The normal chat flow stores raw_result in metadata["references"]. A record is
    RAG-backed when its values are dicts containing a "properties" key.
    """
    if not session:
        return None
    assistant_msgs = [m for m in reversed(session.messages) if m.role == "assistant"]
    for msg in assistant_msgs[:2]:
        refs = (msg.metadata or {}).get("references") or []
        if refs and isinstance(refs[0], dict) and "properties" in refs[0]:
            sections = _raw_result_to_sections(refs)
            if sections:
                return sections
    return None


def _build_clarification_message() -> str:
    types_list = "\n".join(
        f"- **{entry['label']}**"
        for entry in DOCUMENT_TYPE_REGISTRY.values()
    )
    return (
        "Non ho capito che tipo di documento vuoi generare. "
        "Puoi specificare meglio la tua richiesta? I tipi di documento disponibili sono:\n\n"
        f"{types_list}\n\n"
        "Indica quale documento desideri e fornisci i dettagli necessari."
    )


def _run_generation_sync(message: str, session_lang: str, doc_type: str, cached_sections: Optional[list] = None, studio_name: str = "", section_hint: str = "") -> dict:
    citations = None
    if cached_sections is not None:
        sources = sorted({s["document_title"] for s in cached_sections if s.get("document_title")})
    else:
        try:
            # Retrieval only: this call exists to collect the sources the draft
            # will cite, and the intent was already settled upstream. Left
            # unguarded, the calculation gate matches a drafting request on its
            # domain vocabulary alone ("contratto di locazione", "fattura ...
            # IVA") and answers with a tax figure instead of retrieving, so the
            # document would be generated with no sources and no citations.
            rag_state = rag_run(
                message, session_language=session_lang, skip_calculation=True
            )
            raw_result = rag_state.get("raw_result") or []
            retrieved_sections = _raw_result_to_sections(raw_result)
            sources = sorted({s["document_title"] for s in retrieved_sections if s.get("document_title")})
            # _extract_citations takes the Neo4j result ROWS, not the graph
            # state. Handing it the state made it iterate the state's keys and
            # raise on the first one, so this whole block threw on every
            # generation that actually retrieved something: `sources` was
            # computed and then discarded by the handler below, and the draft
            # was written with citations=None. Sources were never lost to the
            # calculation gate alone — they were lost every time.
            citations = _extract_citations(raw_result)
        except Exception as exc:
            logger.warning("RAG retrieval for generation failed: %s", exc)
            sources = []
    gen = generate_document(message, doc_type, session_lang, citations, studio_name, section_hint)
    return {"draft": gen["draft"], "case_details": gen["case_details"], "sources": sources, "doc_type": gen["doc_type"]}


def _run_comparison_sync(message: str, session_lang: str, cached_sections=None) -> dict:
    """Run a document comparison through the RAG graph, returning answer + citations.

    Retrieval-only, like the generation lookup above, and for a sharper reason:
    a calculator is itself a comparator ("confronto gas e luce") and shares
    this flow's opening verb, so the gate can auto-route a request to compare
    two uploaded files into it. Because the
    comparison's answer is returned as the draft, the caller would hand the
    user that calculator's candidate-collection prompt instead of a comparison,
    with no citations. Comparing two offers numerically is a calculation;
    comparing two documents' text is not, and the verb cannot tell them apart.
    """
    rag_state = rag_run(
        message, session_language=session_lang, skip_calculation=True
    )
    return {
        "answer": rag_state.get("answer", ""),
        "citations": rag_state.get("citations", []),
        "is_comparison": bool(rag_state.get("is_comparison")),
    }


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------

@app.post("/api/auth/login", response_model=AuthResponse)
async def login(request: LoginRequest, db: Session = Depends(get_db)):
    user = get_user_by_email(request.email)
    if not user or not user.get("is_active"):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    if not verify_password(request.password, user["hashed_password"]):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    token = create_access_token({
        "sub": str(user["id"]),
        "email": user["email"],
        "role": user["role"],
        "tenant_id": str(user["tenant_id"]),
    })
    tenant_uuid = uuid.UUID(str(user["tenant_id"]))
    tenant = db.query(Tenant).filter(Tenant.id == tenant_uuid).first()
    seats_used = crud.count_active_users_for_tenant(db, tenant_uuid)
    return AuthResponse(
        access_token=token,
        user_id=str(user["id"]),
        email=user["email"],
        studio_name=user.get("studio_name") or "",
        subscription=serialize_subscription(
            crud.get_tenant_subscription(db, tenant_uuid),
            tenant,
            seats_used=seats_used,
        ),
    )


@app.post("/api/auth/register/studio", response_model=RegisterStudioResponse)
async def register_studio(request: RegisterStudioRequest):
    existing = get_user_by_email(request.email)
    if existing:
        raise HTTPException(status_code=409, detail="Email already registered")
    hashed = hash_password(request.password)
    user = create_studio_and_admin(
        email=request.email,
        hashed_password=hashed,
        first_name=request.first_name,
        last_name=request.last_name,
        studio_name=request.studio_name,
    )
    token = create_access_token({
        "sub": user["id"],
        "email": user["email"],
        "role": user["role"],
        "tenant_id": user["tenant_id"],
    })
    return RegisterStudioResponse(
        access_token=token,
        user_id=user["id"],
        email=user["email"],
        studio_name=user["studio_name"],
        invite_code=user["invite_code"],
    )


@app.post("/api/auth/register", response_model=AuthResponse)
async def register_user(request: RegisterUserRequest):
    existing = get_user_by_email(request.email)
    if existing:
        raise HTTPException(status_code=409, detail="Email already registered")
    hashed = hash_password(request.password)
    try:
        user = create_user_with_invite(
            email=request.email,
            hashed_password=hashed,
            first_name=request.first_name,
            last_name=request.last_name,
            invite_code=request.invite_code,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    token = create_access_token({
        "sub": user["id"],
        "email": user["email"],
        "role": user["role"],
        "tenant_id": user["tenant_id"],
    })
    return AuthResponse(
        access_token=token,
        user_id=user["id"],
        email=user["email"],
        studio_name=user.get("studio_name") or "",
    )


@app.get("/api/auth/me", response_model=UserProfileResponse)
async def get_me(current_user: dict = Depends(require_user), db: Session = Depends(get_db)):
    user = get_user_by_id(current_user["sub"])
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    tenant_uuid = uuid.UUID(str(user["tenant_id"]))
    tenant = db.query(Tenant).filter(Tenant.id == tenant_uuid).first()
    seats_used = crud.count_active_users_for_tenant(db, tenant_uuid)
    return UserProfileResponse(
        user_id=str(user["id"]),
        email=user["email"],
        role=user["role"],
        studio_name=user.get("studio_name") or "",
        first_name=user.get("first_name") or "",
        last_name=user.get("last_name") or "",
        tenant_id=str(user["tenant_id"]),
        subscription=serialize_subscription(
            crud.get_tenant_subscription(db, tenant_uuid),
            tenant,
            seats_used=seats_used,
        ),
    )


@app.put("/api/auth/me", response_model=UserProfileResponse)
async def update_me(
    request: UpdateProfileRequest,
    current_user: dict = Depends(require_user),
    db: Session = Depends(get_db),
):
    try:
        user = update_user_profile(
            user_id=current_user["sub"],
            first_name=request.first_name,
            last_name=request.last_name,
            display_name=request.display_name,
            professional_title=request.professional_title,
            phone=request.phone,
        )
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        tenant_uuid = uuid.UUID(str(user["tenant_id"]))
        tenant = db.query(Tenant).filter(Tenant.id == tenant_uuid).first()
        seats_used = crud.count_active_users_for_tenant(db, tenant_uuid)
        return UserProfileResponse(
            user_id=str(user["id"]),
            email=user["email"],
            role=user["role"],
            studio_name=user.get("studio_name") or "",
            first_name=user.get("first_name") or "",
            last_name=user.get("last_name") or "",
            tenant_id=str(user["tenant_id"]),
            subscription=serialize_subscription(
                crud.get_tenant_subscription(db, tenant_uuid),
                tenant,
                seats_used=seats_used,
            ),
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.put("/api/auth/me/credentials")
async def update_credentials(
    request: CredentialsUpdateRequest,
    current_user: dict = Depends(require_user),
):
    if not request.new_password and request.username is None:
        raise HTTPException(status_code=400, detail="Provide new_password and/or username")
    try:
        user = update_user_credentials(
            user_id=current_user["sub"],
            current_password=request.current_password,
            new_password=request.new_password,
            username=request.username,
        )
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        return {"ok": True}
    except ValueError as e:
        if "current_password_invalid" in str(e):
            raise HTTPException(status_code=401, detail="Current password is incorrect")
        if "new_password_too_short" in str(e):
            raise HTTPException(status_code=400, detail="New password must be at least 8 characters")
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


class TenantProfileUpdateRequest(BaseModel):
    legal_name: Optional[str] = None
    display_name: Optional[str] = None
    vat_number: Optional[str] = None
    phone: Optional[str] = None
    website: Optional[str] = None
    address_street: Optional[str] = None
    address_city: Optional[str] = None
    address_postal_code: Optional[str] = None
    address_country: Optional[str] = None
    document_expiry_hours: Optional[int] = None


@app.get("/api/tenant/profile")
async def get_tenant_profile(
    current_user: dict = Depends(require_user),
):
    tenant_id = current_user.get("tenant_id")
    if not tenant_id:
        raise HTTPException(status_code=400, detail="No tenant associated with this account")
    profile = get_tenant_profile_full(tenant_id) or {}
    if "document_expiry_hours" not in profile or profile["document_expiry_hours"] is None:
        profile["document_expiry_hours"] = 24
    return profile


class UserTemplateRequest(BaseModel):
    doc_id: str = Field(..., description="UUID of the user's uploaded template")
    message: str = Field(..., min_length=1, description="Free-text describing what to fill in")
    session_id: Optional[str] = None


@app.post("/api/generate/from-user-template")
async def generate_from_user_template(
    request: UserTemplateRequest,
    current_user: dict = Depends(require_user),
    db: Session = Depends(get_db),
):
    """Fill a user's own uploaded template with data from their message.
    Returns a filled .docx file. PDFs are extracted and rewritten as .docx."""
    user_id = current_user["sub"]
    tenant_id = current_user.get("tenant_id")
    enforce_tenant_product_access(db, uuid.UUID(str(tenant_id)))

    doc_meta = get_user_document_for_generation(user_id, tenant_id, request.doc_id)
    if not doc_meta:
        raise HTTPException(status_code=404, detail="Template not found or access denied")

    storage_path = doc_meta["storage_path"]
    original_filename = doc_meta["original_filename"]

    if not os.path.exists(storage_path):
        raise HTTPException(status_code=500, detail="Template file not found on disk")

    session_id = request.session_id
    _uid = current_user["sub"]
    _tid = current_user.get("tenant_id")
    if not session_id:
        session = chatbot.create_session(user_id=_uid, tenant_id=_tid)
        session_id = session.session_id
    session = chatbot.get_session(session_id, user_id=_uid)
    if not session:
        session = ChatSession(session_id=session_id, user_id=_uid, tenant_id=_tid)
        chatbot._sessions[session_id] = session
    session_lang = session.session_language
    session.add_message("user", request.message)

    carta_intestata = get_tenant_profile_full(tenant_id) if tenant_id else None

    ext = os.path.splitext(storage_path)[1].lower()
    is_pdf = ext == ".pdf"

    try:
        elements = _extract_pdf_elements(storage_path) if is_pdf else _extract_docx_elements(storage_path)
        if not elements:
            raise HTTPException(status_code=422, detail="Could not extract any text from the template")
        session_messages = [{"role": m.role, "content": m.content} for m in session.messages]
        fill_map = _fill_template_gaps(
            elements, request.message, carta_intestata, session_lang, session_messages,
            docx_path=None if is_pdf else storage_path,
        )

        if is_pdf:
            docx_bytes = _build_docx_from_pdf_elements(elements, fill_map)
        else:
            docx_bytes = _apply_fill_to_docx(storage_path, fill_map)

    except HTTPException:
        raise
    except Exception as exc:
        logger.error("User template generation error: %s", exc, exc_info=True)
        raise HTTPException(status_code=500, detail=str(exc))

    base_name = os.path.splitext(original_filename)[0]
    output_filename = f"{base_name}_compilato.docx"

    _missing = _summarise_da_compilare(fill_map, elements, session_lang)
    confirmation = _build_generation_confirmation(session_lang, _missing or None)
    if is_pdf:
        conversion_note = {
            "it": " (il tuo PDF è stato convertito in DOCX)",
            "en": " (your PDF was converted to DOCX)",
            "es": " (tu PDF fue convertido a DOCX)",
        }.get(session_lang, " (il tuo PDF è stato convertito in DOCX)")
        confirmation = confirmation.rstrip(".") + conversion_note + "."
    session.add_message("assistant", confirmation, metadata={})

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{output_filename}"'},
    )


@app.patch("/api/tenant/profile")
async def update_tenant_profile(
    request: TenantProfileUpdateRequest,
    current_user: dict = Depends(require_user),
):
    tenant_id = current_user.get("tenant_id")
    if not tenant_id:
        raise HTTPException(status_code=400, detail="No tenant associated with this account")
    if current_user.get("role") not in ("admin", "superadmin"):
        raise HTTPException(status_code=403, detail="Only admins can update company settings")
    if request.document_expiry_hours is not None and request.document_expiry_hours not in (24, 48, 72):
        raise HTTPException(status_code=400, detail="document_expiry_hours must be 24, 48, or 72")
    fields = {k: v for k, v in request.model_dump().items() if v is not None}
    if not fields:
        raise HTTPException(status_code=400, detail="No fields provided")
    updated = upsert_tenant_profile(tenant_id, fields)
    if "document_expiry_hours" not in updated or updated["document_expiry_hours"] is None:
        updated["document_expiry_hours"] = 24
    return updated


@app.post("/api/tenant/logo")
async def upload_tenant_logo(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user),
):
    """Upload (or replace) the studio's logo, used in generated document
    letterheads. One logo per tenant — re-uploading replaces the existing
    file. No FE UI yet."""
    tenant_id = current_user.get("tenant_id")
    if not tenant_id:
        raise HTTPException(status_code=400, detail="No tenant associated with this account")

    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in ALLOWED_LOGO_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Allowed: {', '.join(sorted(ALLOWED_LOGO_EXTENSIONS))}",
        )

    contents = await file.read()
    if len(contents) > MAX_LOGO_SIZE_BYTES:
        raise HTTPException(status_code=400, detail="Logo file too large (max 5MB)")

    folder = os.path.join(TENANT_LOGOS_BASE, str(tenant_id))
    os.makedirs(folder, exist_ok=True)

    for existing_file in os.listdir(folder):
        try:
            os.remove(os.path.join(folder, existing_file))
        except OSError:
            pass

    logo_path = os.path.join(folder, f"logo{ext}")
    with open(logo_path, "wb") as f:
        f.write(contents)

    updated = upsert_tenant_profile(tenant_id, {"logo_path": logo_path})
    return {"logo_path": updated.get("logo_path")}


@app.get("/api/tenant/logo")
async def get_tenant_logo(
    current_user: dict = Depends(require_user),
):
    tenant_id = current_user.get("tenant_id")
    if not tenant_id:
        raise HTTPException(status_code=400, detail="No tenant associated with this account")
    profile = get_tenant_profile_full(tenant_id)
    logo_path = profile.get("logo_path") if profile else None
    if not logo_path or not os.path.exists(logo_path):
        raise HTTPException(status_code=404, detail="No logo uploaded yet")
    return FileResponse(logo_path)


@app.get("/api/health")
def health_check():
    return {"status": "ok"}


@app.get("/api/debug")
async def debug_check():
    """Check connectivity to LLM and Neo4j."""
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _debug_check_sync)


def _debug_check_sync() -> dict:
    """Synchronous debug checks (runs in thread pool to avoid blocking event loop)."""
    import os
    from dotenv import load_dotenv

    from ..rag.cypher_logger import get_cypher_log_path

    load_dotenv()

    checks = {}

    cypher_log = get_cypher_log_path()
    checks["cypher_log_path"] = cypher_log
    checks["cypher_log_exists"] = os.path.isfile(cypher_log)
    try:
        with open(cypher_log, "a", encoding="utf-8"):
            pass
        checks["cypher_log_writable"] = True
    except OSError as e:
        checks["cypher_log_writable"] = f"error: {e}"

    # Check Neo4j
    try:
        from neo4j import GraphDatabase
        uri = os.getenv("NEO4J_URI")
        user = os.getenv("NEO4J_USER")
        pwd = os.getenv("NEO4J_PASSWORD")
        checks["neo4j_config"] = {"uri": uri, "user": user, "password_set": bool(pwd)}
        driver = GraphDatabase.driver(uri, auth=(user, pwd))
        with driver.session(database=os.getenv("NEO4J_DATABASE", "neo4j")) as session:
            result = session.run("RETURN 1 AS n").single()
            checks["neo4j"] = "ok" if result else "query returned nothing"
        driver.close()
    except Exception as e:
        checks["neo4j"] = f"error: {e}"

    # Check LLM
    try:
        base_url = os.getenv("LLM_BASE_URL")
        model = os.getenv("LLM_MODEL")
        api_key = os.getenv("LLM_API_KEY")
        checks["llm_config"] = {
            "base_url": base_url,
            "model": model,
            "api_key_set": bool(api_key),
        }
        from langchain_openai import ChatOpenAI
        llm = ChatOpenAI(
            model=model, api_key=api_key, base_url=base_url, temperature=0
        )
        resp = llm.invoke("Say hello in one word.")
        checks["llm"] = f"ok: {resp.content[:100]}"
    except Exception as e:
        checks["llm"] = f"error: {e}"

    # Check embeddings
    try:
        from ..rag.ai_chat import embedding_model
        test = embedding_model.embed_query("test")
        checks["embeddings"] = f"ok: dim={len(test)}"
    except Exception as e:
        checks["embeddings"] = f"error: {e}"

    return checks


@app.get("/api/documents")
def list_documents():
    """List all Document nodes with their section counts."""
    try:
        with neo4j_driver.session(database=NEO4J_DATABASE) as session:
            result = session.run(
                "MATCH (d:Document)-[:CONTAINS]->(s:Section) "
                "RETURN d.name AS document_name, d.id AS document_id, count(s) AS section_count "
                "ORDER BY d.name"
            )
            documents = [
                {
                    "document_name": r["document_name"],
                    "document_id": r["document_id"],
                    "section_count": r["section_count"],
                }
                for r in result
            ]
    except Exception as e:
        logger.error("list_documents error: %s", e)
        raise HTTPException(status_code=500, detail=str(e))
    return {"documents": documents}


@app.get("/api/documents/{document_id}/sections/{section_name}")
def get_section(document_id: str, section_name: str):
    """Return full content of a specific section."""
    decoded_doc_id = urllib.parse.unquote(document_id)
    decoded_section = urllib.parse.unquote(section_name)
    parts = decoded_doc_id.split("::")
    doc_hash = parts[1] if len(parts) >= 2 else decoded_doc_id
    try:
        with neo4j_driver.session(database=NEO4J_DATABASE) as session:
            result = session.run(
                "MATCH (d:Document {id: $doc_id})-[:CONTAINS]->(s:Section) "
                "WHERE s.name = $section_name "
                "RETURN d.name AS document_name, s.name AS section_name, "
                "s.abstract AS abstract, s.plain_text AS plain_text "
                "LIMIT 1",
                doc_id=decoded_doc_id,
                section_name=decoded_section,
            )
            row = result.single()
    except Exception as e:
        logger.error("get_section error: %s", e)
        raise HTTPException(status_code=500, detail=str(e))
    if not row:
        raise HTTPException(status_code=404, detail="Section not found")
    return {
        "document_name": row["document_name"],
        "section_name": row["section_name"],
        "abstract": row["abstract"] or "",
        "plain_text": row["plain_text"] or "",
    }


@app.post("/api/sessions", response_model=SessionResponse)
def create_session(current_user: dict = Depends(require_user), db: Session = Depends(get_db)):
    enforce_tenant_product_access(db, uuid.UUID(str(current_user["tenant_id"])))
    session = chatbot.create_session(
        user_id=current_user["sub"],
        tenant_id=current_user.get("tenant_id"),
    )
    return SessionResponse(
        session_id=session.session_id,
        created_at=session.created_at,
        message_count=0,
        title=session.title,
    )


@app.get("/api/sessions")
def list_sessions(current_user: dict = Depends(require_user), db: Session = Depends(get_db)):
    enforce_tenant_product_access(db, uuid.UUID(str(current_user["tenant_id"])))
    return chatbot.list_sessions(user_id=current_user["sub"])


@app.get("/api/sessions/{session_id}")
def get_session(session_id: str, current_user: dict = Depends(require_user), db: Session = Depends(get_db)):
    enforce_tenant_product_access(db, uuid.UUID(str(current_user["tenant_id"])))
    session = chatbot.get_session(session_id, user_id=current_user["sub"])
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    return session.to_dict()


@app.delete("/api/sessions/{session_id}")
def delete_session(session_id: str, current_user: dict = Depends(require_user), db: Session = Depends(get_db)):
    enforce_tenant_product_access(db, uuid.UUID(str(current_user["tenant_id"])))
    if not chatbot.delete_session(session_id, user_id=current_user["sub"]):
        raise HTTPException(status_code=404, detail="Session not found")
    return {"status": "deleted", "session_id": session_id}


def _persist_generated_docx(
    draft: str,
    doc_type: str,
    user_id: str,
    tenant_id: str,
    case_details: dict | None = None,
) -> tuple:
    """Save draft as a DOCX file on disk and create a user_documents record.
    Applies full letterhead (logo, studio name, address) if tenant profile exists.
    Uses case_details to build a meaningful filename (e.g. includes client name).
    Returns (document_id_str, original_filename) or (None, None) on failure."""
    try:
        from docx import Document as _DocxDoc
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io as _io
        from ..db.base import SessionLocal as _SL
        from ..constants import PRIVATE_DOCS_BASE as _BASE

        _catalog_entry = SYSTEM_TEMPLATES_BY_KEY.get(doc_type)
        if _catalog_entry:
            _label = _catalog_entry.get("label") or _catalog_entry.get("tipo_atto", "Documento")
        else:
            _label = doc_type.replace("_", " ").replace("-", " ").title() if doc_type else "Documento"
        _slug = re.sub(r"[^a-z0-9]+", "_", _label.lower()).strip("_")

        # Build meaningful filename from case_details — include client/subject name
        _name_suffix = ""
        if case_details:
            # Try common name fields in priority order
            _name_candidate = (
                case_details.get("conduttore") or
                case_details.get("nome_assistito") or
                case_details.get("nome_difensore") or
                case_details.get("debitore") or
                case_details.get("cliente") or
                case_details.get("nome") or
                case_details.get("richiedente") or
                case_details.get("locatore") or
                ""
            )
            # Only use if it's a real value, not a placeholder
            if _name_candidate and "[DA COMPILARE]" not in _name_candidate:
                # Sanitize: keep only letters, numbers, spaces; truncate to 30 chars
                _clean = re.sub(r"[^a-zA-Z0-9\s]", "", _name_candidate).strip()[:30]
                _clean_slug = re.sub(r"\s+", "_", _clean).lower()
                if _clean_slug:
                    _name_suffix = f"_{_clean_slug}"

        _filename = f"{_slug}{_name_suffix}.docx"

        _doc = _DocxDoc()
        _ph = "[DA COMPILARE]"

        # ── Letterhead ────────────────────────────────────────────────────
        _tenant_profile = get_tenant_profile_full(tenant_id) if tenant_id else None
        if _tenant_profile:
            _logo_path = _tenant_profile.get("logo_path") or ""
            if _logo_path and os.path.exists(_logo_path):
                _logo_para = _doc.add_paragraph()
                _logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _logo_run = _logo_para.add_run()
                _logo_run.add_picture(_logo_path, width=Inches(2.5))
                _logo_para.paragraph_format.space_after = Pt(12)

            _name_val = _tenant_profile.get("legal_name") or _tenant_profile.get("display_name") or _ph
            _name_para = _doc.add_paragraph()
            _name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _name_run = _name_para.add_run(_name_val)
            _name_run.bold = True
            _name_run.font.size = Pt(11)

            _addr_parts = [
                _tenant_profile.get("address_street") or _ph,
                f"{_tenant_profile.get('address_postal_code') or _ph} {_tenant_profile.get('address_city') or _ph}",
                _tenant_profile.get("address_country") or _ph,
            ]
            _addr_para = _doc.add_paragraph(", ".join(_addr_parts))
            _addr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            _contact_para = _doc.add_paragraph(
                f"Tel: {_tenant_profile.get('phone') or _ph}  |  Web: {_tenant_profile.get('website') or _ph}"
            )
            _contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            _vat_para = _doc.add_paragraph(f"P.IVA: {_tenant_profile.get('vat_number') or _ph}")
            _vat_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            _sep_para = _doc.add_paragraph()
            _sep_para.paragraph_format.space_before = Pt(6)
            _sep_run = _sep_para.add_run("─" * 60)
            _sep_run.font.size = Pt(8)
            _sep_run.font.color.rgb = RGBColor(180, 180, 180)
            _sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _sep_para.paragraph_format.space_after = Pt(12)

        # ── Document title ────────────────────────────────────────────────
        _title_para = _doc.add_paragraph()
        _title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _title_run = _title_para.add_run(_label)
        _title_run.bold = True
        _title_run.font.size = Pt(14)
        _doc.add_paragraph()

        # ── Draft content ─────────────────────────────────────────────────
        for _line in _strip_markdown(draft).split("\n"):
            _doc.add_paragraph(_line)

        _folder = os.path.join(_BASE, tenant_id, user_id)
        os.makedirs(_folder, exist_ok=True)
        _storage_path = os.path.join(_folder, f"{uuid.uuid4()}.docx")
        _buf = _io.BytesIO()
        _doc.save(_buf)
        _raw = _buf.getvalue()
        with open(_storage_path, "wb") as _f:
            _f.write(_raw)

        _db = _SL()
        try:
            _rec = create_user_document(
                _db,
                user_id=uuid.UUID(user_id),
                tenant_id=uuid.UUID(tenant_id),
                original_filename=_filename,
                storage_path=_storage_path,
                file_size_bytes=len(_raw),
                scope="personal",
                document_role="generated",
                expires_at=None,
            )
            return str(_rec.id), _filename
        finally:
            _db.close()
    except Exception as _exc:
        logger.warning("_persist_generated_docx failed: %s", _exc)
        return None, None


def _build_generation_confirmation(lang: str, missing_fields: Optional[List[str]] = None) -> str:
    if lang == "es":
        base = "He generado el documento solicitado."
        suffix_tpl = " Los siguientes campos deben completarse manualmente: {fields}."
    elif lang == "en":
        base = "I have generated the requested document."
        suffix_tpl = " The following fields need to be filled in manually: {fields}."
    else:
        base = "Ho generato il documento richiesto."
        suffix_tpl = " I seguenti campi devono essere compilati manualmente: {fields}."
    if not missing_fields:
        return base
    return base + suffix_tpl.format(fields=", ".join(missing_fields))


@app.post("/api/generate", response_model=GenerateResponse)
async def generate(request: GenerateRequest, current_user: Optional[dict] = Depends(get_current_user), db: Session = Depends(get_db)):
    """Generate a legal document draft from a free-text request.

    Classifies the document type, extracts case details, retrieves relevant sections
    from the knowledge base, and returns a structured draft.
    """
    is_comparison_request = bool(re.search(
        r'\b(confronta\b|confronto\s+(tra|fra|dei|di|delle|degli)\b|differenze?\s+tra\b|compara\b|paragona\b|versus\b|vs\.?)\b',
        request.message, re.IGNORECASE
    )) or request.doc_type == "comparison"

    if not is_comparison_request and not is_generation_request(request.message):
        raise HTTPException(status_code=400, detail="Not a generation request")

    if current_user:
        enforce_tenant_product_access(db, uuid.UUID(str(current_user["tenant_id"])))

    session_id = request.session_id
    _uid = current_user["sub"] if current_user else None
    _tid = current_user.get("tenant_id") if current_user else None
    if not session_id:
        session = chatbot.create_session(user_id=_uid, tenant_id=_tid)
        session_id = session.session_id

    session = chatbot.get_session(session_id, user_id=_uid)
    if not session:
        session = ChatSession(session_id=session_id, user_id=_uid, tenant_id=_tid)
        chatbot._sessions[session_id] = session

    session_lang = session.session_language

    if current_user and not request.studio_name:
        _profile = get_user_by_id(current_user["sub"])
        if _profile:
            request = request.model_copy(update={"studio_name": _profile.get("studio_name") or ""})

    if is_comparison_request:
        doc_type = "comparison"
    else:
        doc_type = classify_document_type(request.message, session_lang)
        if doc_type == "unknown":
            doc_type = classify_system_template(request.message, session_lang)
        if doc_type == "unknown":
            clarification = _build_clarification_message()
            session.add_message("user", request.message)
            session.add_message("assistant", clarification)
            return GenerateResponse(
                session_id=session_id, draft=clarification, case_details={}, sources=[], doc_type="unknown"
            )

    cached = _get_cached_sections(session)
    session.add_message("user", request.message)

    try:
        loop = asyncio.get_event_loop()
        if is_comparison_request:
            comparison_result = await loop.run_in_executor(
                None, partial(_run_comparison_sync, request.message, session_lang, cached)
            )
            result = {
                "draft": comparison_result.get("answer", ""),
                "case_details": {},
                "sources": comparison_result.get("citations", []),
                "doc_type": "comparison",
                "studio_name": request.studio_name,
            }
        else:
            result = await loop.run_in_executor(
                None, partial(_run_generation_sync, request.message, session_lang, doc_type, cached, request.studio_name, request.section_hint)
            )
    except Exception as exc:
        logger.error("Generation error: %s", exc, exc_info=True)
        raise HTTPException(status_code=500, detail=str(exc))

    session.add_message(
        "assistant",
        _build_generation_confirmation(session_lang),
        metadata={"sources": result.get("sources", [])},
    )
    return GenerateResponse(session_id=session_id, **result)


@app.post("/api/generate/download")
async def generate_download(request: GenerateRequest, current_user: Optional[dict] = Depends(get_current_user), db: Session = Depends(get_db)):
    """Generate opposition act and return as a downloadable .docx file."""
    try:
        from docx import Document
        from docx.shared import Cm, Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail="python-docx is not installed. Add 'python-docx>=1.1.0' to pyproject.toml and reinstall.",
        )

    is_comparison_request = bool(re.search(
        r'\b(confronta\b|confronto\s+(tra|fra|dei|di|delle|degli)\b|differenze?\s+tra\b|compara\b|paragona\b|versus\b|vs\.?)\b',
        request.message, re.IGNORECASE
    )) or request.doc_type == "comparison"

    if not is_comparison_request and not is_generation_request(request.message):
        raise HTTPException(status_code=400, detail="Not a generation request")

    if current_user:
        enforce_tenant_product_access(db, uuid.UUID(str(current_user["tenant_id"])))

    session_id = request.session_id
    _uid = current_user["sub"] if current_user else None
    _tid = current_user.get("tenant_id") if current_user else None
    if not session_id:
        session = chatbot.create_session(user_id=_uid, tenant_id=_tid)
        session_id = session.session_id
    session = chatbot.get_session(session_id, user_id=_uid)
    if not session:
        session = ChatSession(session_id=session_id, user_id=_uid, tenant_id=_tid)
        chatbot._sessions[session_id] = session
    session_lang = session.session_language

    if current_user and not request.studio_name:
        _profile = get_user_by_id(current_user["sub"])
        if _profile:
            request = request.model_copy(update={"studio_name": _profile.get("studio_name") or ""})

    if is_comparison_request:
        doc_type = "comparison"
    elif request.doc_type and (
        request.doc_type in SYSTEM_TEMPLATES_BY_KEY
        or request.doc_type in {"comparison"}
    ):
        doc_type = request.doc_type
    else:
        doc_type = classify_document_type(request.message, session_lang)
        if doc_type == "unknown":
            doc_type = classify_system_template(request.message, session_lang)
        if doc_type == "unknown":
            raise HTTPException(status_code=400, detail=_build_clarification_message())
    cached = _get_cached_sections(session)

    loop = asyncio.get_event_loop()
    if is_comparison_request and not request.draft:
        comparison_result = await loop.run_in_executor(
            None, partial(_run_comparison_sync, request.message, session_lang)
        )
        request = request.model_copy(update={"draft": comparison_result.get("answer", "")})

    try:
        if request.draft:
            result = {"draft": request.draft, "doc_type": doc_type, "sources": []}
        else:
            result = await loop.run_in_executor(
                None, partial(_run_generation_sync, request.message, session_lang, doc_type, cached, "", request.section_hint)
            )
    except ValueError as exc:
        logger.warning("Generation catalog miss for doc_type=%r: %s", doc_type, exc)
        raise HTTPException(status_code=400, detail=str(exc))
    except Exception as exc:
        logger.error("Generation error: %s", exc, exc_info=True)
        raise HTTPException(status_code=500, detail=str(exc))

    lang_idx = {"it": 0, "es": 1, "en": 2}.get(session_lang, 0)
    _catalog_entry = SYSTEM_TEMPLATES_BY_KEY.get(doc_type)
    if _catalog_entry:
        doc_title = _catalog_entry.get("label") or _catalog_entry.get("tipo_atto", "Documento")
        _slug = re.sub(r"[^a-z0-9]+", "_", doc_title.lower()).strip("_")
        filename = f"{_slug}.docx"
    elif doc_type in _DOC_FILENAMES:
        doc_info = _DOC_FILENAMES[doc_type]
        doc_title = doc_info[lang_idx]
        filename = f"{doc_info[3]}.docx"
    else:
        logger.error(
            "generate_download: doc_type=%r not in SYSTEM_TEMPLATES_BY_KEY or _DOC_FILENAMES "
            "— catalog key mismatch. Returning 400.",
            doc_type,
        )
        raise HTTPException(
            status_code=400,
            detail=_build_clarification_message(),
        )
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    tenant_profile = None
    if current_user and current_user.get("tenant_id"):
        tenant_profile = get_tenant_profile_full(current_user["tenant_id"])

    if tenant_profile:
        ph = _placeholder(session_lang)

        logo_path_val = tenant_profile.get("logo_path") or ""
        if logo_path_val and os.path.exists(logo_path_val):
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_run = logo_para.add_run()
            logo_run.add_picture(logo_path_val, width=Inches(2.5))
            # Space after logo before studio name
            logo_para.paragraph_format.space_after = Pt(12)

        name_val = tenant_profile.get("legal_name") or tenant_profile.get("display_name") or ph
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(name_val)
        name_run.bold = True
        name_run.font.size = Pt(11)

        address_parts = [
            tenant_profile.get("address_street") or ph,
            f"{tenant_profile.get('address_postal_code') or ph} {tenant_profile.get('address_city') or ph}",
            tenant_profile.get("address_country") or ph,
        ]
        address_para = doc.add_paragraph(", ".join(address_parts))
        address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        contact_para = doc.add_paragraph(
            f"Tel: {tenant_profile.get('phone') or ph}  |  Web: {tenant_profile.get('website') or ph}"
        )
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        vat_para = doc.add_paragraph(f"P.IVA: {tenant_profile.get('vat_number') or ph}")
        vat_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Separator line and spacing before document title
        sep_para = doc.add_paragraph()
        sep_para.paragraph_format.space_before = Pt(6)
        sep_run = sep_para.add_run("─" * 60)
        sep_run.font.size = Pt(8)
        sep_run.font.color.rgb = RGBColor(180, 180, 180)
        sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sep_para.paragraph_format.space_after = Pt(12)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(doc_title)
    title_run.bold = True
    title_run.font.size = Pt(14)
    doc.add_paragraph()

    for line in _strip_markdown(result["draft"]).split("\n"):
        para = doc.add_paragraph()
        parts = _PH_PATTERN.split(line)
        matches = _PH_PATTERN.findall(line)
        for i, part in enumerate(parts):
            if part:
                para.add_run(part)
            if i < len(matches):
                hl_run = para.add_run(matches[i])
                hl_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# Document-aware intent detection helper
# ---------------------------------------------------------------------------

import re as _re_doc

# Two patterns:
# 1. Filenames inside quotes (single or double) — captures spaces and special chars
# 2. Filenames without quotes — word chars and hyphens only
# Two quoted patterns:
# 1. Quoted with extension — most reliable
# 2. Quoted without extension — only if name is >= 5 chars (avoids matching
#    short quoted strings like "ok", "sì", article references etc.)
_FILENAME_QUOTED_RE = _re_doc.compile(
    r'"([^"<>\n]+?\.(?:pdf|docx?|txt))"', _re_doc.IGNORECASE
)
_FILENAME_QUOTED_NO_EXT_RE = _re_doc.compile(
    r'"([A-Za-z][^"<>\n]{4,}?)"', _re_doc.IGNORECASE
)
_FILENAME_BARE_RE = _re_doc.compile(
    r'\b([\w][\w\-]*\.(?:pdf|docx?|txt))\b', _re_doc.IGNORECASE
)

# Words that strongly signal "I want to READ this document"
_ANALYSE_SIGNALS = {
    # Italian — verb forms
    "riassumi", "riassumimi", "analizza", "analizzami", "spiegami",
    "interpreta", "interpretami", "estrai", "dimmi", "descrivimi",
    "cosa dice", "cosa prevede", "cosa contiene", "cosa stabilisce",
    "chi sono le parti", "qual è", "quali sono", "come funziona",
    # Italian — noun forms (e.g. "fammi un riassunto", "voglio un'analisi")
    "riassunto", "analisi", "sintesi", "estratto", "riepilogo",
    "fammi un", "dammi un", "vorrei un", "vorrei sapere",
    # English
    "summarise", "summarize", "explain", "what does", "who are",
    "summary", "analysis", "extract", "tell me",
    # Spanish/French
    "résume", "explique", "analiza",
}

# Words that strongly signal "I want to GENERATE a document FROM this"
_GENERATE_SIGNALS = {
    "compila", "compilami", "riempi", "riempimi", "usa", "usalo", "usala",
    "usa questo", "usa questa", "genera con", "crea con", "usa come template",
    "usa come modello", "fill", "use this", "use as template",
    # Duplication / copy-with-changes — very common for form workflows
    "duplica", "duplicami", "copia", "copiami",
    "fanne una copia", "fai una copia", "fammi una copia",
    "stesso modulo", "stesso documento", "stessa richiesta",
    "uguale ma con", "uguale con", "identico ma con",
    # Generation verbs
    "generami", "genera", "rigenerami", "rigenera",
    "rifai", "rifammi", "ricreami", "ricrea",
    # Data substitution signals
    "modifica i dati", "cambia i dati", "aggiorna i dati", "sostituisci i dati",
    "con i miei dati", "con i dati di", "con i dati del", "con i dati della",
    "cambia il nome", "cambia solo il nome", "sostituisci il nome", "aggiorna il nome",
    "cambia nome", "cambia cognome", "cambia il cognome",
    "con nome", "con cognome",
    # English equivalents
    "generate", "regenerate", "duplicate", "copy with", "same form", "same document",
    "change the data", "update the data", "replace the data",
    "with my details", "with the details of",
}


def _detect_document_intent(message: str) -> list:
    """
    Return all filenames mentioned in the message as a list.
    Priority:
      1. Quoted strings with file extension (.pdf, .docx, .txt) — most reliable
      2. Bare word filenames with extension
      3. Quoted strings WITHOUT extension — only used as fallback when (1) and (2)
         find nothing, to handle cases like "documento X" without .pdf suffix.
         These are passed to the fuzzy DB lookup which confirms if they match a real file.
    """
    seen = set()
    result = []

    # Step 1: quoted with extension
    for m in _FILENAME_QUOTED_RE.findall(message):
        clean = m.strip()
        if clean and clean.lower() not in seen:
            seen.add(clean.lower())
            result.append(clean)

    # Step 2: bare with extension (on message with quoted filenames removed)
    _msg_without_quoted = _FILENAME_QUOTED_RE.sub("", message)
    for m in _FILENAME_BARE_RE.findall(_msg_without_quoted):
        clean = m.strip()
        if clean and clean.lower() not in seen:
            seen.add(clean.lower())
            result.append(clean)

    # Step 3: quoted without extension — only if steps 1+2 found nothing
    # This handles "documento X" references where user omitted the extension
    if not result:
        for m in _FILENAME_QUOTED_NO_EXT_RE.findall(message):
            clean = m.strip()
            # Skip very common Italian words and short strings
            _SKIP = {"documento", "file", "atto", "contratto", "sì", "no",
                     "ok", "vero", "falso", "questo", "quello", "testo"}
            if clean and clean.lower() not in seen and clean.lower() not in _SKIP:
                seen.add(clean.lower())
                result.append(clean)

    return result


def _classify_doc_intent(
    message: str, session_lang: str, filename: str = "", default: str = "rag"
) -> str:
    """
    Given a message that references an uploaded document, decide whether the
    user wants to:
      - 'analyse'  → read, summarise, extract info from the document
      - 'generate' → use the document as a template to produce a new one
      - 'rag'      → the file mention was incidental; treat as normal RAG query

    Uses signal words first (fast, no LLM). Falls back to a single LLM call
    only when signals are ambiguous.
    """
    lower = message.lower()

    analyse_hit = any(s in lower for s in _ANALYSE_SIGNALS)
    generate_hit = any(s in lower for s in _GENERATE_SIGNALS)

    if analyse_hit and not generate_hit:
        return "analyse"
    if generate_hit and not analyse_hit:
        return "generate"

    # Ambiguous or no signal — ask the LLM
    system = (
        "Sei un classificatore di intenzioni. L'utente ha menzionato un documento caricato. "
        "Rispondi con UNA SOLA parola:\n"
        "- 'analyse' se l'utente vuole leggere, riassumere o estrarre informazioni dal documento\n"
        "- 'generate' se l'utente vuole usare il documento come template per generarne uno nuovo, "
        "compilarlo con dati diversi, duplicarlo con modifiche, creare una copia con dati aggiornati, "
        "o produrre un nuovo documento simile con informazioni diverse\n"
        "- 'rag' se la menzione del file è incidentale e la domanda riguarda altro\n"
        "Rispondi SOLO con una di queste tre parole."
    )
    _human = message
    if filename:
        _human += f"\nNome file: {filename}"
    try:
        result = _call_chat(
            [SystemMessage(content=system), HumanMessage(content=_human)],
            max_tokens=5,
        ).strip().lower()
        if result in ("analyse", "generate", "rag"):
            return result
    except Exception:
        pass
    return default


def _classify_top_level_intent(message: str, session_lang: str) -> str:
    """
    Classify the top-level intent of a message as one of:
      - 'generate'  → user wants a legal document created from scratch
      - 'compare'   → user wants to compare two specific uploaded documents
      - 'rag'       → question, analysis, or anything else

    Strategy:
    1. Fast path: strong unambiguous generation triggers (no LLM needed)
    2. Fast path: explicit comparison — must have comparison verb AND two quoted names
    3. LLM call only for ambiguous cases that have some signal
    4. Default to 'rag' when uncertain
    """
    lower = message.lower()

    # ── Fast path: strong generation triggers ─────────────────────────────
    _GEN_STRONG = [
        "redigimi", "generami", "scrivimi", "preparami", "creami", "fammi",
        "elaborami", "stendimi", "formulami", "producimi",
        "scrivi un", "scrivi una", "redigi un", "redigi una",
        "genera un", "genera una", "crea un", "crea una",
        "prepara un", "prepara una",
        "scrivi atto", "scrivi contratto", "scrivi memoria", "scrivi ricorso",
        "scrivi istanza", "scrivi diffida", "scrivi dichiarazione",
        "redigi atto", "redigi contratto", "redigi memoria", "redigi ricorso",
        "genera atto", "genera documento", "genera contratto",
        "draft a", "draft an", "write a contract", "write a letter",
        "redacta un", "redacta una",
        "ho bisogno di un contratto", "ho bisogno di una lettera",
        "ho bisogno di un atto", "ho bisogno di una memoria",
        "voglio un contratto", "voglio una lettera", "voglio un atto",
        "vorrei un contratto", "vorrei una lettera", "vorrei un atto",
    ]
    if any(t in lower for t in _GEN_STRONG):
        return "generate"

    # ── Fast path: comparison — verb AND two quoted document names ─────────
    _has_comparison_verb = bool(re.search(
        r'\b(confronta|paragona|compara|compare|versus|vs\.?)\b'
        r'|confronto\s+(tra|fra|dei|di|delle|degli)\b'
        r'|differenze?\s+(tra|fra|dei|di)\b',
        message, re.IGNORECASE
    ))
    _quoted_names = re.findall(r'"[^"]{3,}"', message)
    if _has_comparison_verb and len(_quoted_names) >= 2:
        return "compare"

    # ── No signal at all → RAG immediately, no LLM call ───────────────────
    _has_any_signal = bool(re.search(
        r'\b(scrivi|redigi|genera|crea|prepara|elabora|formula|'
        r'stendi|produce|draft|write|create|voglio|vorrei|'
        r'ho bisogno|confronta|paragona|compara|differenze)\b',
        message, re.IGNORECASE
    ))
    if not _has_any_signal:
        return "rag"

    # ── LLM fallback for ambiguous cases ──────────────────────────────────
    system = (
        "Sei un classificatore di intenzioni per un assistente legale italiano. "
        "Classifica il messaggio in UNA di queste categorie:\n\n"
        "- 'generate': l'utente vuole creare un documento legale da zero "
        "(contratto, memoria difensiva, atto, ricorso, istanza, nomina, ecc.)\n"
        "- 'compare': l'utente vuole confrontare DUE documenti specifici già "
        "caricati — richiede esplicitamente due nomi di file tra virgolette\n"
        "- 'rag': domanda su temi legali, analisi, spiegazione, o qualsiasi "
        "altro caso\n\n"
        "REGOLE IMPORTANTI:\n"
        "- Parole come 'confronto', 'differenze', 'paragone' usate in contesto "
        "legale generico (es. 'disponibile ad ogni confronto') NON sono richieste "
        "di confronto documenti → usa 'rag'\n"
        "- Solo 'generate' quando l'utente vuole un documento NUOVO creato per lui\n"
        "- In caso di dubbio usa 'rag'\n\n"
        "Rispondi SOLO con una parola: generate, compare, o rag."
    )
    try:
        result = _call_chat(
            [SystemMessage(content=system), HumanMessage(content=message[:600])],
            max_tokens=5,
        ).strip().lower()
        if result in ("generate", "compare", "rag"):
            return result
    except Exception:
        pass
    return "rag"


@app.post("/api/chat", response_model=ChatResponse)
async def chat(request: ChatRequest, current_user: Optional[dict] = Depends(get_current_user), db: Session = Depends(get_db)):
    """Send a message and get a response.

    If session_id is provided, continues the conversation.
    If omitted, a new session is created automatically.
    If the message is a generation request, redirects to the opposition act generation flow.
    """
    if current_user:
        enforce_tenant_product_access(db, uuid.UUID(str(current_user["tenant_id"])))

    session_id = request.session_id
    _uid = current_user["sub"] if current_user else None
    _tid = current_user.get("tenant_id") if current_user else None
    if not session_id:
        session = chatbot.create_session(user_id=_uid, tenant_id=_tid)
        session_id = session.session_id

    _req_start = time.time()
    vlog("request_start", {"session_id": session_id, "message_length": len(request.message)})

    # ── Document-aware branch ─────────────────────────────────────────────
    # Detects filename(s) in the message and routes to analyse, compare,
    # or generate — before the generation and RAG paths.
    if current_user:
        import os as _os
        from ..utils.document import extract_text_from_file as _extract
        from ..db.base import get_db as _get_db

        _mentioned_names = _detect_document_intent(request.message)

        # Fallback: anaphoric document reference ("questo file", "il documento", etc.)
        # — resolve to the last document the user interacted with in this session.
        _ANAPHORIC_DOC_RE = re.compile(
            r'\b(questo file|questa file|questo documento|questa documento|'
            r'il file|il documento|quel file|quel documento|'
            r'questo modulo|il modulo|quel modulo|'
            r'this file|this document|the file|the document)\b',
            re.IGNORECASE,
        )
        if not _mentioned_names and _ANAPHORIC_DOC_RE.search(request.message):
            _anaphoric_session = chatbot.get_session(session_id, user_id=_uid)
            if _anaphoric_session:
                for _prev_msg in reversed(_anaphoric_session.messages):
                    _doc_name = (_prev_msg.metadata or {}).get("document_name")
                    if _doc_name:
                        _mentioned_names = [_doc_name]
                        break

        # Fallback: if no filename detected but FE passed an explicit document_id,
        # treat it as a single-document analyse request
        if not _mentioned_names and request.document_id:
            _db_gen2 = _get_db()
            _db2 = next(_db_gen2)
            try:
                _explicit_doc = get_user_document(
                    _db2,
                    uuid.UUID(request.document_id),
                    uuid.UUID(_uid),
                    uuid.UUID(_tid),
                )
            finally:
                try:
                    _db_gen2.close()
                except Exception:
                    pass
            if _explicit_doc:
                _mentioned_names = [_explicit_doc.original_filename]

        # Explicit multi-doc list: look up each UUID directly and add filenames
        if request.document_ids:
            _db_gen_ids = _get_db()
            _db_ids = next(_db_gen_ids)
            try:
                for _did in request.document_ids:
                    try:
                        _edoc = get_user_document(
                            _db_ids,
                            uuid.UUID(_did),
                            uuid.UUID(_uid),
                            uuid.UUID(_tid),
                        )
                        if _edoc and _edoc.original_filename not in _mentioned_names:
                            _mentioned_names.append(_edoc.original_filename)
                    except (ValueError, Exception):
                        pass
            finally:
                try:
                    _db_gen_ids.close()
                except Exception:
                    pass

        if _mentioned_names:
            _db_gen = _get_db()
            _db = next(_db_gen)
            try:
                _matched_docs = []
                _not_found = []
                for _name in _mentioned_names:
                    _doc = find_user_document_by_name(
                        _db,
                        uuid.UUID(_uid),
                        uuid.UUID(_tid),
                        _name,
                    )
                    if _doc:
                        # Avoid duplicates if two filename mentions resolve to same doc
                        if not any(d.id == _doc.id for d in _matched_docs):
                            _matched_docs.append(_doc)
                    else:
                        _not_found.append(_name)
            finally:
                try:
                    _db_gen.close()
                except Exception:
                    pass

            # ── Expiry check ──────────────────────────────────────────────
            # If filename(s) were mentioned but no live doc found, check if
            # any of them exist but are expired — return a clear message
            # instead of falling through to RAG or document_id fallback.
            if not _matched_docs and _not_found:
                _db_exp_gen = _get_db()
                _db_exp = next(_db_exp_gen)
                _expired_msgs = []
                try:
                    for _exp_name in _not_found:
                        _exp_doc = find_expired_user_document_by_name(
                            _db_exp,
                            uuid.UUID(_uid),
                            uuid.UUID(_tid),
                            _exp_name,
                        )
                        if _exp_doc and _exp_doc.expires_at:
                            _exp_date = _exp_doc.expires_at.strftime("%d/%m/%Y")
                            _expired_msgs.append(
                                f"Il documento «{_exp_doc.original_filename}» "
                                f"è scaduto il {_exp_date}. "
                                f"Per continuare l'analisi è necessario caricarlo di nuovo."
                            )
                finally:
                    try:
                        _db_exp_gen.close()
                    except Exception:
                        pass

                if _expired_msgs:
                    _session = chatbot.get_session(session_id, user_id=_uid)
                    if not _session:
                        _session = ChatSession(
                            session_id=session_id, user_id=_uid, tenant_id=_tid
                        )
                        chatbot._sessions[session_id] = _session
                    _exp_answer = "\n\n".join(_expired_msgs)
                    _session.add_message("user", request.message)
                    _session.add_message("assistant", _exp_answer)
                    return ChatResponse(
                        session_id=session_id,
                        answer=_exp_answer,
                        original_query=request.message,
                        resolved_query=request.message,
                        session_language=_session.session_language,
                        status_messages=["document_expired"],
                    )

            if _matched_docs:
                _session = chatbot.get_session(session_id, user_id=_uid)
                if not _session:
                    _session = ChatSession(
                        session_id=session_id, user_id=_uid, tenant_id=_tid
                    )
                    chatbot._sessions[session_id] = _session
                _session_lang = _session.session_language
                _doc_settings = get_user_settings(_uid) if _uid else {"tone": 2, "standing": 2, "response_length": 2}

                _lang_note = {
                    "es": "Rispondi in spagnolo.",
                    "en": "Reply in English.",
                }.get(_session_lang, "Rispondi in italiano.")

                _MAX_CHARS_PER_DOC = 50_000  # full doc — windowing handles truncation

                # ── TWO OR MORE DOCS → comparison path ────────────────────
                if len(_matched_docs) >= 2:
                    _doc_texts = []
                    _missing = []
                    for _doc in _matched_docs:
                        if not _os.path.exists(_doc.storage_path):
                            _missing.append(_doc.original_filename)
                            continue
                        try:
                            _text = _extract(_doc.storage_path)
                            if _text and _text.strip():
                                _doc_texts.append((_doc.original_filename, _text[:_MAX_CHARS_PER_DOC]))
                            else:
                                _missing.append(_doc.original_filename)
                        except Exception as _exc:
                            logger.warning(
                                "chat: failed to extract text from %s: %s",
                                _doc.storage_path, _exc,
                            )
                            _missing.append(_doc.original_filename)

                    if not _doc_texts:
                        answer = (
                            "Non riesco a estrarre il testo dai documenti indicati. "
                            "Prova a ricaricarli."
                        )
                    elif len(_doc_texts) == 1:
                        answer = (
                            f"Non riesco a trovare o leggere il documento "
                            f"'{_missing[0]}'. Puoi ricaricarlo?"
                        )
                    else:
                        # Sequential extraction: one LLM call per document to extract
                        # only the relevant section, then a final call to compare.
                        # This keeps every call within the 10k token context limit
                        # regardless of document size.
                        _CHUNK = 3_000  # chars per doc chunk (~750 tokens)

                        _WINDOW = 500

                        def _get_dynamic_windows(text: str, keywords: list) -> str:
                            """
                            Heading-aware extraction:
                            1. Split document into named article/section blocks
                            2. Return blocks whose heading or body contains a keyword
                            3. Fall back to keyword-window approach if no headings found
                            """
                            import re as _re

                            # Match article/section headings in Italian legal docs.
                            # Handles: ART. 6, Art. 6, ARTICOLO 6, 6., 6.1, SEZIONE 5
                            # with various dash types (-, –, —) and trailing text.
                            # Only match explicit ART./ARTICOLO headings — avoids
                            # false positives from postcodes, isolated numbers, etc.
                            _HEADING_RE = _re.compile(
                                r'ART(?:ICOLO)?\.?\s*\d+\w*'
                                r'(?:\s*[\-–—]\s*[A-ZÀÈÌÒÙ][^\n]{0,80})?'
                                r'(?:\s{2,}|\n)',
                                _re.IGNORECASE
                            )

                            _matches = list(_HEADING_RE.finditer(text))
                            _blocks = []
                            for i, m in enumerate(_matches):
                                h_start = m.start()
                                h_end = m.end()
                                b_end = _matches[i + 1].start() if i + 1 < len(_matches) else len(text)
                                heading = m.group().strip()
                                body = text[h_end:b_end].strip()
                                # Count keyword hits in this block
                                combined = (heading + " " + body).lower()
                                hits = sum(combined.count(kw.lower()) for kw in keywords)
                                _blocks.append((h_start, hits, heading, body))

                            # Stem Italian keywords to first 7 chars for fuzzy matching.
                            # Use 7 chars minimum to avoid false positives from short stems
                            # (e.g. 'manage' matching 'managed operations' in preambles).
                            # Only keep stems from keywords >= 7 chars long.
                            _GENERIC_STEMS = {
                                'manage', 'gestio', 'clause', 'claus', 'breach',
                                'penalt', 'paymen', 'remune',
                            }
                            _stems = list({
                                kw[:7].lower()
                                for kw in keywords
                                if len(kw) >= 7 and kw[:7].lower() not in _GENERIC_STEMS
                            })
                            # Always include these high-signal Italian legal stems
                            _stems += ['compenso', 'corrispo', 'pagamen', 'inadem',
                                       'risoluz', 'recesso', 'sanzion', 'penale']
                            _stems = list(set(_stems))

                            def _hits(text_lower: str) -> int:
                                return sum(text_lower.count(s) for s in _stems)

                            if _blocks:
                                _relevant = []
                                for pos, _, heading, body in _blocks:
                                    heading_lower = heading.lower()
                                    body_lower = body.lower()
                                    heading_hits = _hits(heading_lower)
                                    body_hits = _hits(body_lower)
                                    total = heading_hits + body_hits
                                    if total == 0:
                                        continue
                                    # Exclude blocks where keywords appear only incidentally:
                                    # heading has no keyword hit AND density is < 1 per 300 chars
                                    density = body_hits / max(len(body), 1) * 300
                                    if heading_hits == 0 and density < 1.0:
                                        continue
                                    _relevant.append((pos, total, f"{heading}\n{body[:800]}"))
                                if _relevant:
                                    # Sort by hit density (hits per char) not raw hits,
                                    # so long preamble blocks don't outrank short precise articles
                                    _relevant.sort(key=lambda x: (
                                        -(_hits(x[2].lower()) / max(len(x[2]), 1)),
                                        x[0]
                                    ))
                                    return "\n\n---\n\n".join(
                                        block for _, _, block in _relevant
                                    )[:4_000]

                            # Fallback: stem-based keyword-window search
                            _wins = []
                            _seen = []
                            _text_lower = text.lower()
                            for stem in _stems:
                                idx = 0
                                while True:
                                    idx = _text_lower.find(stem, idx)
                                    if idx == -1:
                                        break
                                    s = max(0, idx - 200)
                                    e = min(len(text), idx + _WINDOW)
                                    if not any(max(s, a) < min(e, b) for a, b in _seen):
                                        _wins.append((s, text[s:e].strip()))
                                        _seen.append((s, e))
                                    idx += len(stem)
                            _wins.sort(key=lambda x: x[0])
                            return "\n\n---\n\n".join(w for _, w in _wins)[:4_000]

                        try:
                            loop = asyncio.get_event_loop()

                            # ── Call 1: extract search terms from the user's question ──
                            # Tiny call — just needs 5-8 Italian legal keywords to search for.
                            _kw_system = (
                                "Sei un assistente legale italiano. "
                                "Data una domanda su documenti legali, restituisci "
                                "una lista di 6-10 parole chiave italiane (e inglesi se il "
                                "documento potrebbe essere in inglese) da cercare nei documenti "
                                "per trovare le clausole rilevanti. "
                                "Restituisci SOLO le parole chiave separate da virgola, "
                                "senza testo aggiuntivo. "
                                "Esempio: corrispettivo, pagamento, fattura, inadempimento, risoluzione"
                            )
                            _kw_human = f"Domanda: {request.message}"

                            _kw_raw = await loop.run_in_executor(
                                None,
                                partial(
                                    _call_chat,
                                    [
                                        SystemMessage(content=_kw_system),
                                        HumanMessage(content=_kw_human),
                                    ],
                                    80,
                                ),
                            )
                            _keywords = [
                                k.strip()
                                for k in _kw_raw.replace("\n", ",").split(",")
                                if k.strip() and len(k.strip()) > 2
                            ]
                            logger.warning(
                                "chat: multi-doc comparison keywords for question %r: %r",
                                request.message[:60], _keywords,
                            )

                            # ── Keyword search (no LLM) ───────────────────────────────
                            _doc_extracts = []
                            for fname, text in _doc_texts:
                                _windows = _get_dynamic_windows(text, _keywords)
                                if not _windows:
                                    # Fallback: first 2000 chars if no keyword hits
                                    _windows = text[:2_000]
                                _doc_extracts.append((fname, _windows))

                            # ── Call 2: comparison ────────────────────────────────────
                            # The model cannot reliably acknowledge verbatim context.
                            # Instead: present the extracted windows directly as the answer,
                            # then ask the model to write only a brief concluding comparison
                            # sentence — not to search or summarise the extracts.
                            _extract_block = "\n\n".join(
                                f"**{fname}**\n\n{extract}"
                                for fname, extract in _doc_extracts
                            )
                            _compare_system = (
                                legal_consultant_system_prefix(
                                    _session_lang,
                                    tone=_doc_settings["tone"],
                                    standing=_doc_settings["standing"],
                                ) + " "
                                "Ti vengono forniti due blocchi di testo già estratti da documenti "
                                "legali. Sintetizza "
                                "le principali differenze tra i due documenti relativamente alla "
                                "domanda dell'utente. "
                                "Non ripetere il contenuto degli estratti — solo le differenze chiave. "
                                + _lang_note
                                + "\n\n" + _LENGTH.get(_doc_settings["response_length"], _LENGTH[2])
                            )
                            _compare_human = (
                                "\n\n".join(
                                    f"DOCUMENTO {i+1} — {fname}:\n{extract}"
                                    for i, (fname, extract) in enumerate(_doc_extracts)
                                )
                                + f"\n\nDomanda: {request.message}"
                            )

                            try:
                                _summary = await loop.run_in_executor(
                                    None,
                                    partial(
                                        _call_chat,
                                        [
                                            SystemMessage(content=_compare_system),
                                            HumanMessage(content=_compare_human),
                                        ],
                                        2500,
                                    ),
                                )
                            except Exception:
                                _summary = ""

                            # Chat answer: show only the summary — clause extraction stays server-side
                            _doc_names = " e ".join(f"«{f}»" for f, _ in _doc_extracts)
                            if _summary:
                                answer = _summary
                            else:
                                answer = (
                                    f"Ho analizzato i documenti {_doc_names} "
                                    "ma non è stato possibile generare una sintesi. "
                                    "Prova a riformulare la domanda."
                                )
                        except Exception as _exc:
                            _exc_str = str(_exc)
                            if "maximum context length" in _exc_str or "input_tokens" in _exc_str or "400" in _exc_str:
                                logger.warning("chat: multi-doc comparison context overflow: %s", _exc)
                                answer = (
                                    "I documenti selezionati sono troppo lunghi per essere confrontati insieme. "
                                    "Prova a fare domande più specifiche su singole clausole, "
                                    "oppure carica versioni più brevi dei documenti."
                                )
                            else:
                                logger.error(
                                    "chat: multi-doc comparison LLM call failed: %s",
                                    _exc, exc_info=True,
                                )
                                answer = (
                                    "Si è verificato un errore durante il confronto dei documenti. "
                                    "Riprova tra qualche istante."
                                )

                    _session.add_message("user", request.message, metadata={
                        "documents": [{"document_id": str(d.id), "document_name": d.original_filename} for d in _matched_docs],
                    })
                    if len(_session.messages) == 1:
                        _session.title = _generate_session_title(request.message)
                    _session.add_message("assistant", answer)
                    chatbot._save_sessions()
                    return ChatResponse(
                        session_id=session_id,
                        answer=answer,
                        original_query=request.message,
                        resolved_query=request.message,
                        session_language=_session_lang,
                        status_messages=["document_analyse_mode"],
                        title=_session.title,
                    )

                # ── ONE DOC → analyse or generate ─────────────────────────
                _matched_doc = _matched_docs[0]
                doc_intent = _classify_doc_intent(
                    request.message,
                    _session_lang,
                    filename=_matched_doc.original_filename,
                    default="analyse" if getattr(_matched_doc, "document_role", "document") == "document" else "rag",
                )

                if doc_intent == "generate":
                    # User wants to fill/duplicate their own uploaded document.
                    # Bypass the from-scratch template picker and generate directly.
                    if not _os.path.exists(_matched_doc.storage_path):
                        _fill_err = (
                            "Non riesco a trovare il file sul server. "
                            "Prova a caricarlo di nuovo."
                        )
                        _session.add_message("user", request.message, metadata={
                            "document_id": str(_matched_doc.id),
                            "document_name": _matched_doc.original_filename,
                        })
                        if len(_session.messages) == 1:
                            _session.title = _generate_session_title(request.message)
                        _session.add_message("assistant", _fill_err)
                        chatbot._save_sessions()
                        return ChatResponse(
                            session_id=session_id,
                            answer=_fill_err,
                            original_query=request.message,
                            resolved_query=request.message,
                            session_language=_session_lang,
                            status_messages=["generation_mode"],
                            title=_session.title,
                        )
                    _fill_ext = _os.path.splitext(_matched_doc.storage_path)[1].lower()
                    _fill_is_pdf = _fill_ext == ".pdf"
                    _fill_carta = get_tenant_profile_full(_tid) if _tid else None
                    _fill_session_msgs = [{"role": m.role, "content": m.content} for m in _session.messages]
                    try:
                        _fill_elements = (
                            _extract_pdf_elements(_matched_doc.storage_path) if _fill_is_pdf
                            else _extract_docx_elements(_matched_doc.storage_path)
                        )
                        if not _fill_elements:
                            raise ValueError("Nessun elemento estratto dal documento.")
                        _fill_map = _fill_template_gaps(
                            _fill_elements, request.message, _fill_carta,
                            _session_lang, _fill_session_msgs,
                            docx_path=None if _fill_is_pdf else _matched_doc.storage_path,
                        )
                        if _fill_is_pdf:
                            _fill_bytes = _build_docx_from_pdf_elements(_fill_elements, _fill_map)
                        else:
                            _fill_bytes = _apply_fill_to_docx(_matched_doc.storage_path, _fill_map)
                    except Exception as _fill_exc:
                        logger.error("chat: user-template fill failed: %s", _fill_exc, exc_info=True)
                        _fill_err = (
                            "Si è verificato un errore durante la generazione del documento. "
                            "Riprova tra qualche istante."
                        )
                        _session.add_message("user", request.message, metadata={
                            "document_id": str(_matched_doc.id),
                            "document_name": _matched_doc.original_filename,
                        })
                        if len(_session.messages) == 1:
                            _session.title = _generate_session_title(request.message)
                        _session.add_message("assistant", _fill_err)
                        chatbot._save_sessions()
                        return ChatResponse(
                            session_id=session_id,
                            answer=_fill_err,
                            original_query=request.message,
                            resolved_query=request.message,
                            session_language=_session_lang,
                            status_messages=["generation_mode"],
                            title=_session.title,
                        )
                    # Persist the filled document as a new user_document record
                    _fill_base = _os.path.splitext(_matched_doc.original_filename)[0]
                    _fill_outname = f"{_fill_base}_compilato.docx"
                    _fill_doc_id, _fill_doc_name = None, None
                    if _uid and _tid:
                        try:
                            from ..constants import PRIVATE_DOCS_BASE as _FILL_BASE
                            _fill_folder = _os.path.join(_FILL_BASE, _tid, _uid)
                            _os.makedirs(_fill_folder, exist_ok=True)
                            _fill_storage = _os.path.join(_fill_folder, f"{uuid.uuid4()}.docx")
                            with open(_fill_storage, "wb") as _ff:
                                _ff.write(_fill_bytes)
                            _fill_db_gen = _get_db()
                            _fill_db = next(_fill_db_gen)
                            try:
                                _fill_rec = create_user_document(
                                    _fill_db,
                                    user_id=uuid.UUID(_uid),
                                    tenant_id=uuid.UUID(_tid),
                                    original_filename=_fill_outname,
                                    storage_path=_fill_storage,
                                    file_size_bytes=len(_fill_bytes),
                                    scope="personal",
                                    document_role="generated",
                                    expires_at=None,
                                )
                                _fill_doc_id = str(_fill_rec.id)
                                _fill_doc_name = _fill_outname
                            finally:
                                try:
                                    _fill_db_gen.close()
                                except Exception:
                                    pass
                        except Exception as _persist_exc:
                            logger.warning("chat: failed to persist filled doc: %s", _persist_exc)
                    _fill_missing = _summarise_da_compilare(_fill_map, _fill_elements, _session_lang)
                    _fill_confirmation = _build_generation_confirmation(_session_lang, _fill_missing or None)
                    _session.add_message("user", request.message, metadata={
                        "document_id": str(_matched_doc.id),
                        "document_name": _matched_doc.original_filename,
                    })
                    if len(_session.messages) == 1:
                        _session.title = _generate_session_title(request.message)
                    _session.add_message(
                        "assistant",
                        _fill_confirmation,
                        metadata={
                            **({"generated_document_id": _fill_doc_id, "generated_document_name": _fill_doc_name} if _fill_doc_id else {}),
                        },
                    )
                    chatbot._save_sessions()
                    return ChatResponse(
                        session_id=session_id,
                        answer=_fill_confirmation,
                        original_query=request.message,
                        resolved_query=request.message,
                        session_language=_session_lang,
                        status_messages=["generation_mode"],
                        title=_session.title,
                        generated_document_id=_fill_doc_id,
                        generated_document_name=_fill_doc_name,
                    )

                elif doc_intent == "analyse":
                    if not _os.path.exists(_matched_doc.storage_path):
                        answer = (
                            "Non riesco a trovare il file sul server. "
                            "Prova a caricarlo di nuovo."
                        )
                    else:
                        try:
                            _text = _extract(_matched_doc.storage_path)
                        except Exception as _exc:
                            _text = None
                            logger.warning(
                                "chat: failed to extract text from %s: %s",
                                _matched_doc.storage_path, _exc,
                            )

                        if not _text or not _text.strip():
                            answer = (
                                f"Il documento '{_matched_doc.original_filename}' "
                                "non contiene testo estraibile."
                            )
                        else:
                            _truncated = _text[:12_000]
                            _was_cut = len(_text) > 12_000
                            _system = (
                                legal_consultant_system_prefix(
                                    _session_lang,
                                    tone=_doc_settings["tone"],
                                    standing=_doc_settings["standing"],
                                ) + " "
                                "L'utente ti ha fornito il testo di un documento privato "
                                f"('{_matched_doc.original_filename}'). "
                                "Rispondi alla domanda dell'utente basandoti ESCLUSIVAMENTE "
                                "sul contenuto del documento. "
                                "Non inventare fatti non presenti nel testo. "
                                "Se l'informazione richiesta non è nel documento, dillo chiaramente. "
                                "Rispondi in modo conversazionale e naturale — non usare JSON, "
                                "non aggiungere intestazioni non necessarie. "
                                "STRUTTURA MODULI ITALIANI: nei moduli italiani alcune frasi fisse "
                                "sono etichette di campo — il valore segue immediatamente sulla stessa riga "
                                "(e può essere anch'esso in maiuscolo). "
                                "Etichette comuni: 'COGNOME E NOME', 'COGNOME', 'NOME', "
                                "'DATA E LUOGO DI NASCITA', 'DATA DI NASCITA', 'LUOGO DI NASCITA', "
                                "'RESIDENZA', 'DOMICILIO', 'CODICE FISCALE', 'FIRMA', 'DATA', "
                                "'IL SOTTOSCRITTO', 'LA SOTTOSCRITTA', 'INDIRIZZO', 'COMUNE', 'PROVINCIA'. "
                                "Anche un testo seguito da ':' è sempre un'etichetta. "
                                "Esempi: 'COGNOME E NOME CICCIO ELLE' → nome: CICCIO ELLE; "
                                "'DATA E LUOGO DI NASCITA 14 LUGLIO 1990 - FOGGIA' → nato il 14/07/1990 a Foggia; "
                                "'RESIDENZA VIA ROMA 1' → residente in Via Roma 1; "
                                "'NOME: MARIO ROSSI' → nome: MARIO ROSSI (MARIO ROSSI è il valore, non l'etichetta). "
                                "Non includere mai l'etichetta del campo come parte del valore quando rispondi. "
                                + _lang_note
                                + "\n\n" + _LENGTH.get(_doc_settings["response_length"], _LENGTH[2])
                            )
                            _human = (
                                f"Documento:\n\n{_truncated}"
                                + ("\n\n[Il documento è stato troncato per limiti di lunghezza.]"
                                   if _was_cut else "")
                                + f"\n\nDomanda dell'utente: {request.message}"
                            )
                            try:
                                loop = asyncio.get_event_loop()
                                answer = await loop.run_in_executor(
                                    None,
                                    partial(
                                        _call_chat,
                                        [
                                            SystemMessage(content=_system),
                                            HumanMessage(content=_human),
                                        ],
                                        2500,
                                    ),
                                )
                            except Exception as _exc:
                                _exc_str = str(_exc)
                                if "maximum context length" in _exc_str or "input_tokens" in _exc_str or "400" in _exc_str:
                                    logger.warning("chat: document analyse context overflow: %s", _exc)
                                    answer = (
                                        "Il documento è troppo lungo per essere analizzato interamente. "
                                        "Prova a fare una domanda più specifica su una sezione particolare."
                                    )
                                else:
                                    logger.error(
                                        "chat: document analyse LLM call failed: %s",
                                        _exc, exc_info=True,
                                    )
                                    answer = (
                                        "Si è verificato un errore durante l'analisi del documento. "
                                        "Riprova tra qualche istante."
                                    )

                    _session.add_message("user", request.message, metadata={
                        "document_id": str(_matched_doc.id),
                        "document_name": _matched_doc.original_filename,
                    })
                    if len(_session.messages) == 1:
                        _session.title = _generate_session_title(request.message)
                    _session.add_message("assistant", answer)
                    chatbot._save_sessions()
                    return ChatResponse(
                        session_id=session_id,
                        answer=answer,
                        original_query=request.message,
                        resolved_query=request.message,
                        session_language=_session_lang,
                        status_messages=["document_analyse_mode"],
                        title=_session.title,
                    )
                # ── RAG intent: fall through to normal pipeline ────────────

            elif _not_found:
                # Filenames mentioned but no matching document found in user_documents.
                # Fall through to RAG — corpus may have something relevant.
                logger.info(
                    "chat: filename(s) %r mentioned but not found in user_documents "
                    "— falling through to RAG",
                    _not_found,
                )

    _top_intent = _classify_top_level_intent(request.message, "it")
    if _top_intent == "generate" or is_generation_request(request.message):
        session = chatbot.get_session(session_id, user_id=_uid)
        if not session:
            session = ChatSession(session_id=session_id, user_id=_uid, tenant_id=_tid)
            chatbot._sessions[session_id] = session
        session_lang = session.session_language
        # A generation turn must not silently drop a calculation that is still
        # collecting inputs: this branch returns without ever reaching the RAG
        # graph, and session.py looks for pending_calculation only on the last
        # assistant message. Carrying it across is safe — calculation_node
        # escapes to normal RAG when the next message turns out not to answer
        # the open slot.
        _pending_calc = last_pending_calculation(session)
        _carry_calc = {"pending_calculation": _pending_calc} if _pending_calc else {}
        doc_type = classify_document_type(request.message, session_lang)
        if doc_type == "unknown":
            _template_result = classify_system_template(request.message, session_lang, top_k=5)
            logger.info("DEBUG picker: top_k=5 result: %s", _template_result)
            if isinstance(_template_result, str):
                # Defensive: classify_system_template(top_k>1) is documented to always
                # return a list, but if it ever returns a string, treat it as a single
                # match and skip the picker logic entirely.
                doc_type = _template_result
            elif len(_template_result) >= 2:
                _variant_prompt = "Quale variante preferisci?"
                session.add_message("user", request.message)
                if len(session.messages) == 1:
                    session.title = _generate_session_title(request.message)
                session.add_message("assistant", _variant_prompt, metadata=_carry_calc or None)
                chatbot._save_sessions()
                return ChatResponse(
                    session_id=session_id,
                    answer=_variant_prompt,
                    original_query=request.message,
                    resolved_query=request.message,
                    session_language=session_lang,
                    status_messages=["Quale variante preferisci?"],
                    title=session.title,
                    generation_candidates=_template_result,
                )
            elif len(_template_result) == 1:
                doc_type = _template_result[0]["key"]
            else:
                doc_type = "unknown"
        if doc_type == "unknown":
            clarification = _build_clarification_message()
            session.add_message("user", request.message)
            if len(session.messages) == 1:
                session.title = _generate_session_title(request.message)
            session.add_message("assistant", clarification, metadata=_carry_calc or None)
            chatbot._save_sessions()
            return ChatResponse(
                session_id=session_id,
                answer=clarification,
                original_query=request.message,
                resolved_query=request.message,
                session_language=session_lang,
                status_messages=["generation_mode"],
                title=session.title,
            )
        cached = _get_cached_sections(session)
        session.add_message("user", request.message)
        if len(session.messages) == 1:
            session.title = _generate_session_title(request.message)
        try:
            loop = asyncio.get_event_loop()
            gen_result = await loop.run_in_executor(
                None, partial(_run_generation_sync, request.message, session_lang, doc_type, cached)
            )
        except Exception as exc:
            logger.error("Generation error: %s", exc, exc_info=True)
            raise HTTPException(status_code=500, detail=str(exc))
        gen_result["draft"] = _strip_markdown(gen_result["draft"])
        _gen_doc_id, _gen_doc_name = None, None
        logger.info("chat: generation complete doc_type=%r uid=%r tid=%r draft_len=%d",
                    doc_type, _uid, _tid, len(gen_result["draft"]))
        if _uid and _tid:
            _gen_doc_id, _gen_doc_name = _persist_generated_docx(
                gen_result["draft"], doc_type, _uid, _tid,
                case_details=gen_result.get("case_details"),
            )
            logger.info("chat: persist result doc_id=%r name=%r", _gen_doc_id, _gen_doc_name)
        else:
            logger.warning("chat: skipping persist — uid=%r tid=%r", _uid, _tid)
        _confirmation = _build_generation_confirmation(session_lang)
        session.add_message(
            "assistant",
            _confirmation,
            metadata={
                "sources": gen_result["sources"],
                **_carry_calc,
                **({"generated_document_id": _gen_doc_id, "generated_document_name": _gen_doc_name} if _gen_doc_id else {}),
            },
        )
        chatbot._save_sessions()
        return ChatResponse(
            session_id=session_id,
            answer=_confirmation,
            original_query=request.message,
            resolved_query=request.message,
            session_language=session_lang,
            status_messages=["generation_mode"],
            title=session.title,
            draft=gen_result["draft"],
            generated_document_id=_gen_doc_id,
            generated_document_name=_gen_doc_name,
        )

    try:
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(
            None, partial(chatbot.chat, session_id, request.message,
                          user_id=current_user.get("sub") if current_user else None,
                          tenant_id=current_user.get("tenant_id") if current_user else None)
        )
    except PermissionError:
        raise HTTPException(status_code=403, detail="This session does not belong to your account")
    except Exception as e:
        logger.error("Chat error: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

    vlog("request_end", {"session_id": session_id, "message_length": len(request.message)}, (time.time() - _req_start) * 1000)
    return ChatResponse(**result)


# ---------------------------------------------------------------------------
# Server runner
# ---------------------------------------------------------------------------

def start_server(host: str = "0.0.0.0", port: int = 8000):
    """Start the FastAPI server with uvicorn."""
    import uvicorn
    uvicorn.run(app, host=host, port=port)
