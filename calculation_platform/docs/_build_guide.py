#!/usr/bin/env python3
"""Build the Italian project-manager guide for the calculation platform."""

from __future__ import annotations

from datetime import date
from pathlib import Path
import shutil
import subprocess
import textwrap

import yaml
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor, Twips


DOCS_DIR = Path(__file__).resolve().parent
PLATFORM_DIR = DOCS_DIR.parent
FORMULA_DIR = PLATFORM_DIR / "formula_packs"
PARAMETER_DIR = PLATFORM_DIR / "parameters"
ASSETS_DIR = DOCS_DIR / "assets"
OUTPUT_PATH = DOCS_DIR / "guida-piattaforma.docx"

FLOW_IMAGE = ASSETS_DIR / "diagramma-flusso-richiesta.png"
ROLES_IMAGE = ASSETS_DIR / "diagramma-divisione-compiti.png"

INK = "0B2545"
BLUE = "2E74B5"
TEAL = "147D8C"
LIGHT_BLUE = "E8F1F8"
LIGHT_TEAL = "E7F4F3"
LIGHT_GRAY = "F4F6F9"
MID_GRAY = "667085"
GRID = "CAD3DD"
GOLD = "C58B2A"
GREEN = "2C6E49"
GREEN_FILL = "E9F5EE"
AMBER = "8A5A00"
AMBER_FILL = "FFF4D6"
WHITE = "FFFFFF"

TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


# The descriptions are intentionally written for a non-technical PM. The ID set is
# validated against the actual formula-pack files before the document is built.
CATALOG = {
    "legal_it.imu": (
        "Fiscale",
        "IMU da rendita catastale, moltiplicatore e aliquota deliberata dal comune.",
    ),
    "legal_it.irpef": (
        "Fiscale",
        "IRPEF nazionale lorda per scaglioni; non include detrazioni, addizionali o regimi speciali.",
    ),
    "legal_it.ravvedimento_operoso": (
        "Fiscale",
        "Totale del ravvedimento: tributo, sanzione ridotta e interessi legali.",
    ),
    "legal_it.registration_tax_leases": (
        "Fiscale",
        "Imposta di registro di una locazione abitativa in regime ordinario.",
    ),
    "legal_it.inps_contributions": (
        "Lavoro",
        "Contributi INPS su una retribuzione lorda, separando quota lavoratore e quota datore.",
    ),
    "legal_it.notice_indemnity": (
        "Lavoro",
        "Indennità per mancato preavviso da retribuzione mensile e mesi previsti dal CCNL.",
    ),
    "legal_it.tfr": (
        "Lavoro",
        "Quota annua di TFR e rivalutazione del fondo già accantonato.",
    ),
    "legal_it.compensi_dm55": (
        "Contenzioso / giudiziario",
        "Compensi dell’avvocato per fasi e valore della causa, con spese, CPA e IVA. Gli scaglioni DM 55 disponibili sono sintetici.",
    ),
    "legal_it.contributo_unificato_civile": (
        "Contenzioso / giudiziario",
        "Contributo unificato per un processo civile ordinario, in base al valore e al grado.",
    ),
    "legal_it.late_payment_interest": (
        "Contenzioso / giudiziario",
        "Interessi di mora sui pagamenti commerciali in ritardo, usando il tasso BCE fornito.",
    ),
    "legal_it.legal_interest": (
        "Contenzioso / giudiziario",
        "Interessi legali semplici su un capitale e un periodo, anche quando il tasso cambia.",
    ),
    "legal_it.rivalutazione_interessi_1712": (
        "Contenzioso / giudiziario",
        "Rivalutazione e interessi su debiti di valore secondo Cass. SS.UU. 1712/1995. I dati FOI attuali sono sintetici.",
    ),
    "legal_it.rivalutazione_istat": (
        "Contenzioso / giudiziario",
        "Rivalutazione monetaria tra due date con indice ISTAT FOI. I dati FOI attuali sono sintetici.",
    ),
    "legal_it.termini_processuali_civili": (
        "Contenzioso / giudiziario",
        "Scadenza di termini processuali civili a giorni, considerando ferie e festività.",
    ),
    "legal_it.furto_pena_draft": (
        "Penale",
        "Forbice di pena per furto semplice, con circostanze e rito abbreviato. Non è una previsione della sentenza.",
    ),
    "legal_it.omicidio_pena_draft": (
        "Penale",
        "Forbice di pena per omicidio volontario, con circostanze e rito abbreviato. Non è una previsione della sentenza.",
    ),
    "legal_it.rapina_pena_draft": (
        "Penale",
        "Forbice di pena per rapina semplice, con circostanze e rito abbreviato. Non è una previsione della sentenza.",
    ),
    "business.invoice_total": (
        "Business",
        "Totale fattura a partire da imponibile, sconto e aliquota IVA.",
    ),
    "business.loan_payment": (
        "Business",
        "Rata mensile fissa di un finanziamento ammortizzato.",
    ),
}

AREA_ORDER = {
    "Fiscale": 0,
    "Lavoro": 1,
    "Contenzioso / giudiziario": 2,
    "Penale": 3,
    "Business": 4,
}

SYNTHETIC_DATA_CALCULATORS = {
    "legal_it.compensi_dm55",
    "legal_it.rivalutazione_interessi_1712",
    "legal_it.rivalutazione_istat",
}


def load_yaml(path: Path):
    return yaml.safe_load(path.read_text(encoding="utf-8"))


def load_source_data():
    packs = {}
    for path in sorted(FORMULA_DIR.rglob("*.yml")):
        data = load_yaml(path)
        packs[data["id"]] = data

    if set(packs) != set(CATALOG):
        missing = sorted(set(packs) - set(CATALOG))
        stale = sorted(set(CATALOG) - set(packs))
        raise RuntimeError(
            "Il catalogo PM non corrisponde ai formula pack. "
            f"Nuovi/non mappati: {missing}; mancanti nei file: {stale}"
        )

    draft_ids = {pack_id for pack_id, pack in packs.items() if "draft" in str(pack.get("version", "")).lower()}
    expected_drafts = {
        "legal_it.furto_pena_draft",
        "legal_it.omicidio_pena_draft",
        "legal_it.rapina_pena_draft",
    }
    if draft_ids != expected_drafts:
        raise RuntimeError(
            f"Stato bozze inatteso nei formula pack: trovato {sorted(draft_ids)}"
        )

    parameters = {}
    for path in sorted(PARAMETER_DIR.rglob("*.yml")):
        parameters[path.name] = load_yaml(path)

    dm55_entry = parameters["dm55_compensi.yml"]["values"][0]
    foi_entries = parameters["foi_indices.yml"]["values"]
    if not (
        dm55_entry.get("placeholder") is True
        and dm55_entry.get("verified") is False
        and all(item.get("placeholder") is True and item.get("verified") is False for item in foi_entries)
    ):
        raise RuntimeError("I dati sintetici FOI/DM55 non hanno più lo stato atteso; rivedere la guida.")

    return packs, parameters, draft_ids


def years_covered(entries):
    years = set()
    for item in entries:
        start = item.get("effective_from")
        end = item.get("effective_to")
        if start is None:
            continue
        start_year = start.year if hasattr(start, "year") else int(str(start)[:4])
        end_year = (
            end.year if hasattr(end, "year") else int(str(end)[:4])
        ) if end is not None else start_year
        years.update(range(start_year, end_year + 1))
    return sorted(years)


def italian_date(value: date) -> str:
    months = [
        "gennaio",
        "febbraio",
        "marzo",
        "aprile",
        "maggio",
        "giugno",
        "luglio",
        "agosto",
        "settembre",
        "ottobre",
        "novembre",
        "dicembre",
    ]
    return f"{value.day} {months[value.month - 1]} {value.year}"


def format_years(years):
    return "–".join(str(year) for year in (years[0], years[-1])) if len(years) > 1 else str(years[0])


def format_foi_months(entries):
    months = [
        "gen",
        "feb",
        "mar",
        "apr",
        "mag",
        "giu",
        "lug",
        "ago",
        "set",
        "ott",
        "nov",
        "dic",
    ]
    return ", ".join(f"{months[int(item['month']) - 1]} {item['year']}" for item in entries)


def _run_swift_diagram_builder():
    swift = shutil.which("swift")
    if not swift:
        return False

    cache_dir = ASSETS_DIR / ".swift-diagram-cache"
    cache_dir.mkdir(parents=True, exist_ok=True)
    swift_source = r'''
import AppKit

let flowOutput = CommandLine.arguments[1]
let rolesOutput = CommandLine.arguments[2]

func color(_ hex: String) -> NSColor {
    let clean = hex.replacingOccurrences(of: "#", with: "")
    let value = UInt64(clean, radix: 16)!
    return NSColor(
        calibratedRed: CGFloat((value >> 16) & 255) / 255.0,
        green: CGFloat((value >> 8) & 255) / 255.0,
        blue: CGFloat(value & 255) / 255.0,
        alpha: 1
    )
}

func canvas(width: Int, height: Int) -> NSBitmapImageRep {
    let rep = NSBitmapImageRep(
        bitmapDataPlanes: nil,
        pixelsWide: width,
        pixelsHigh: height,
        bitsPerSample: 8,
        samplesPerPixel: 4,
        hasAlpha: true,
        isPlanar: false,
        colorSpaceName: .deviceRGB,
        bytesPerRow: 0,
        bitsPerPixel: 0
    )!
    NSGraphicsContext.current = NSGraphicsContext(bitmapImageRep: rep)
    NSGraphicsContext.current?.imageInterpolation = .high
    color("#FFFFFF").setFill()
    NSRect(x: 0, y: 0, width: width, height: height).fill()
    return rep
}

func roundedBox(_ rect: NSRect, fill: NSColor, stroke: NSColor, radius: CGFloat = 26, lineWidth: CGFloat = 4) {
    let path = NSBezierPath(roundedRect: rect, xRadius: radius, yRadius: radius)
    fill.setFill()
    path.fill()
    stroke.setStroke()
    path.lineWidth = lineWidth
    path.stroke()
}

func drawText(
    _ text: String,
    in rect: NSRect,
    size: CGFloat,
    color textColor: NSColor,
    weight: NSFont.Weight = .regular,
    alignment: NSTextAlignment = .center,
    lineSpacing: CGFloat = 7
) {
    let paragraph = NSMutableParagraphStyle()
    paragraph.alignment = alignment
    paragraph.lineSpacing = lineSpacing
    let attributes: [NSAttributedString.Key: Any] = [
        .font: NSFont.systemFont(ofSize: size, weight: weight),
        .foregroundColor: textColor,
        .paragraphStyle: paragraph
    ]
    let attributed = NSAttributedString(string: text, attributes: attributes)
    let bounds = attributed.boundingRect(
        with: NSSize(width: rect.width, height: 1000),
        options: [.usesLineFragmentOrigin, .usesFontLeading]
    )
    let y = rect.minY + max(0, (rect.height - bounds.height) / 2)
    attributed.draw(
        with: NSRect(x: rect.minX, y: y, width: rect.width, height: bounds.height + 4),
        options: [.usesLineFragmentOrigin, .usesFontLeading]
    )
}

func arrow(points: [NSPoint], color strokeColor: NSColor = color("#667085"), width: CGFloat = 7) {
    guard points.count >= 2 else { return }
    let path = NSBezierPath()
    path.move(to: points[0])
    for point in points.dropFirst() {
        path.line(to: point)
    }
    strokeColor.setStroke()
    path.lineWidth = width
    path.lineCapStyle = .round
    path.lineJoinStyle = .round
    path.stroke()

    let end = points[points.count - 1]
    let previous = points[points.count - 2]
    let angle = atan2(end.y - previous.y, end.x - previous.x)
    let head: CGFloat = 25
    let wing: CGFloat = 0.55
    let triangle = NSBezierPath()
    triangle.move(to: end)
    triangle.line(to: NSPoint(
        x: end.x - head * cos(angle - wing),
        y: end.y - head * sin(angle - wing)
    ))
    triangle.line(to: NSPoint(
        x: end.x - head * cos(angle + wing),
        y: end.y - head * sin(angle + wing)
    ))
    triangle.close()
    strokeColor.setFill()
    triangle.fill()
}

func save(_ rep: NSBitmapImageRep, to path: String) {
    let data = rep.representation(using: .png, properties: [:])!
    try! data.write(to: URL(fileURLWithPath: path))
}

func buildFlow() {
    let rep = canvas(width: 1800, height: 1260)
    let navy = color("#0B2545")
    let blue = color("#2E74B5")
    let teal = color("#147D8C")
    let gray = color("#667085")
    let lightBlue = color("#E8F1F8")
    let lightTeal = color("#E7F4F3")
    let lightGray = color("#F4F6F9")

    let question = NSRect(x: 590, y: 1080, width: 620, height: 125)
    roundedBox(question, fill: navy, stroke: navy)
    drawText("DOMANDA DELL’UTENTE", in: question.insetBy(dx: 25, dy: 12), size: 38, color: .white, weight: .semibold)

    let decision = NSRect(x: 470, y: 820, width: 860, height: 165)
    arrow(points: [NSPoint(x: 900, y: 1080), NSPoint(x: 900, y: 985)])
    roundedBox(decision, fill: lightBlue, stroke: blue)
    drawText("Il sistema riconosce:\nè una richiesta di calcolo?", in: decision.insetBy(dx: 32, dy: 16), size: 38, color: navy, weight: .semibold)

    let normal = NSRect(x: 70, y: 430, width: 590, height: 170)
    arrow(points: [
        NSPoint(x: 650, y: 820),
        NSPoint(x: 650, y: 700),
        NSPoint(x: 365, y: 700),
        NSPoint(x: 365, y: 600)
    ])
    drawText("NO", in: NSRect(x: 455, y: 620, width: 120, height: 55), size: 30, color: gray, weight: .bold)
    roundedBox(normal, fill: lightGray, stroke: gray)
    drawText("RISPOSTA NORMALE\nDEL CHATBOT", in: normal.insetBy(dx: 30, dy: 20), size: 36, color: navy, weight: .semibold)

    let ai = NSRect(x: 1080, y: 600, width: 640, height: 170)
    arrow(points: [
        NSPoint(x: 1200, y: 820),
        NSPoint(x: 1200, y: 795),
        NSPoint(x: 1400, y: 795),
        NSPoint(x: 1400, y: 770)
    ])
    roundedBox(ai, fill: lightBlue, stroke: blue)
    drawText("INTELLIGENZA ARTIFICIALE\nEstrae dal testo i dati necessari", in: ai.insetBy(dx: 30, dy: 18), size: 32, color: navy, weight: .semibold)
    let yesBadge = NSRect(x: 1245, y: 774, width: 100, height: 42)
    roundedBox(yesBadge, fill: .white, stroke: blue, radius: 14, lineWidth: 2)
    drawText("SÌ", in: yesBadge, size: 24, color: blue, weight: .bold, lineSpacing: 0)

    let engine = NSRect(x: 1080, y: 330, width: 640, height: 170)
    arrow(points: [NSPoint(x: 1400, y: 600), NSPoint(x: 1400, y: 500)])
    roundedBox(engine, fill: navy, stroke: navy)
    drawText("MOTORE DI CALCOLO\nApplica regole e parametri", in: engine.insetBy(dx: 30, dy: 18), size: 33, color: .white, weight: .semibold)

    let result = NSRect(x: 1010, y: 55, width: 710, height: 180)
    arrow(points: [NSPoint(x: 1400, y: 330), NSPoint(x: 1400, y: 235)])
    roundedBox(result, fill: lightTeal, stroke: teal)
    drawText("RISULTATO\nImporto + Fonti + Avvisi", in: result.insetBy(dx: 30, dy: 20), size: 38, color: navy, weight: .bold)

    save(rep, to: flowOutput)
}

func buildRoles() {
    let rep = canvas(width: 1800, height: 700)
    let navy = color("#0B2545")
    let blue = color("#2E74B5")
    let teal = color("#147D8C")
    let lightBlue = color("#E8F1F8")
    let lightTeal = color("#E7F4F3")
    let lightGray = color("#F4F6F9")
    let gold = color("#C58B2A")

    let left = NSRect(x: 50, y: 180, width: 520, height: 420)
    let middle = NSRect(x: 640, y: 180, width: 520, height: 420)
    let right = NSRect(x: 1230, y: 180, width: 520, height: 420)

    roundedBox(left, fill: lightBlue, stroke: blue)
    drawText("1", in: NSRect(x: 78, y: 520, width: 65, height: 55), size: 30, color: blue, weight: .bold)
    drawText("INTELLIGENZA\nARTIFICIALE", in: NSRect(x: 95, y: 395, width: 430, height: 125), size: 34, color: navy, weight: .bold)
    drawText("Capisce la richiesta\ne raccoglie i dati", in: NSRect(x: 95, y: 245, width: 430, height: 130), size: 31, color: navy)

    roundedBox(middle, fill: navy, stroke: navy)
    drawText("2", in: NSRect(x: 668, y: 520, width: 65, height: 55), size: 30, color: .white, weight: .bold)
    drawText("MOTORE DI\nCALCOLO", in: NSRect(x: 685, y: 395, width: 430, height: 125), size: 34, color: .white, weight: .bold)
    drawText("Esegue l’aritmetica\nesatta e verificabile", in: NSRect(x: 685, y: 245, width: 430, height: 130), size: 31, color: .white)

    roundedBox(right, fill: lightTeal, stroke: teal)
    drawText("3", in: NSRect(x: 1258, y: 520, width: 65, height: 55), size: 30, color: teal, weight: .bold)
    drawText("FONTI\nNORMATIVE", in: NSRect(x: 1275, y: 395, width: 430, height: 125), size: 34, color: navy, weight: .bold)
    drawText("Ogni risultato regolato\nriporta le fonti", in: NSRect(x: 1275, y: 245, width: 430, height: 130), size: 31, color: navy)

    let banner = NSRect(x: 260, y: 45, width: 1280, height: 82)
    roundedBox(banner, fill: lightGray, stroke: gold, radius: 20, lineWidth: 3)
    drawText("L’IA NON INVENTA I NUMERI: LI AFFIDA AL MOTORE.", in: banner.insetBy(dx: 20, dy: 8), size: 29, color: navy, weight: .bold)

    save(rep, to: rolesOutput)
}

buildFlow()
buildRoles()
'''
    env = dict(**__import__("os").environ)
    env["CLANG_MODULE_CACHE_PATH"] = str(cache_dir)
    env["SWIFT_MODULECACHE_PATH"] = str(cache_dir)
    env["TMPDIR"] = str(cache_dir)
    try:
        result = subprocess.run(
            [swift, "-", str(FLOW_IMAGE), str(ROLES_IMAGE)],
            input=swift_source,
            text=True,
            cwd=DOCS_DIR,
            env=env,
            capture_output=True,
            timeout=180,
            check=False,
        )
        if result.returncode != 0:
            print("Swift/AppKit non ha generato i diagrammi:")
            print(result.stdout)
            print(result.stderr)
            return False
        return FLOW_IMAGE.exists() and ROLES_IMAGE.exists()
    finally:
        shutil.rmtree(cache_dir, ignore_errors=True)


def _run_dot_diagram_builder():
    dot = shutil.which("dot")
    if not dot:
        return False
    flow_dot = ASSETS_DIR / ".diagramma-flusso.dot"
    roles_dot = ASSETS_DIR / ".diagramma-ruoli.dot"
    flow_dot.write_text(
        textwrap.dedent(
            """
            digraph flow {
              graph [rankdir=TB, bgcolor="white", pad="0.25", nodesep="0.45", ranksep="0.65"];
              node [shape=box, style="rounded,filled", fontname="Arial", fontsize=22, margin="0.22,0.16", color="#2E74B5", fillcolor="#E8F1F8", fontcolor="#0B2545", penwidth=2];
              edge [fontname="Arial", fontsize=18, color="#667085", penwidth=2, arrowsize=0.9];
              q [label="Domanda dell’utente", fillcolor="#0B2545", fontcolor="white", color="#0B2545"];
              d [label="Il sistema riconosce:\\nè una richiesta di calcolo?"];
              normal [label="Risposta normale\\ndel chatbot", fillcolor="#F4F6F9", color="#667085"];
              ai [label="Intelligenza Artificiale\\nEstrae i dati dal testo"];
              engine [label="Motore di Calcolo\\nApplica regole e parametri", fillcolor="#0B2545", fontcolor="white", color="#0B2545"];
              result [label="Risultato\\nImporto + Fonti + Avvisi", fillcolor="#E7F4F3", color="#147D8C"];
              q -> d;
              d -> normal [label="NO"];
              d -> ai [label="SÌ"];
              ai -> engine -> result;
            }
            """
        ).strip(),
        encoding="utf-8",
    )
    roles_dot.write_text(
        textwrap.dedent(
            """
            digraph roles {
              graph [rankdir=LR, bgcolor="white", pad="0.25", nodesep="0.35"];
              node [shape=box, style="rounded,filled", fontname="Arial", fontsize=20, margin="0.22,0.18", color="#2E74B5", fillcolor="#E8F1F8", fontcolor="#0B2545", penwidth=2];
              edge [style=invis];
              ai [label="INTELLIGENZA ARTIFICIALE\\nCapisce ed estrae i dati"];
              engine [label="MOTORE DI CALCOLO\\nAritmetica esatta e verificabile", fillcolor="#0B2545", fontcolor="white", color="#0B2545"];
              sources [label="FONTI NORMATIVE\\nOgni risultato regolato è citato", fillcolor="#E7F4F3", color="#147D8C"];
              ai -> engine -> sources;
            }
            """
        ).strip(),
        encoding="utf-8",
    )
    try:
        subprocess.run([dot, "-Tpng", "-Gdpi=180", str(flow_dot), "-o", str(FLOW_IMAGE)], check=True)
        subprocess.run([dot, "-Tpng", "-Gdpi=180", str(roles_dot), "-o", str(ROLES_IMAGE)], check=True)
        return FLOW_IMAGE.exists() and ROLES_IMAGE.exists()
    except subprocess.CalledProcessError:
        return False
    finally:
        flow_dot.unlink(missing_ok=True)
        roles_dot.unlink(missing_ok=True)


def _run_matplotlib_diagram_builder():
    try:
        import matplotlib.pyplot as plt
        from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
    except ImportError:
        return False

    def box(ax, xy, wh, text, fill, edge, text_color=INK, fontsize=14):
        x, y = xy
        width, height = wh
        ax.add_patch(
            FancyBboxPatch(
                (x, y),
                width,
                height,
                boxstyle="round,pad=0.02,rounding_size=0.03",
                linewidth=2,
                edgecolor=f"#{edge}",
                facecolor=f"#{fill}",
            )
        )
        ax.text(
            x + width / 2,
            y + height / 2,
            text,
            ha="center",
            va="center",
            fontsize=fontsize,
            color=f"#{text_color}",
            weight="semibold",
        )

    fig, ax = plt.subplots(figsize=(12, 8))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    box(ax, (0.33, 0.82), (0.34, 0.1), "Domanda dell’utente", INK, INK, WHITE)
    box(ax, (0.28, 0.61), (0.44, 0.12), "È una richiesta di calcolo?", LIGHT_BLUE, BLUE)
    box(ax, (0.04, 0.31), (0.30, 0.11), "Risposta normale del chatbot", LIGHT_GRAY, MID_GRAY)
    box(ax, (0.63, 0.42), (0.33, 0.11), "IA: estrae i dati", LIGHT_BLUE, BLUE)
    box(ax, (0.63, 0.24), (0.33, 0.11), "Motore: calcola", INK, INK, WHITE)
    box(ax, (0.60, 0.05), (0.36, 0.12), "Risultato + Fonti + Avvisi", LIGHT_TEAL, TEAL)
    arrows = [
        ((0.5, 0.82), (0.5, 0.73)),
        ((0.36, 0.61), (0.2, 0.42)),
        ((0.64, 0.61), (0.79, 0.53)),
        ((0.79, 0.42), (0.79, 0.35)),
        ((0.79, 0.24), (0.79, 0.17)),
    ]
    for start, end in arrows:
        ax.add_patch(FancyArrowPatch(start, end, arrowstyle="-|>", mutation_scale=18, linewidth=2, color=f"#{MID_GRAY}"))
    fig.savefig(FLOW_IMAGE, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(12, 4.5))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    box(ax, (0.02, 0.25), (0.29, 0.55), "INTELLIGENZA ARTIFICIALE\nCapisce ed estrae i dati", LIGHT_BLUE, BLUE, fontsize=13)
    box(ax, (0.355, 0.25), (0.29, 0.55), "MOTORE DI CALCOLO\nAritmetica esatta e verificabile", INK, INK, WHITE, fontsize=13)
    box(ax, (0.69, 0.25), (0.29, 0.55), "FONTI NORMATIVE\nOgni risultato regolato è citato", LIGHT_TEAL, TEAL, fontsize=13)
    ax.text(0.5, 0.08, "L’IA NON INVENTA I NUMERI: LI AFFIDA AL MOTORE.", ha="center", va="center", fontsize=13, weight="bold", color=f"#{INK}")
    fig.savefig(ROLES_IMAGE, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return FLOW_IMAGE.exists() and ROLES_IMAGE.exists()


def generate_diagrams():
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    FLOW_IMAGE.unlink(missing_ok=True)
    ROLES_IMAGE.unlink(missing_ok=True)

    if _run_dot_diagram_builder():
        print("Diagrammi generati con Graphviz/dot.")
        return
    print("Graphviz/dot non disponibile; provo matplotlib.")

    if _run_matplotlib_diagram_builder():
        print("Diagrammi generati con matplotlib.")
        return
    print("Matplotlib non disponibile; uso Swift/AppKit presente nel sistema.")

    if _run_swift_diagram_builder():
        print("Diagrammi generati con Swift/AppKit.")
        return

    raise RuntimeError(
        "Impossibile generare immagini diagramma: Graphviz/dot, matplotlib e Swift/AppKit non disponibili."
    )


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, **margins):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin in ["top", "start", "bottom", "end"]:
        if margin in margins:
            node = tc_mar.find(qn(f"w:{margin}"))
            if node is None:
                node = OxmlElement(f"w:{margin}")
                tc_mar.append(node)
            node.set(qn("w:w"), str(margins[margin]))
            node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        if edge not in edges:
            continue
        data = edges[edge]
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in data:
                element.set(qn(f"w:{key}"), str(data[key]))


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"Larghezze tabella non valide: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Twips(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell, **CELL_MARGINS)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name, size, color=None, bold=None, italic=None):
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def add_custom_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = parse_xml(
        f"""
        <w:abstractNum {nsdecls('w')} w:abstractNumId="{abstract_id}">
          <w:multiLevelType w:val="singleLevel"/>
          <w:lvl w:ilvl="0">
            <w:start w:val="1"/>
            <w:numFmt w:val="bullet"/>
            <w:lvlText w:val="•"/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
              <w:tabs><w:tab w:val="num" w:pos="270"/></w:tabs>
              <w:ind w:left="540" w:hanging="270"/>
              <w:spacing w:before="0" w:after="80" w:line="300" w:lineRule="auto"/>
            </w:pPr>
            <w:rPr>
              <w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>
              <w:color w:val="{BLUE}"/>
            </w:rPr>
          </w:lvl>
        </w:abstractNum>
        """
    )
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_bullet(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, "Calibri", 11, "222222")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    set_style_font(title, "Calibri", 30, INK, bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.line_spacing = 1.05

    subtitle = styles["Subtitle"]
    set_style_font(subtitle, "Calibri", 14, MID_GRAY)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.paragraph_format.line_spacing = 1.15

    h1 = styles["Heading 1"]
    set_style_font(h1, "Calibri", 16, BLUE, bold=True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing = 1.0
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.keep_together = True

    h2 = styles["Heading 2"]
    set_style_font(h2, "Calibri", 13, BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    set_style_font(h3, "Calibri", 12, INK, bold=True)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    set_style_font(caption, "Calibri", 9, MID_GRAY, italic=True)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(5)
    caption.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("PIATTAFORMA DI CALCOLO  |  GUIDA PM")
    set_run_font(r, size=8.5, color=MID_GRAY, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    label = fp.add_run("PAGINA ")
    set_run_font(label, size=8.5, color=MID_GRAY)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run = fp.add_run()
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, size=8.5, color=MID_GRAY)


def add_cover(doc):
    for _ in range(3):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(24)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("GUIDA ALLA PIATTAFORMA")
    set_run_font(run, size=10.5, color=GOLD, bold=True)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Piattaforma di Calcolo\n— Guida introduttiva")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Cosa fa, come protegge l’affidabilità dei numeri e quali limiti governare."
    )

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.space_before = Pt(16)
    rule.paragraph_format.space_after = Pt(34)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), GOLD)
    border.append(bottom)
    rule._p.get_or_add_pPr().append(border)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_after = Pt(4)
    date_run = date_p.add_run(italian_date(date.today()))
    set_run_font(date_run, size=11.5, color=INK, bold=True)

    edition = doc.add_paragraph()
    edition.alignment = WD_ALIGN_PARAGRAPH.CENTER
    edition_run = edition.add_run("Panoramica per stakeholder non tecnici")
    set_run_font(edition_run, size=9.5, color=MID_GRAY, italic=True)

    doc.add_page_break()


def add_summary_box(doc, bullet_num_id, total, ready, draft):
    heading = doc.add_heading("In breve", level=1)
    heading.paragraph_format.space_before = Pt(0)

    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_borders(
        cell,
        top={"val": "single", "sz": "8", "color": GRID},
        bottom={"val": "single", "sz": "8", "color": GRID},
        start={"val": "single", "sz": "20", "color": BLUE},
        end={"val": "single", "sz": "8", "color": GRID},
    )
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    lead = cell.paragraphs[0]
    lead.paragraph_format.space_after = Pt(8)
    lead_run = lead.add_run("La piattaforma in una pagina")
    set_run_font(lead_run, size=11.5, color=INK, bold=True)

    bullets = [
        "È un motore deterministico per calcoli legali, fiscali, di lavoro, giudiziari e business.",
        "Serve a trasformare richieste espresse in linguaggio naturale in risultati spiegabili e ripetibili.",
        "È affidabile perché separa l’interpretazione dell’IA dall’aritmetica eseguita dal motore.",
        "Rende visibili fonti, assunzioni e avvisi: il limite del risultato non resta nascosto.",
        f"Stato attuale: {total} calcolatori a catalogo; {ready} Pronti e {draft} in Bozza per dati sintetici o validazione ancora mancante.",
    ]
    for text in bullets:
        p = cell.add_paragraph()
        apply_bullet(p, bullet_num_id)
        p.add_run(text)

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0)


def add_body_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first, rest = text[: len(bold_prefix)], text[len(bold_prefix) :]
        run = p.add_run(first)
        run.bold = True
        p.add_run(rest)
    else:
        p.add_run(text)
    return p


def add_diagram(doc, image_path, caption_text, alt_text, width=6.3):
    caption = doc.add_paragraph(caption_text, style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    shape = p.add_run().add_picture(str(image_path), width=Inches(width))
    shape._inline.docPr.set("descr", alt_text)


def add_catalog_table(doc, packs, draft_ids):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [1800, 6180, 1380]
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    headers = ["Area", "Che cosa calcola", "Stato"]
    for cell, label in zip(header.cells, headers):
        cell.text = label
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, INK)
        set_cell_borders(
            cell,
            top={"val": "single", "sz": "8", "color": INK},
            bottom={"val": "single", "sz": "8", "color": INK},
            start={"val": "single", "sz": "8", "color": INK},
            end={"val": "single", "sz": "8", "color": INK},
        )
        paragraph = cell.paragraphs[0]
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if label == "Stato" else WD_ALIGN_PARAGRAPH.LEFT
        )
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.runs[0]
        set_run_font(run, size=9.5, color=WHITE, bold=True)

    ordered = sorted(
        CATALOG.items(),
        key=lambda item: (
            AREA_ORDER[item[1][0]],
            packs[item[0]]["name"].lower(),
        ),
    )
    bozza_ids = draft_ids | SYNTHETIC_DATA_CALCULATORS
    area_fills = {
        "Fiscale": "EEF4FA",
        "Lavoro": "EEF7F2",
        "Contenzioso / giudiziario": "F5F2FA",
        "Penale": "FFF2EF",
        "Business": "FFF8E8",
    }

    for calculator_id, (area, description) in ordered:
        row = table.add_row()
        set_row_cant_split(row)
        status = "Bozza" if calculator_id in bozza_ids else "Pronto"
        values = [area, description, status]
        for index, (cell, value) in enumerate(zip(row.cells, values)):
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(
                cell,
                top={"val": "single", "sz": "5", "color": GRID},
                bottom={"val": "single", "sz": "5", "color": GRID},
                start={"val": "single", "sz": "5", "color": GRID},
                end={"val": "single", "sz": "5", "color": GRID},
            )
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.12
            if index == 2:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0]
            set_run_font(run, size=9.1, color=INK, bold=(index in {0, 2}))

        set_cell_shading(row.cells[0], area_fills[area])
        if status == "Pronto":
            set_cell_shading(row.cells[2], GREEN_FILL)
            row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(GREEN)
        else:
            set_cell_shading(row.cells[2], AMBER_FILL)
            row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(AMBER)

    set_table_geometry(table, widths)
    return table


def add_definition(doc, term, definition):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    label = p.add_run(f"{term}: ")
    set_run_font(label, color=INK, bold=True)
    p.add_run(definition)


def build_document():
    packs, parameters, draft_ids = load_source_data()
    generate_diagrams()

    bozza_ids = draft_ids | SYNTHETIC_DATA_CALCULATORS
    total = len(packs)
    draft = len(bozza_ids)
    ready = total - draft

    legal_interest_years = years_covered(parameters["legal_interest_rates.yml"]["values"])
    irpef_years = years_covered(parameters["irpef_brackets.yml"]["values"])
    foi_entries = parameters["foi_indices.yml"]["values"]
    holidays = parameters["giorni_festivi.yml"]["values"][0]["value"]

    doc = Document()
    configure_document(doc)
    bullet_num_id = add_custom_bullet_numbering(doc)

    doc.core_properties.title = "Piattaforma di Calcolo — Guida introduttiva"
    doc.core_properties.subject = "Guida non tecnica alla piattaforma di calcolo"
    doc.core_properties.author = "Calculation Platform"
    doc.core_properties.keywords = "calcolo deterministico, project management, fonti normative"

    add_cover(doc)
    add_summary_box(doc, bullet_num_id, total, ready, draft)

    doc.add_heading("Che cos’è", level=1)
    add_body_paragraph(
        doc,
        "La piattaforma è un motore deterministico: calcola importi legali e fiscali — per esempio imposte, interessi, sanzioni e contributi — oltre ad alcune operazioni di lavoro, giudiziarie e business. “Deterministico” significa che, a parità di dati e regole applicabili, il numero ottenuto è sempre lo stesso.",
    )
    add_body_paragraph(
        doc,
        "L’intelligenza artificiale può capire la domanda e riconoscere i dati utili, ma non esegue mai i calcoli. L’aritmetica resta nel motore, dove le regole sono dichiarate e verificabili. Nei calcoli regolati, le fonti normative accompagnano il risultato; le funzioni puramente aritmetiche di business non richiedono una fonte normativa.",
    )

    doc.add_heading("Come funziona una richiesta", level=1)
    add_body_paragraph(
        doc,
        "Il sistema prima distingue una normale domanda da una richiesta di calcolo. Solo nel secondo caso l’IA estrae i dati dal testo e li consegna al motore. Il risultato torna con le fonti applicabili e con eventuali avvisi.",
    )
    add_diagram(
        doc,
        FLOW_IMAGE,
        "Diagramma 1 — Dalla domanda al risultato",
        "Flusso di una richiesta: la domanda viene classificata; una richiesta di calcolo passa dall’estrazione dei dati al motore e produce un risultato con fonti e avvisi, mentre le altre domande ricevono una risposta normale del chatbot.",
        width=6.2,
    )

    doc.add_heading("Chi fa cosa", level=1)
    add_body_paragraph(
        doc,
        "La fiducia nasce dalla separazione dei ruoli. L’IA interpreta; il motore calcola; le fonti documentano le regole applicate. Questo confine evita che un modello linguistico improvvisi un importo.",
    )
    add_diagram(
        doc,
        ROLES_IMAGE,
        "Diagramma 2 — Divisione dei compiti",
        "Tre blocchi affiancati mostrano i compiti dell’intelligenza artificiale, del motore di calcolo e delle fonti normative. Messaggio conclusivo: l’intelligenza artificiale non inventa i numeri.",
        width=6.25,
    )

    catalog_heading = doc.add_heading("Che cosa sa calcolare", level=1)
    catalog_heading.paragraph_format.page_break_before = True
    add_body_paragraph(
        doc,
        "Il catalogo attuale comprende le aree sotto indicate. “Pronto” significa che il calcolatore è disponibile entro il perimetro e le esclusioni dichiarate. “Bozza” segnala invece una validazione ancora mancante oppure dati sintetici non utilizzabili operativamente.",
    )
    add_catalog_table(doc, packs, draft_ids)
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(6)
    note.paragraph_format.space_after = Pt(4)
    run = note.add_run(
        f"Stato complessivo: {total} calcolatori — {ready} Pronti e {draft} in Bozza."
    )
    set_run_font(run, size=9.5, color=MID_GRAY, italic=True)

    doc.add_heading("Affidabilità e trasparenza", level=1)
    add_body_paragraph(
        doc,
        "Un risultato non è presentato come un numero isolato. La piattaforma mostra anche le Assunzioni, cioè le ipotesi usate, e gli Avvisi, cioè i limiti da conoscere prima di decidere. Può quindi dichiarare, per esempio, che l’IRPEF calcolata è solo quella lorda e non comprende detrazioni o addizionali.",
    )
    add_body_paragraph(
        doc,
        "Questa trasparenza è un punto di forza: un limite visibile può essere gestito, verificato o portato a uno specialista. Un limite nascosto, invece, rischia di trasformarsi in una decisione sbagliata.",
    )
    disclaimer = doc.add_paragraph()
    disclaimer.paragraph_format.space_before = Pt(6)
    disclaimer.paragraph_format.space_after = Pt(8)
    disclaimer.paragraph_format.left_indent = Inches(0.18)
    disclaimer.paragraph_format.right_indent = Inches(0.18)
    label = disclaimer.add_run("Nota di utilizzo — ")
    set_run_font(label, color=AMBER, bold=True)
    text = disclaimer.add_run(
        "ogni risultato è una stima indicativa e non sostituisce il parere di un professionista."
    )
    set_run_font(text, color=INK, italic=True)

    doc.add_heading("Stato attuale e limiti", level=1)
    add_body_paragraph(
        doc,
        "La copertura è utile ma non ancora completa. Il comportamento corretto, quando manca un parametro necessario, non è indovinare: il motore rifiuta il calcolo oppure emette un avviso di verifica.",
    )
    limits = [
        f"Indici FOI/ISTAT: sono presenti solo {format_foi_months(foi_entries)}. Tutti questi valori sono sintetici e non sono dati ISTAT reali.",
        "Compensi DM 55: solo lo scaglione da 26.000,01 a 52.000 euro contiene valori medi sintetici; gli altri scaglioni non sono risolti.",
        "Calcolatori penali: furto, omicidio volontario e rapina sono bozze dimostrative non validate legalmente e non prevedono la sentenza.",
        f"Finestre temporali: interessi legali {format_years(legal_interest_years)}; IRPEF {format_years(irpef_years)}; calendario delle festività {holidays['covers_from'][:4]}–{holidays['covers_through'][:4]}; regime del ravvedimento dalle violazioni del 1° settembre 2024.",
        "Manutenzione: le tabelle che cambiano nel tempo devono essere aggiornate e ricontrollate. Le tabelle ufficiali presenti non riportano ancora una data di verifica automatica, quindi il motore aggiunge un avviso.",
    ]
    for text in limits:
        p = doc.add_paragraph()
        apply_bullet(p, bullet_num_id)
        p.add_run(text)

    add_body_paragraph(
        doc,
        "Sul piano operativo il punto è semplice: l’aggiornamento annuale dei parametri va pianificato come attività di governance, non lasciato implicito.",
    )

    doc.add_heading("Glossario minimo", level=1)
    add_definition(
        doc,
        "Deterministico",
        "un sistema che restituisce lo stesso risultato quando riceve gli stessi dati e applica le stesse regole.",
    )
    add_definition(
        doc,
        "Parametro / tabella parametri",
        "un dato esterno alla formula che può cambiare nel tempo, come un tasso, uno scaglione o un indice mensile.",
    )
    add_definition(
        doc,
        "Fonte normativa",
        "il riferimento ufficiale che documenta la regola usata nel calcolo.",
    )
    add_definition(
        doc,
        "Bozza",
        "un calcolatore o un insieme di dati non ancora pronto per l’uso operativo.",
    )
    add_definition(
        doc,
        "Avviso",
        "un limite, un’esclusione o una cautela da leggere insieme al risultato.",
    )

    doc.save(OUTPUT_PATH)

    reopened = Document(OUTPUT_PATH)
    headings = [
        paragraph.text
        for paragraph in reopened.paragraphs
        if paragraph.style and paragraph.style.name.startswith("Heading")
    ]
    print(f"DOCX: {OUTPUT_PATH}")
    print(f"Paragrafi: {len(reopened.paragraphs)}")
    print(f"Intestazioni: {headings}")
    print(f"Tabelle: {len(reopened.tables)}")
    print(f"Immagini incorporate: {len(reopened.inline_shapes)}")
    print(f"Asset diagramma 1: {FLOW_IMAGE.exists()} ({FLOW_IMAGE})")
    print(f"Asset diagramma 2: {ROLES_IMAGE.exists()} ({ROLES_IMAGE})")


if __name__ == "__main__":
    build_document()
